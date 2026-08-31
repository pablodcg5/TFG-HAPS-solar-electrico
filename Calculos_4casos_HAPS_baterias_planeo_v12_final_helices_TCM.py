# -*- coding: utf-8 -*-
"""
Dimensionado HAPS: comparación de 4 casos completos
- Geometría / aerodinámica en crucero
- Potencia propulsiva real
- Captación solar diaria (modelo tipo Gundlach + Hall digitizado)
- Balance energético diario simple
- Balance detallado de almacenamiento para baterías (1, 30 y 60 días)
- Exportación opcional a CSV / TSV / HTML / DOCX (dejada comentada por defecto)
"""

import math
from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd

# -----------------------------------------------------------------------------
# 0. IMPORT OPCIONAL PARA WORD
# -----------------------------------------------------------------------------
DOCX_AVAILABLE = True
try:
    from docx import Document
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt
except ModuleNotFoundError:
    DOCX_AVAILABLE = False

# -----------------------------------------------------------------------------
# 1. CONSTANTES GLOBALES
# -----------------------------------------------------------------------------
g = 9.80665
R = 287.05
PI = math.pi

# Condición general de operación
h_crucero = 20000.0           # m
v_crucero = 23.0              # m/s
lat_deg = 40.0                # deg
lat_rad = math.radians(lat_deg)

# Aerodinámica / arrastre
C_fe = 0.0030
Rel_sups = 2.3
e_S = 0.98
e_fus = 0.99
e_visc = 0.85
C_D0 = 1.15 * (C_fe * Rel_sups)
e = e_S * e_fus * e_visc

# Propulsión (ajústalo si haces sensibilidad)
eta_prop_global = 0.65

# Solar
R_earth_km = 6378.1
pole_tilt_deg = 23.44
pole_tilt_rad = math.radians(pole_tilt_deg)
orbit_ecc = 0.0167
I_mean = 1361.0               # W/m², irradiancia media extraterrestre
h_cloud_km = 0.0
dt_min = 5.0

# Cobertura / eficiencia array
f_cobertura_max = 0.80
S_cola_extra = 0.0            # m², de momento sin cola
eta_cell = 0.22               # Hipótesis propia inicial
eta_covering = 0.95           # Hipótesis propia inicial
eta_wiring = 0.98             # Hipótesis propia inicial
eta_tracker = 0.98            # Hipótesis propia inicial
packing_factor = 0.91         # Cota superior geométrica inicial sobre el área ya definida
eta_array = eta_cell * eta_covering * eta_wiring * eta_tracker * packing_factor


# Consumos de sistemas y payload que ya se meten en los cálculos
"""
HIPÓTESIS INICIALES DE POTENCIA NO PROPULSIVA.

ORIGEN DE LOS VALORES:
- En tu Word aparece explícitamente que el Honeywell VersaWave puede consumir
  hasta 53 W, pero que ese consumo máximo no sería continuo, sino que se usaría
  solo en ciertos momentos del día.
- Para el resto de la aviónica se ha adoptado un presupuesto preliminar coherente
  con el conjunto de equipos que has ido fijando (autopilot, ADS-B, cámaras 360º,
  actuadores, sensores y comunicaciones), de modo que el modelo deje de usar 0 W
  para la parte no propulsiva.
- La payload se representa como una cámara sencilla de misión que solo se enciende
  durante una ventana diurna concreta. Esto es más realista que repartir su consumo
  uniformemente a lo largo de toda la noche, ya que la lógica operativa normal es
  grabar cuando interesa la observación y cuando además hay iluminación solar.

CÓMO SE USA EN EL CÓDIGO:
- P_sistemas_const_W se suma SIEMPRE a la potencia propulsiva, porque representa
  aviónica, comunicaciones, control y demás cargas del bus que deben mantenerse.
- La payload de cámara se aplica mediante un perfil horario, no como potencia media
  constante. Así el cálculo de batería nocturna no se penaliza artificialmente con
  una cámara que en realidad estaría apagada de noche.
"""

P_sistemas_const_W = 60.0

"""
HIPÓTESIS SIMPLE DE PAYLOAD (PL) PARA NO COMPLICAR EL CÓDIGO EN ESTA FASE.

IMPORTANTE:
- Aquí NO se modela una misión concreta de payload ni una ventana horaria de uso.
- Como todavía no tienes una misión definida, la PL se representa simplemente
  como una carga eléctrica constante y editable.
- Esto es una hipótesis de trabajo, NO una ecuación de Gundlach.

POR QUÉ SE HACE ASÍ:
- Tú mismo comentaste que, por ahora, la PL podría ser casi "una piedra" de 5-10 kg.
- Por tanto, complicar el programa con una función temporal de cámara no aporta valor
  en esta fase.
- Lo útil ahora es que el consumo de la PL exista como término simple en el balance
  del bus y que luego puedas cambiarlo fácilmente.

CÓMO SE USA:
- P_payload_const_W se suma siempre a la potencia total del bus, igual que la aviónica.
- Si más adelante defines una misión real, este valor constante podrá sustituirse por
  un perfil temporal más detallado.
"""
P_payload_const_W = 5.0

"""
BLOQUE NUEVO: hipótesis para el dimensionado preliminar de baterías.

ORIGEN CONCEPTUAL:
- Gundlach, Cap. 3, Ecs. (3.36) a (3.39): la energía utilizable de la batería
  se relaciona con la potencia media demandada por la aeronave.
- Gundlach, Cap. 8, Ecs. (8.104) a (8.109): para aeronaves solares multidiurnas
  la batería no se dimensiona solo con un "consumo diario total", sino con el
  balance temporal entre lo que sobra de día y lo que falta de noche / amanecer /
  atardecer.
- Gundlach, Tabla 8.2: aporta rangos de energía específica práctica (Wh/kg) y
  potencia específica (W/kg) para distintas químicas.
- En tus notas del TFG aparece la idea de reservar margen para no descargar la
  batería por completo y alargar su vida. Por eso se deja f_usable_batt como una
  hipótesis explícita y fácil de cambiar.

POR QUÉ SE USA AQUÍ:
- Tu script ya calcula correctamente la potencia solar disponible a lo largo del día.
  El siguiente paso lógico es traducir ese perfil de potencia en un perfil de carga
  y descarga de la batería.
- Para un HAPS solar de 1-2 meses NO tiene sentido multiplicar simplemente la
  energía de una noche por 30 o 60 si el balance diario es no-negativo. En ese caso,
  la batería se recarga cada día y el tamaño del pack lo fija la "oscilación diaria"
  de energía almacenada.
- Solo si el balance diario de almacenamiento es negativo, la capacidad necesaria
  sí crece con los días de misión, porque cada ciclo deja la batería un poco más vacía.
"""

eta_power_conditioning = 1.0
eta_batt_charge = 0.95
eta_batt_discharge = 0.95
f_usable_batt = 0.80
mission_days_list = [1, 30, 60]

"""
BLOQUE NUEVO: tecnologías de batería a comparar.

ORIGEN:
- Gundlach, Tabla 8.2 "Characteristics of Battery Chemistries".
- Se usan directamente los rangos prácticos de energía específica y potencia
  específica publicados en esa tabla.

POR QUÉ SE USA ASÍ:
- En fase conceptual todavía no has seleccionado una celda concreta ni un pack
  comercial concreto. Por eso es más honesto trabajar con rangos conservador /
  optimista de la literatura base que fijar un único número demasiado pronto.
- Para cada química se calcularán dos masas:
    1) la que impone la energía almacenada necesaria,
    2) la que impone la potencia pico que debe entregar la batería,
  y se tomará la mayor de ambas, porque el pack debe cumplir simultáneamente
  ambas restricciones.

NOTA IMPORTANTE:
- Gundlach también advierte que Li-S ofrece gran energía específica práctica,
  pero con limitaciones de ciclos de carga/descarga. Por eso esta química puede
  salir muy bien en masa, pero luego conviene contrastarla con una celda real o
  con el fabricante si quieres cerrar una elección de diseño.
"""

battery_technologies = {
    "Li-Ion": {
        "Espec_min_Whkg": 100.0,
        "Espec_max_Whkg": 135.0,
        "Pspec_min_Wkg": 250.0,
        "Pspec_max_Wkg": 340.0,
    },
    "Li-Po": {
        "Espec_min_Whkg": 50.7,
        "Espec_max_Whkg": 220.0,
        "Pspec_min_Wkg": 200.0,
        "Pspec_max_Wkg": 1900.0,
    },
    "Li-S": {
        "Espec_min_Whkg": 350.0,
        "Espec_max_Whkg": 350.0,
        "Pspec_min_Wkg": 600.0,
        "Pspec_max_Wkg": 700.0,
    },
}


"""
BLOQUE NUEVO: hipótesis de tecnología propulsiva para el punto 8 del sizing.

QUÉ REPRESENTA CADA DATO:
- Pspec_motor_Wkg: potencia específica del motor [W/kg].
- Pspec_controller_Wkg: potencia específica del ESC/controlador [W/kg].
- Pspec_propeller_Wkg: potencia específica equivalente de hélice / integración
  propulsiva [W/kg]. No es un dato "clásico" de catálogo, pero sirve como
  aproximación conceptual cuando se quiere repartir masa del sistema propulsivo.
- f_install: factor multiplicador para capturar soportes, herrajes, cableado,
  instalación mecánica y penalizaciones de integración.

IMPORTANTE:
- Estos valores NO salen del modelo aerodinámico ni del balance solar.
- Son datos de benchmark / tecnología y, por tanto, deben completarse cuando
  revises motores, ESC y hélices reales o literatura comparable.
- Se dejan inicialmente como NaN para que el código no confunda una hipótesis
  pendiente con un valor ya validado.

HIPÓTESIS AUXILIAR:
- eta_propulsor_para_sizing convierte la potencia de empuje calculada por el
  modelo en una potencia de eje de referencia para pesar motor + controlador +
  hélice. Esto es una hipótesis de modelado práctica porque el script actual
  usa una única eta_prop_global agregada.
"""
eta_propulsor_para_sizing = 0.80  # editable; hipótesis provisional para sizing de hardware

propulsion_technologies = {
    "Hobbywing_C6225_ref": {
        "Pspec_motor_Wkg": 580.0 / 0.530,           #potencia nominal de salida del C6225/masa motor
        "Pspec_controller_Wkg": 580.0 / 0.150,      #potencia de eje equivalente/masa ESC aprox.
        "Pspec_propeller_Wkg": 580.0 / 0.200,       #potencia de eje equivalente/masa hélice estimada
        "f_install": 1.20,                          #margen de instalación: soportes, cables, conectores, fijaciones
    },
}


"""
BLOQUE NUEVO: opciones de cierre preliminar de WTO (punto 9 del esquema de Gundlach).

QUÉ SE QUIERE HACER:
- Usar los 8 bloques previos como conjunto coherente y traducirlos a un WTO
  esperado mediante la ecuación de sizing inicial.
- Facilitar una iteración sencilla sin tener que rehacer a mano todas las cuentas.

DOS MODOS DE TRABAJO:
1) fixed_weights_mode = "implied_from_current_mtow"
   - Usa el MTOW actual del caso como punto de anclaje.
   - A partir de MF_Struct, MF_Subs, MF_Energy y un MF_Prop de referencia,
     reconstruye el total de pesos fijos que sería consistente con ese MTOW.
   - Es muy útil cuando todavía no has cerrado en masa W_PL, W_Avion y W_Other.

2) fixed_weights_mode = "manual"
   - Usa directamente masas fijas introducidas a mano por plataforma.
   - Es el modo recomendable cuando ya tengas mejor cerrados payload, aviónica
     y otros pesos no variables.

NOTA IMPORTANTE:
- El cierre de WTO que se añade aquí es un cierre PRELIMINAR de nivel conceptual.
- Todavía no rehace toda la aerodinámica, el solar y la batería en cada iteración.
- Sirve para ordenar el sizing y explorar sensibilidad de MF_Prop de forma rápida.
"""

wto_energy_reference_days = 30
wto_fixed_weights_mode = "implied_from_current_mtow"   # opciones: "implied_from_current_mtow", "manual"
wto_reference_mf_prop_for_implied_fixed = 0.10         # caso medio por defecto para reconstruir W_fixed
wto_secant_tol_kg = 1e-6
wto_secant_max_iter = 50

"""
PESOS FIJOS MANUALES PARA EL CIERRE PRELIMINAR DE WTO.

IMPORTANTE:
- Estos valores SOLO se usan si wto_fixed_weights_mode = "manual".
- Se dejan como ejemplo editable para que puedas activarlos cuando tengas una
  mejor idea de las masas fijas reales.
- Si mantienes el modo "implied_from_current_mtow", estos números no intervienen.

RECUERDA:
- W_PL_kg: payload fija
- W_Avion_kg: aviónica
- W_Other_kg: otros pesos fijos no recogidos en categorías anteriores
"""
fixed_weights_manual = {
    "24m": {"W_PL_kg": 8.0, "W_Avion_kg": 4.0, "W_Other_kg": 3.0},
    "32m": {"W_PL_kg": 12.0, "W_Avion_kg": 5.0, "W_Other_kg": 4.0},
}

"""
CASOS SUPUESTOS DE MF_Prop PARA CUANDO AÚN NO HAY BENCHMARK DE MOTOR/ESC/HÉLICE.

ORIGEN DE LOS VALORES:
- 0.05:
    * sale directamente del Ejemplo 3.1 de Gundlach para el caso de motor
      de dos tiempos, donde MF_Prop = 0.05.
- 0.10:
    * sale del mismo Ejemplo 3.1 para el caso de cuatro tiempos;
    * además es coherente con el Ejemplo 3.2, donde aparece una fracción
      de masa de motor+propeller del orden de 0.1.
- 0.15:
    * NO es un valor impuesto por Gundlach.
    * se introduce aquí como sensibilidad conservadora por encima de los
      ejemplos del libro, para ver cómo se degrada el cierre del WTO si la
      propulsión resulta más pesada de lo esperado.

LECTURA:
- "bajo": escenario optimista / ligero.
- "medio": escenario base razonable mientras no haya benchmark real.
- "conservador": sensibilidad penalizada.
"""
mf_prop_assumed_cases = {
    "Asumido_bajo_0p05": {
        "MF_Prop": 0.05,
        "Origen_MF_Prop": "Gundlach Ej. 3.1 | caso ligero ~0.05",
    },
    "Asumido_medio_0p10": {
        "MF_Prop": 0.10,
        "Origen_MF_Prop": "Gundlach Ej. 3.1 / Ej. 3.2 | caso base ~0.10",
    },
    "Asumido_conservador_0p15": {
        "MF_Prop": 0.15,
        "Origen_MF_Prop": "Sensibilidad conservadora propia > ejemplos Gundlach",
    },
}

# -----------------------------------------------------------------------------
# 2. CASOS DE ESTUDIO
# -----------------------------------------------------------------------------
casos_haps = {
    "Caso_1.1_24m_equinoccio": {
        "W_TO": 85.0,
        "b": 24.0,
        "day_of_year": 79,
        "days_since_vernal": 0,
    },
    "Caso_1.2_24m_invierno": {
        "W_TO": 85.0,
        "b": 24.0,
        "day_of_year": 355,
        "days_since_vernal": 355 - 79,
    },
    "Caso_2.1_32m_equinoccio": {
        "W_TO": 150.0,
        "b": 32.0,
        "day_of_year": 79,
        "days_since_vernal": 0,
    },
    "Caso_2.2_32m_invierno": {
        "W_TO": 150.0,
        "b": 32.0,
        "day_of_year": 355,
        "days_since_vernal": 355 - 79,
    },
}


"""
ESCENARIOS OPERATIVOS A COMPARAR.

1) Crucero constante a 20 km:
   - Replica el planteamiento base que ya venías usando.
   - ROC = 0 durante todo el día.
   - La potencia propulsiva se mantiene aproximadamente constante.

2) Descenso nocturno 20 km -> 15 km y reascenso diurno 15 km -> 20 km:
   - Se inspira directamente en la observación de Gundlach de que, durante la
     noche, una aeronave solar puede actuar como planeador y usar la energía
     potencial acumulada durante el día para reducir o incluso anular la potencia
     propulsiva nocturna.
   - Durante las horas sin aporte solar suficiente, el avión desciende linealmente
     hasta la cota mínima fijada.
   - Durante las horas diurnas útiles de panel horizontal vuelve a ascender
     linealmente hasta la cota alta.
   - Esto NO pretende ser todavía un modelo completo de misión con fase inicial de
     ascenso o despegue; es solo una comparación conceptual entre dos estrategias
     de crucero multidiurno.

LIMITACIÓN IMPORTANTE:
- El modelo solar sigue usando la atenuación de Hall digitalizada para 20 km,
  que es la que ya venías empleando en tu script. Por tanto, en el escenario con
  descenso nocturno, la comparación aísla sobre todo el efecto de cambiar la
  potencia propulsiva requerida, no el posible cambio de captación solar entre
  15 km y 20 km.
"""

escenarios_operativos = {
    "Constante_20km": {
        "tipo": "crucero_constante",
        "h_max_m": 20000.0,
        "h_min_m": 20000.0,
    },
    "DescensoNocturno_20a15km": {
        "tipo": "descenso_nocturno_y_reascenso_diurno",
        "h_max_m": 20000.0,
        "h_min_m": 15000.0,
    },
}

# -----------------------------------------------------------------------------
# 3. FUNCIONES AUXILIARES
# -----------------------------------------------------------------------------
def atmosfera_isa(h):
    if h <= 11000.0:
        T = 288.15 - 0.0065 * h
        rho = 1.225 * (T / 288.15) ** ((g / (R * 0.0065)) - 1)
    else:
        T = 216.65
        rho = 0.3639 * math.exp(-g * (h - 11000.0) / (R * 216.65))
    return rho


def clamp(x, xmin, xmax):
    """Evita problemas numéricos de redondeo en acos/asin."""
    return max(xmin, min(xmax, x))


def rad2deg(x):
    return math.degrees(x)


def solar_day_angle_rad(day_of_year):
    """
    Gundlach Eq. (8.92)
    u = 2*pi*(n - 4)/365
    donde n es el número de días transcurridos desde el 31 de diciembre.
    En la práctica, n se toma como el día del año (Jan 1 -> 1).
    """
    return 2.0 * PI * (day_of_year - 4.0) / 365.0


def extraterrestrial_irradiance(day_of_year):
    """
    Gundlach Eq. (8.91)
    I = I_mean * ((1 + e*cos(u)) / (1 - e^2))^2

    Devuelve la irradiancia solar fuera de la atmósfera terrestre para el día dado.
    """
    u = solar_day_angle_rad(day_of_year)
    return I_mean * ((1.0 + orbit_ecc * math.cos(u)) / (1.0 - orbit_ecc ** 2)) ** 2


def declination_rad(days_since_vernal_equinox):
    """
    Gundlach Eq. (8.94)
    d = i_pole * sin(2*pi*day/365)

    day = días desde el último equinoccio de primavera.
    En Gundlach el equinoccio vernal se toma el día 79 del año (20 de marzo).
    """
    return pole_tilt_rad * math.sin(2.0 * PI * days_since_vernal_equinox / 365.0)


def hour_angle_rad(time_solar_hours):
    """
    Gundlach Eq. (8.95)
    v = pi - 2*pi*time/24

    time en horas desde medianoche solar.
    A las 12:00 solares -> v = 0.
    """
    return PI - 2.0 * PI * time_solar_hours / 24.0


def solar_elevation_rad(lat_rad_local, decl_rad_local, hour_ang_rad_local):
    """
    Gundlach Eq. (8.96)
    El = asin( sin(lat)*sin(d) + cos(lat)*cos(d)*cos(v) )
    """
    arg = (
        math.sin(lat_rad_local) * math.sin(decl_rad_local)
        + math.cos(lat_rad_local) * math.cos(decl_rad_local) * math.cos(hour_ang_rad_local)
    )
    return math.asin(clamp(arg, -1.0, 1.0))


def horizon_elevation_rad(h_uav_km_local, h_cloud_km_local=0.0):
    """
    Gundlach Eq. (8.98)
    El_horiz = asin((R_E + h_cloud)/(R_E + h)) - pi/2

    Sale NEGATIVO porque, a altitud, el horizonte cae por debajo de la horizontal local.
    """
    ratio = (R_earth_km + h_cloud_km_local) / (R_earth_km + h_uav_km_local)
    return math.asin(clamp(ratio, -1.0, 1.0)) - PI / 2.0


def sunrise_hour_angle_rad(lat_rad_local, decl_rad_local, elev_ref_rad):
    """
    Forma general de Gundlach Eq. (8.99)
    v_SR = acos((sin(El_ref) - sin(lat)*sin(d)) / (cos(lat)*cos(d)))

    Si elev_ref_rad = El_horiz  -> amanecer/atardecer GEOMÉTRICO a altitud
    Si elev_ref_rad = 0         -> amanecer/atardecer útil para panel horizontal
                                  (cuando el Sol cruza la horizontal local)
    """
    denom = math.cos(lat_rad_local) * math.cos(decl_rad_local)

    # Casos polares / casi polares (no se esperan en 40ºN, pero por robustez)
    if abs(denom) < 1e-12:
        return None

    arg = (math.sin(elev_ref_rad) - math.sin(lat_rad_local) * math.sin(decl_rad_local)) / denom

    if arg <= -1.0:
        # Sol siempre por encima del elev_ref
        return PI
    if arg >= 1.0:
        # Sol siempre por debajo del elev_ref
        return 0.0

    return math.acos(clamp(arg, -1.0, 1.0))


def sunrise_time_from_hour_angle(v_sr_rad):
    """
    Gundlach Eq. (8.100)
    time_SR = ((pi - v_SR)/(2*pi))*24
    """
    return ((PI - v_sr_rad) / (2.0 * PI)) * 24.0


def sunset_time_from_sunrise(time_sr_hours):
    """
    Gundlach Eq. (8.101)
    time_SS = 24 - time_SR
    """
    return 24.0 - time_sr_hours


def interp_lineal(x, x_pts, y_pts):
    if len(x_pts) != len(y_pts):
        raise ValueError("x_pts e y_pts deben tener la misma longitud")

    if x <= x_pts[0]:
        return y_pts[0]
    if x >= x_pts[-1]:
        return y_pts[-1]

    for i in range(len(x_pts) - 1):
        x0, x1 = x_pts[i], x_pts[i + 1]
        y0, y1 = y_pts[i], y_pts[i + 1]
        if x0 <= x <= x1:
            frac = (x - x0) / (x1 - x0)
            return y0 + frac * (y1 - y0)

    return y_pts[-1]


def c_atten_hall_20km(el_rad):
    """
    Coeficiente de atenuación atmosférica C_atten para h = 20 km.

    ORIGEN:
    - Hall et al., NASA CR-3699, Figure 9:
      "Atmospheric Attenuation Coefficient"
    - La figura da C_a en función de altitud y elevación solar.
    - Estos valores son una digitalización manual de la curva a 20 km.
    - Atmósfera usada por Hall: "mid-latitude, winter"
    - Banda espectral: 0.25 a 4.0 micras
    """
    el_deg = math.degrees(el_rad)
    # Puntos digitizados de Hall Fig. 9 para ALTITUD = 20 km
    el_pts = [0.0, 5.0, 10.0, 15.0, 20.0, 30.0, 45.0, 90.0]
    ca_pts = [0.583, 0.827, 0.890, 0.913, 0.924, 0.939, 0.948, 0.960]

    if el_deg <= 0.0:
        return ca_pts[0]
    return interp_lineal(el_deg, el_pts, ca_pts)


def direct_irradiance_on_horizontal_panel(day_of_year, decl_rad_local, time_solar_h):
    """
    Irradiancia directa sobre un panel horizontal orientado hacia arriba.

    Partimos de Gundlach Eq. (8.93):
        I_tot = C_atten * I * cos(c)

    Para un panel horizontal sobre el extradós:
        cos(c) = sin(El)

    Además, si El < 0, la radiación directa sobre una superficie horizontal
    orientada hacia arriba es nula, por eso se aplica max(0, sin(El)).
    """
    v = hour_angle_rad(time_solar_h)
    el = solar_elevation_rad(lat_rad, decl_rad_local, v)
    i_space = extraterrestrial_irradiance(day_of_year)
    c_atten = c_atten_hall_20km(el)
    cos_inc = max(0.0, math.sin(el))
    i_panel = c_atten * i_space * cos_inc
    return i_panel, el


def integrar_energia_wh(times_h, powers_w):
    """Integración trapezoidal simple en Wh cuando t está en horas y P en W."""
    e_wh = 0.0
    for i in range(len(times_h) - 1):
        dt_h = times_h[i + 1] - times_h[i]
        e_wh += 0.5 * (powers_w[i] + powers_w[i + 1]) * dt_h
    return e_wh


def formatear_valor_tabla_word(magnitud, valor):
    if pd.isna(valor):
        return ""
    if magnitud == "DOY":
        return f"{int(round(float(valor)))}"
    return f"{float(valor):.4f}"


def preparar_tabla_word(df_resumen):
    df_word = df_resumen.reset_index().rename(columns={"index": "Magnitud"})
    df_word_txt = df_word.copy().astype(object)

    for i in df_word_txt.index:
        magnitud = df_word_txt.at[i, "Magnitud"]
        for col in df_word_txt.columns[1:]:
            df_word_txt.at[i, col] = formatear_valor_tabla_word(magnitud, df_word_txt.at[i, col])

    return df_word_txt


def sombrear_celda(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def escribir_celda(cell, text, bold=False, align="left"):
    cell.text = str(text)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        if align == "center":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == "right":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        for run in paragraph.runs:
            run.bold = bold
            run.font.size = Pt(10)


def exportar_tabla_resumen_word(df_tabla, ruta_docx, titulo, subtitulo=None):
    if not DOCX_AVAILABLE:
        return

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run(titulo)
    r_title.bold = True
    r_title.font.size = Pt(13)

    if subtitulo:
        p_sub = doc.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_sub = p_sub.add_run(subtitulo)
        r_sub.font.size = Pt(10)

    table = doc.add_table(rows=1, cols=len(df_tabla.columns))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    widths_cm = [7.0] + [3.2] * (len(df_tabla.columns) - 1)

    header_cells = table.rows[0].cells
    for j, col_name in enumerate(df_tabla.columns):
        escribir_celda(header_cells[j], col_name, bold=True, align="center")
        sombrear_celda(header_cells[j], "D9EAF7")
        header_cells[j].width = Cm(widths_cm[j])

    for _, row in df_tabla.iterrows():
        row_cells = table.add_row().cells
        for j, col_name in enumerate(df_tabla.columns):
            align = "left" if j == 0 else "center"
            escribir_celda(row_cells[j], row[col_name], bold=False, align=align)
            row_cells[j].width = Cm(widths_cm[j])

    doc.save(ruta_docx)


def exportar_tabla_resumen_html(df_tabla, ruta_html, titulo):
    html_table = df_tabla.to_html(index=False, border=0, justify="center")
    html = f"""<!DOCTYPE html>
<html lang='es'>
<head>
<meta charset='utf-8'>
<title>{titulo}</title>
<style>
body {{ font-family: Arial, sans-serif; margin: 24px; }}
h2 {{ text-align: center; margin-bottom: 18px; }}
table {{ border-collapse: collapse; margin: 0 auto; font-size: 11pt; }}
th, td {{ border: 1px solid #666; padding: 6px 10px; }}
th {{ background: #D9EAF7; text-align: center; }}
td:first-child {{ text-align: left; white-space: nowrap; }}
td:not(:first-child) {{ text-align: center; }}
</style>
</head>
<body>
<h2>{titulo}</h2>
{html_table}
</body>
</html>"""
    ruta_html.write_text(html, encoding="utf-8")


"""
BLOQUE NUEVO: utilidades de presentación compacta y lectura rápida de resultados.

POR QUÉ SE AÑADE:
- La tabla antigua para Word tenía muchas filas y demasiados decimales, lo que la
  hacía más pesada de leer de lo que te interesa ahora.
- En esta fase del TFG lo más útil es tener:
    1) una tabla energética compacta,
    2) una tabla de batería que traduzca masa -> capacidad nominal/reserva/SOC/DoD,
    3) una lectura rápida automática de lo que significan los números.
- Esto NO cambia la física del modelo; solo reorganiza y presenta mejor los
  resultados para consola y Word.
"""


def etiqueta_plataforma(caso):
    return "24 m" if "_24m_" in caso else "32 m"


def etiqueta_momento(caso):
    return "Equinoccio" if "equinoccio" in caso.lower() else "Invierno"


def etiqueta_escenario_corta(escenario):
    if escenario == "Constante_20km":
        return "20 km cte"
    if escenario == "DescensoNocturno_20a15km":
        return "20→15 km"
    return escenario


def etiqueta_caso_breve(caso, escenario=None):
    base = f"{etiqueta_plataforma(caso)} | {etiqueta_momento(caso)}"
    if escenario is None:
        return base
    return f"{base} | {etiqueta_escenario_corta(escenario)}"


def flag_a_texto(flag):
    return "Sí" if int(flag) == 1 else "No"


def ordenar_dataframe_presentacion(df):
    df_ord = df.copy()
    mapa_momento = {"Equinoccio": 0, "Invierno": 1}
    mapa_estrategia = {"20 km cte": 0, "20→15 km": 1}
    df_ord["_ord_plat"] = df_ord["Plat."].map({"24 m": 0, "32 m": 1})
    df_ord["_ord_momento"] = df_ord["Momento"].map(mapa_momento)
    df_ord["_ord_estrategia"] = df_ord["Estrategia"].map(mapa_estrategia)
    df_ord = df_ord.sort_values(["_ord_plat", "_ord_momento", "_ord_estrategia"]).reset_index(drop=True)
    return df_ord.drop(columns=["_ord_plat", "_ord_momento", "_ord_estrategia"])


def aplicar_formato_columnas(df, decimales_por_columna):
    df_fmt = df.copy().astype(object)
    for col, nd in decimales_por_columna.items():
        if col in df_fmt.columns:
            df_fmt[col] = df_fmt[col].map(
                lambda v: "" if pd.isna(v) else f"{float(v):.{nd}f}"
            )
    return df_fmt


def lectura_rapida_caso(row):
    """
    LECTURA RÁPIDA BASADA SOLO EN LOS RESULTADOS DEL MODELO.

    QUÉ HACE:
    - Resume en una frase corta si el día representativo cierra o no.
    - Si 30 y 60 días coinciden prácticamente con 1 día, interpreta que el sistema
      no arrastra déficit acumulativo dentro del modelo actual.
    - Si 30 y 60 días crecen mucho respecto a 1 día, interpreta que existe un
      drenaje neto día tras día.

    IMPORTANTE:
    - Esta frase NO sustituye al análisis físico detallado.
    - Solo compacta la lectura de los números que el propio script ya ha calculado.
    """
    e1 = float(row["E_batt_usable_1dia (Wh)"])
    e60 = float(row["E_batt_usable_60dias (Wh)"])
    autosostenible = int(row["Autosostenible diario (1/0)"]) == 1

    if autosostenible:
        if abs(e60 - e1) <= max(0.01 * max(e1, 1.0), 50.0):
            return "Cierra el día; 30/60 d no crecen"
        return "Cierra el día; revisar deriva"
    return "No cierra el día; déficit acumulativo"


def construir_tabla_presentacion_energetica(df_resultados):
    filas = []
    for _, row in df_resultados.iterrows():
        filas.append({
            "Plat.": etiqueta_plataforma(row["Caso"]),
            "Momento": etiqueta_momento(row["Caso"]),
            "Estrategia": etiqueta_escenario_corta(row["Escenario"]),
            "E array/día (kWh)": row["Energia electrica diaria array (Wh)"] / 1000.0,
            "Balance batt/día (kWh)": row["Balance diario almacenamiento (Wh)"] / 1000.0,
            "1 d": flag_a_texto(row["Viable_preliminar_1d (1/0)"]),
            "30 d": flag_a_texto(row["Viable_preliminar_30d (1/0)"]),
            "60 d": flag_a_texto(row["Viable_preliminar_60d (1/0)"]),
            "Lectura": lectura_rapida_caso(row),
        })
    df = ordenar_dataframe_presentacion(pd.DataFrame(filas))
    return aplicar_formato_columnas(
        df,
        {
            "E array/día (kWh)": 2,
            "Balance batt/día (kWh)": 2,
        },
    )


def construir_tabla_presentacion_bateria_30d(df_resultados):
    filas = []
    for _, row in df_resultados.iterrows():
        filas.append({
            "Plat.": etiqueta_plataforma(row["Caso"]),
            "Momento": etiqueta_momento(row["Caso"]),
            "Estrategia": etiqueta_escenario_corta(row["Escenario"]),
            "Tec. 30 d": row["Tecnologia_min_masa_30d"],
            "M pack 30 d (kg)": row["Masa_min_conservadora_30d (kg)"],
            "E nom. 30 d (kWh)": row["E_nominal_pack_30d (Wh)"] / 1000.0,
            "Reserva 30 d (kWh)": row["E_reserva_pack_30d (Wh)"] / 1000.0,
            "SOC min día (%)": row["SOC_min_diario_pack_30d (%)"],
            "DoD día (%)": row["DoD_diario_pack_30d (%)"],
            "EFC/día": row["EFC_diarios_pack_30d (-)"],
        })
    df = ordenar_dataframe_presentacion(pd.DataFrame(filas))
    return aplicar_formato_columnas(
        df,
        {
            "M pack 30 d (kg)": 2,
            "E nom. 30 d (kWh)": 2,
            "Reserva 30 d (kWh)": 2,
            "SOC min día (%)": 1,
            "DoD día (%)": 1,
            "EFC/día": 3,
        },
    )


def construir_tabla_presentacion_comparacion(df_comparacion):
    if df_comparacion.empty:
        return df_comparacion.copy()

    filas = []
    for _, row in df_comparacion.iterrows():
        plat = etiqueta_plataforma(row["Caso"])
        momento = etiqueta_momento(row["Caso"])
        delta_bal = row["Delta_Balance_almacenamiento_Wh (planeo-constante)"] / 1000.0
        delta_m30 = row["Delta_Masa_min_30d_kg (planeo-constante)"]
        delta_m60 = row["Delta_Masa_min_60d_kg (planeo-constante)"]

        if momento == "Invierno":
            lectura = "Ayuda, pero no rescata invierno"
        else:
            lectura = "Reduce batería con día ya cerrado"

        filas.append({
            "Plat.": plat,
            "Momento": momento,
            "Δ balance día (kWh)": delta_bal,
            "Δ M pack 30 d (kg)": delta_m30,
            "Δ M pack 60 d (kg)": delta_m60,
            "Lectura": lectura,
        })

    df = pd.DataFrame(filas)
    df["_ord_plat"] = df["Plat."].map({"24 m": 0, "32 m": 1})
    df["_ord_momento"] = df["Momento"].map({"Equinoccio": 0, "Invierno": 1})
    df = df.sort_values(["_ord_plat", "_ord_momento"]).drop(columns=["_ord_plat", "_ord_momento"]).reset_index(drop=True)
    return aplicar_formato_columnas(
        df,
        {
            "Δ balance día (kWh)": 2,
            "Δ M pack 30 d (kg)": 2,
            "Δ M pack 60 d (kg)": 2,
        },
    )



def clave_plataforma_desde_caso(nombre_caso):
    if "24m" in nombre_caso:
        return "24m"
    if "32m" in nombre_caso:
        return "32m"
    return "desconocido"


def tecnologia_propulsiva_completa(tech_dict):
    claves = ["Pspec_motor_Wkg", "Pspec_controller_Wkg", "Pspec_propeller_Wkg", "f_install"]
    return not any(pd.isna(tech_dict[c]) for c in claves)


def construir_tabla_presentacion_eficiencia_aero(df_resultados):
    filas = []
    for _, row in df_resultados.iterrows():
        filas.append({
            "Plat.": etiqueta_plataforma(row["Caso"]),
            "Momento": etiqueta_momento(row["Caso"]),
            "Estrategia": etiqueta_escenario_corta(row["Escenario"]),
            "CL": row["CL_diseno_20km (-)"],
            "CD": row["CD_diseno_20km (-)"],
            "L/D diseño": row["E_diseno_20km = CL/CD (-)"],
        })
    df = ordenar_dataframe_presentacion(pd.DataFrame(filas))
    return aplicar_formato_columnas(
        df,
        {
            "CL": 3,
            "CD": 4,
            "L/D diseño": 2,
        },
    )


def construir_tabla_presentacion_potencias_referencia(df_resultados):
    filas = []
    for _, row in df_resultados.iterrows():
        filas.append({
            "Plat.": etiqueta_plataforma(row["Caso"]),
            "Momento": etiqueta_momento(row["Caso"]),
            "Estrategia": etiqueta_escenario_corta(row["Escenario"]),
            "P elec cruc 20 km (W)": row["P_elec_crucero_20km_base (W)"],
            "P thrust max perfil (W)": row["P_thrust_max_perfil (W)"],
            "P elec max perfil (W)": row["P_propulsion_max (W)"],
        })
    df = ordenar_dataframe_presentacion(pd.DataFrame(filas))
    return aplicar_formato_columnas(
        df,
        {
            "P elec cruc 20 km (W)": 2,
            "P thrust max perfil (W)": 2,
            "P elec max perfil (W)": 2,
        },
    )


def construir_resumen_potencias_por_plataforma(df_resultados):
    filas = []
    for plat in ["24m", "32m"]:
        sub = df_resultados[df_resultados["Caso"].str.contains(plat)].copy()
        if len(sub) == 0:
            continue

        p_cruc = float(sub["P_elec_crucero_20km_base (W)"].iloc[0])
        idx_max = sub["P_propulsion_max (W)"].idxmax()
        fila_max = sub.loc[idx_max]

        filas.append({
            "Plataforma": plat,
            "P elec cruc 20 km (W)": p_cruc,
            "P elec max peor caso (W)": float(fila_max["P_propulsion_max (W)"]),
            "Caso peor": fila_max["Caso"],
            "Escenario peor": fila_max["Escenario"],
        })

    df = pd.DataFrame(filas)
    if len(df) == 0:
        return df

    return aplicar_formato_columnas(
        df,
        {
            "P elec cruc 20 km (W)": 2,
            "P elec max peor caso (W)": 2,
        },
    )


def construir_tabla_presentacion_mf_energy(df_resultados):
    filas = []
    for _, row in df_resultados.iterrows():
        filas.append({
            "Plat.": etiqueta_plataforma(row["Caso"]),
            "Momento": etiqueta_momento(row["Caso"]),
            "Estrategia": etiqueta_escenario_corta(row["Escenario"]),
            "Tec. 30 d": row["Tecnologia_min_masa_30d"],
            "M bat 30 d (kg)": row["Masa_min_conservadora_30d (kg)"],
            "MF_Energy 30 d": row["MF_Energy_30d (-)"],
            "Tec. 60 d": row["Tecnologia_min_masa_60d"],
            "M bat 60 d (kg)": row["Masa_min_conservadora_60d (kg)"],
            "MF_Energy 60 d": row["MF_Energy_60d (-)"],
        })
    df = ordenar_dataframe_presentacion(pd.DataFrame(filas))
    return aplicar_formato_columnas(
        df,
        {
            "M bat 30 d (kg)": 2,
            "MF_Energy 30 d": 3,
            "M bat 60 d (kg)": 2,
            "MF_Energy 60 d": 3,
        },
    )


def construir_tabla_presentacion_mf_struct_subs(df_resultados):
    filas = []
    for _, row in df_resultados.iterrows():
        filas.append({
            "Plat.": etiqueta_plataforma(row["Caso"]),
            "Momento": etiqueta_momento(row["Caso"]),
            "Estrategia": etiqueta_escenario_corta(row["Escenario"]),
            "MF_Struct asumida": row["MF_Struct_asumida (-)"],
            "MF_Subs asumida": row["MF_Subs_asumida (-)"],
            "MF_Energy 30 d": row["MF_Energy_30d (-)"],
        })
    df = ordenar_dataframe_presentacion(pd.DataFrame(filas))
    return aplicar_formato_columnas(
        df,
        {
            "MF_Struct asumida": 3,
            "MF_Subs asumida": 3,
            "MF_Energy 30 d": 3,
        },
    )


def construir_tabla_presentacion_pw_tw(df_resultados):
    filas = []
    for _, row in df_resultados.iterrows():
        filas.append({
            "Plat.": etiqueta_plataforma(row["Caso"]),
            "Momento": etiqueta_momento(row["Caso"]),
            "Estrategia": etiqueta_escenario_corta(row["Escenario"]),
            "P/W cruc (W/kg)": row["P_W_elec_crucero_Wkg"],
            "P/W max (W/kg)": row["P_W_elec_max_perfil_Wkg"],
            "T/W cruc": row["T_W_crucero_20km_base (-)"],
            "T/W max": row["T_W_max_perfil (-)"],
        })
    df = ordenar_dataframe_presentacion(pd.DataFrame(filas))
    return aplicar_formato_columnas(
        df,
        {
            "P/W cruc (W/kg)": 2,
            "P/W max (W/kg)": 2,
            "T/W cruc": 3,
            "T/W max": 3,
        },
    )


def evaluar_tecnologia_propulsiva(row, tech_name, tech_dict, eta_propulsor_para_sizing_local):
    pspec_motor = tech_dict["Pspec_motor_Wkg"]
    pspec_controller = tech_dict["Pspec_controller_Wkg"]
    pspec_propeller = tech_dict["Pspec_propeller_Wkg"]
    f_install = tech_dict["f_install"]

    p_thrust_ref_w = float(row["P_thrust_max_perfil (W)"])
    p_shaft_ref_w = p_thrust_ref_w / eta_propulsor_para_sizing_local

    m_motor = p_shaft_ref_w / pspec_motor
    m_controller = p_shaft_ref_w / pspec_controller
    m_propeller = p_shaft_ref_w / pspec_propeller

    m_prop_inst = f_install * (m_motor + m_controller + m_propeller)
    mf_prop = m_prop_inst / float(row["MTOW (kg)"])

    return {
        "Caso": row["Caso"],
        "Escenario": row["Escenario"],
        "Tecnologia_propulsiva": tech_name,
        "P_shaft_ref_W": p_shaft_ref_w,
        "M_motor_kg": m_motor,
        "M_controller_kg": m_controller,
        "M_propeller_kg": m_propeller,
        "M_propulsion_instalada_kg": m_prop_inst,
        "MF_Prop (-)": mf_prop,
    }


def construir_tabla_presentacion_propulsion(df_propulsion_sizing):
    if df_propulsion_sizing.empty:
        return df_propulsion_sizing.copy()

    filas = []
    for _, row in df_propulsion_sizing.iterrows():
        filas.append({
            "Plat.": etiqueta_plataforma(row["Caso"]),
            "Momento": etiqueta_momento(row["Caso"]),
            "Estrategia": etiqueta_escenario_corta(row["Escenario"]),
            "Tec. prop.": row["Tecnologia_propulsiva"],
            "P shaft ref (W)": row["P_shaft_ref_W"],
            "M prop inst (kg)": row["M_propulsion_instalada_kg"],
            "MF_Prop": row["MF_Prop (-)"],
        })
    df = ordenar_dataframe_presentacion(pd.DataFrame(filas))
    return aplicar_formato_columnas(
        df,
        {
            "P shaft ref (W)": 2,
            "M prop inst (kg)": 3,
            "MF_Prop": 3,
        },
    )



def seleccionar_mf_energy_referencia(row, mission_days_ref):
    mapa = {
        1: "MF_Energy_1d (-)",
        30: "MF_Energy_30d (-)",
        60: "MF_Energy_60d (-)",
    }
    col = mapa[mission_days_ref]
    return float(row[col]), col


def obtener_fixed_weights_total_preliminar(
    row,
    fixed_weights_mode_local,
    mission_days_ref,
    mf_prop_ref_local,
    fixed_weights_manual_local,
):
    """
    CIERRE PRELIMINAR DE W_FIXED.

    MODO 1: implied_from_current_mtow
    - Reconstruye el total de pesos fijos consistente con el MTOW actual del caso.
    - Es útil mientras todavía no has fijado en masa W_PL, W_Avion y W_Other.

    MODO 2: manual
    - Usa las masas fijas introducidas a mano por plataforma.

    IMPORTANTE:
    - En el modo "implied" la descomposición W_PL/W_Avion/W_Other no es física;
      solo se recupera el TOTAL de pesos fijos.
    """
    plat_key = clave_plataforma_desde_caso(row["Caso"])
    mf_energy_ref, mf_energy_col = seleccionar_mf_energy_referencia(row, mission_days_ref)
    mf_struct = float(row["MF_Struct_asumida (-)"])
    mf_subs = float(row["MF_Subs_asumida (-)"])
    mtow_input = float(row["MTOW (kg)"])

    if fixed_weights_mode_local == "manual":
        fw = fixed_weights_manual_local[plat_key]
        w_pl = float(fw["W_PL_kg"])
        w_avion = float(fw["W_Avion_kg"])
        w_other = float(fw["W_Other_kg"])
        w_fixed_total = w_pl + w_avion + w_other
        return {
            "W_PL_kg": w_pl,
            "W_Avion_kg": w_avion,
            "W_Other_kg": w_other,
            "W_fixed_total_kg": w_fixed_total,
            "MF_Energy_ref": mf_energy_ref,
            "MF_Energy_col": mf_energy_col,
            "Metodo_fixed_weights": "manual",
            "Observacion_fixed": "W_fixed introducido manualmente",
        }

    if fixed_weights_mode_local == "implied_from_current_mtow":
        w_fixed_total = mtow_input * (1.0 - mf_struct - mf_subs - mf_prop_ref_local - mf_energy_ref)
        return {
            "W_PL_kg": float("nan"),
            "W_Avion_kg": float("nan"),
            "W_Other_kg": float("nan"),
            "W_fixed_total_kg": w_fixed_total,
            "MF_Energy_ref": mf_energy_ref,
            "MF_Energy_col": mf_energy_col,
            "Metodo_fixed_weights": "implied_from_current_mtow",
            "Observacion_fixed": (
                f"W_fixed reconstruido desde MTOW actual usando MF_Prop_ref={mf_prop_ref_local:.2f}"
            ),
        }

    raise ValueError("fixed_weights_mode no reconocido.")


def calcular_wto_directo_preliminar(w_fixed_total_kg, mf_struct, mf_subs, mf_prop, mf_energy):
    suma_mf = mf_struct + mf_subs + mf_prop + mf_energy
    denom = 1.0 - suma_mf

    if (w_fixed_total_kg <= 0.0) or (denom <= 0.0):
        return {
            "WTO_direct_kg": float("nan"),
            "WEF (-)": float("nan"),
            "Suma_MF (-)": suma_mf,
            "Denominador (-)": denom,
            "Factible (1/0)": 0,
            "Diagnostico_WTO": "No hay espacio factible: W_fixed<=0 o denominador<=0",
        }

    wto_direct = w_fixed_total_kg / denom
    wef = 1.0 / denom

    diagnostico = "Factible"
    if wef > 10.0:
        diagnostico = "Factible pero WEF>10 (muy sensible)"
    elif wef > 8.0:
        diagnostico = "Factible pero WEF alto"

    return {
        "WTO_direct_kg": wto_direct,
        "WEF (-)": wef,
        "Suma_MF (-)": suma_mf,
        "Denominador (-)": denom,
        "Factible (1/0)": 1,
        "Diagnostico_WTO": diagnostico,
    }


def diff_wto_lineal(wto_input_kg, w_fixed_total_kg, suma_mf):
    """
    Forma iterativa linealizada de la ecuación de Gundlach.

    WTO_calc = W_fixed + (MFStruct + MFSubs + MFProp + MFEnergy) * WTO_input
    Diff = WTO_calc - WTO_input

    Si Diff = 0, se ha encontrado el WTO consistente con el conjunto de hipótesis.
    """
    wto_calc = w_fixed_total_kg + suma_mf * wto_input_kg
    return wto_calc - wto_input_kg


def resolver_wto_secante_lineal(
    w_fixed_total_kg,
    suma_mf,
    wto_guess_1_kg,
    wto_guess_2_kg,
    tol_kg,
    max_iter,
):
    """
    ITERACIÓN TIPO SECANTE PARA EL CIERRE PRELIMINAR DE WTO.

    IMPORTANTE:
    - Como aquí las fracciones se mantienen congeladas dentro de una iteración,
      el problema es lineal y la secante converge muy rápido.
    - Se deja igualmente implementado porque conceptualmente enlaza con la
      discusión de Gundlach sobre convergencia de WTO.
    """
    x0 = float(wto_guess_1_kg)
    x1 = float(wto_guess_2_kg)

    f0 = diff_wto_lineal(x0, w_fixed_total_kg, suma_mf)
    f1 = diff_wto_lineal(x1, w_fixed_total_kg, suma_mf)

    historial = [(x0, f0), (x1, f1)]

    for i in range(max_iter):
        if abs(f1 - f0) < 1e-12:
            return {
                "WTO_secante_kg": float("nan"),
                "Iter_secante": i + 2,
                "Conv_secante (1/0)": 0,
                "Historial_secante": historial,
            }

        x2 = x1 - f1 * (x1 - x0) / (f1 - f0)
        f2 = diff_wto_lineal(x2, w_fixed_total_kg, suma_mf)
        historial.append((x2, f2))

        if abs(f2) <= tol_kg:
            return {
                "WTO_secante_kg": x2,
                "Iter_secante": i + 3,
                "Conv_secante (1/0)": 1,
                "Historial_secante": historial,
            }

        x0, f0 = x1, f1
        x1, f1 = x2, f2

    return {
        "WTO_secante_kg": x1,
        "Iter_secante": max_iter + 2,
        "Conv_secante (1/0)": 0,
        "Historial_secante": historial,
    }


def construir_dataframe_wto_preliminar(
    df_resultados,
    df_propulsion_sizing,
    fixed_weights_mode_local,
    mission_days_ref,
    mf_prop_assumed_cases_local,
    fixed_weights_manual_local,
    mf_prop_ref_local,
    tol_kg_local,
    max_iter_local,
):
    """
    CONSTRUYE EL BLOQUE FINAL DEL ESQUEMA DE GUNDLACH PARA WTO PRELIMINAR.

    QUÉ HACE:
    - Usa MF_Struct, MF_Subs, MF_Energy y MF_Prop para cerrar un WTO esperado.
    - Genera tres casos asumidos de MF_Prop si no hay benchmark de motor.
    - Si sí existe benchmark tecnológico, añade además los casos calculados de
      MF_Prop a partir del bloque 8.
    """
    filas = []

    for _, row in df_resultados.iterrows():
        base_fixed = obtener_fixed_weights_total_preliminar(
            row=row,
            fixed_weights_mode_local=fixed_weights_mode_local,
            mission_days_ref=mission_days_ref,
            mf_prop_ref_local=mf_prop_ref_local,
            fixed_weights_manual_local=fixed_weights_manual_local,
        )

        mf_struct = float(row["MF_Struct_asumida (-)"])
        mf_subs = float(row["MF_Subs_asumida (-)"])
        mf_energy = float(base_fixed["MF_Energy_ref"])
        mtow_input = float(row["MTOW (kg)"])
        w_fixed_total = float(base_fixed["W_fixed_total_kg"])

        # Casos asumidos de MF_Prop
        for nombre_mf_prop, datos_mf_prop in mf_prop_assumed_cases_local.items():
            mf_prop = float(datos_mf_prop["MF_Prop"])
            origen_mf_prop = datos_mf_prop["Origen_MF_Prop"]

            calc = calcular_wto_directo_preliminar(
                w_fixed_total_kg=w_fixed_total,
                mf_struct=mf_struct,
                mf_subs=mf_subs,
                mf_prop=mf_prop,
                mf_energy=mf_energy,
            )

            suma_mf = calc["Suma_MF (-)"]
            sec = resolver_wto_secante_lineal(
                w_fixed_total_kg=w_fixed_total,
                suma_mf=suma_mf,
                wto_guess_1_kg=max(w_fixed_total, 1.0),
                wto_guess_2_kg=max(2.0 * mtow_input, max(w_fixed_total * 2.0, 10.0)),
                tol_kg=tol_kg_local,
                max_iter=max_iter_local,
            ) if calc["Factible (1/0)"] == 1 else {
                "WTO_secante_kg": float("nan"),
                "Iter_secante": 0,
                "Conv_secante (1/0)": 0,
                "Historial_secante": [],
            }

            filas.append({
                "Caso": row["Caso"],
                "Escenario": row["Escenario"],
                "Modo_fixed_weights": base_fixed["Metodo_fixed_weights"],
                "Observacion_fixed": base_fixed["Observacion_fixed"],
                "W_PL_kg": base_fixed["W_PL_kg"],
                "W_Avion_kg": base_fixed["W_Avion_kg"],
                "W_Other_kg": base_fixed["W_Other_kg"],
                "W_fixed_total_kg": w_fixed_total,
                "Horizonte_MF_Energy_dias": mission_days_ref,
                "MF_Energy_ref (-)": mf_energy,
                "MF_Struct (-)": mf_struct,
                "MF_Subs (-)": mf_subs,
                "Caso_MF_Prop": nombre_mf_prop,
                "MF_Prop (-)": mf_prop,
                "Origen_MF_Prop": origen_mf_prop,
                "MTOW_input_kg": mtow_input,
                "WTO_direct_kg": calc["WTO_direct_kg"],
                "WEF (-)": calc["WEF (-)"],
                "Suma_MF (-)": calc["Suma_MF (-)"],
                "Denominador (-)": calc["Denominador (-)"],
                "Factible (1/0)": calc["Factible (1/0)"],
                "Diagnostico_WTO": calc["Diagnostico_WTO"],
                "WTO_secante_kg": sec["WTO_secante_kg"],
                "Iter_secante": sec["Iter_secante"],
                "Conv_secante (1/0)": sec["Conv_secante (1/0)"],
                "Delta_WTO_vs_input_kg": (
                    calc["WTO_direct_kg"] - mtow_input if not pd.isna(calc["WTO_direct_kg"]) else float("nan")
                ),
                "Delta_WTO_vs_input_pct": (
                    100.0 * (calc["WTO_direct_kg"] - mtow_input) / mtow_input
                    if (not pd.isna(calc["WTO_direct_kg"]) and mtow_input > 0.0) else float("nan")
                ),
            })

        # Casos benchmark de MF_Prop, si existen
        if not df_propulsion_sizing.empty:
            sub_prop = df_propulsion_sizing[
                (df_propulsion_sizing["Caso"] == row["Caso"]) &
                (df_propulsion_sizing["Escenario"] == row["Escenario"])
            ].copy()

            for _, prow in sub_prop.iterrows():
                mf_prop = float(prow["MF_Prop (-)"])

                calc = calcular_wto_directo_preliminar(
                    w_fixed_total_kg=w_fixed_total,
                    mf_struct=mf_struct,
                    mf_subs=mf_subs,
                    mf_prop=mf_prop,
                    mf_energy=mf_energy,
                )

                suma_mf = calc["Suma_MF (-)"]
                sec = resolver_wto_secante_lineal(
                    w_fixed_total_kg=w_fixed_total,
                    suma_mf=suma_mf,
                    wto_guess_1_kg=max(w_fixed_total, 1.0),
                    wto_guess_2_kg=max(2.0 * mtow_input, max(w_fixed_total * 2.0, 10.0)),
                    tol_kg=tol_kg_local,
                    max_iter=max_iter_local,
                ) if calc["Factible (1/0)"] == 1 else {
                    "WTO_secante_kg": float("nan"),
                    "Iter_secante": 0,
                    "Conv_secante (1/0)": 0,
                    "Historial_secante": [],
                }

                filas.append({
                    "Caso": row["Caso"],
                    "Escenario": row["Escenario"],
                    "Modo_fixed_weights": base_fixed["Metodo_fixed_weights"],
                    "Observacion_fixed": base_fixed["Observacion_fixed"],
                    "W_PL_kg": base_fixed["W_PL_kg"],
                    "W_Avion_kg": base_fixed["W_Avion_kg"],
                    "W_Other_kg": base_fixed["W_Other_kg"],
                    "W_fixed_total_kg": w_fixed_total,
                    "Horizonte_MF_Energy_dias": mission_days_ref,
                    "MF_Energy_ref (-)": mf_energy,
                    "MF_Struct (-)": mf_struct,
                    "MF_Subs (-)": mf_subs,
                    "Caso_MF_Prop": prow["Tecnologia_propulsiva"],
                    "MF_Prop (-)": mf_prop,
                    "Origen_MF_Prop": "Calculado desde bloque 8 / benchmark tecnológico",
                    "MTOW_input_kg": mtow_input,
                    "WTO_direct_kg": calc["WTO_direct_kg"],
                    "WEF (-)": calc["WEF (-)"],
                    "Suma_MF (-)": calc["Suma_MF (-)"],
                    "Denominador (-)": calc["Denominador (-)"],
                    "Factible (1/0)": calc["Factible (1/0)"],
                    "Diagnostico_WTO": calc["Diagnostico_WTO"],
                    "WTO_secante_kg": sec["WTO_secante_kg"],
                    "Iter_secante": sec["Iter_secante"],
                    "Conv_secante (1/0)": sec["Conv_secante (1/0)"],
                    "Delta_WTO_vs_input_kg": (
                        calc["WTO_direct_kg"] - mtow_input if not pd.isna(calc["WTO_direct_kg"]) else float("nan")
                    ),
                    "Delta_WTO_vs_input_pct": (
                        100.0 * (calc["WTO_direct_kg"] - mtow_input) / mtow_input
                        if (not pd.isna(calc["WTO_direct_kg"]) and mtow_input > 0.0) else float("nan")
                    ),
                })

    return pd.DataFrame(filas)


def construir_tabla_presentacion_wto_preliminar(df_wto_preliminar):
    if df_wto_preliminar.empty:
        return df_wto_preliminar.copy()

    filas = []
    for _, row in df_wto_preliminar.iterrows():
        filas.append({
            "Plat.": etiqueta_plataforma(row["Caso"]),
            "Momento": etiqueta_momento(row["Caso"]),
            "Estrategia": etiqueta_escenario_corta(row["Escenario"]),
            "Caso MF_Prop": row["Caso_MF_Prop"],
            "MF_Energy": row["MF_Energy_ref (-)"],
            "MF_Prop": row["MF_Prop (-)"],
            "W_fixed (kg)": row["W_fixed_total_kg"],
            "WEF": row["WEF (-)"],
            "WTO input (kg)": row["MTOW_input_kg"],
            "WTO calc (kg)": row["WTO_direct_kg"],
            "Δ WTO (%)": row["Delta_WTO_vs_input_pct"],
            "Factible": "Sí" if int(row["Factible (1/0)"]) == 1 else "No",
            "Lectura": row["Diagnostico_WTO"],
        })

    df = ordenar_dataframe_presentacion(pd.DataFrame(filas))
    return aplicar_formato_columnas(
        df,
        {
            "MF_Energy": 3,
            "MF_Prop": 3,
            "W_fixed (kg)": 2,
            "WEF": 2,
            "WTO input (kg)": 2,
            "WTO calc (kg)": 2,
            "Δ WTO (%)": 1,
        },
    )


def construir_conclusiones_por_caso(df_resultados):
    """
    BLOQUE NUEVO: conclusiones cortas, directamente deducidas de los datos.

    QUÉ HACE:
    - Agrupa por caso base (plataforma + momento del año) y compara sus dos
      estrategias operativas.
    - Resume en lenguaje natural:
        * si el caso cierra energéticamente,
        * si el descenso nocturno ayuda,
        * y si esa ayuda es suficiente o no.

    IMPORTANTE:
    - Las conclusiones salen solo de los números del script.
    - No introducen hipótesis nuevas.
    """
    conclusiones = []

    for caso in sorted(df_resultados["Caso"].unique()):
        df_caso = df_resultados[df_resultados["Caso"] == caso].copy()
        fila_const = df_caso[df_caso["Escenario"] == "Constante_20km"].iloc[0]
        fila_desc = df_caso[df_caso["Escenario"] == "DescensoNocturno_20a15km"].iloc[0]

        base = f"{etiqueta_plataforma(caso)} | {etiqueta_momento(caso)}"
        mejora_balance_kwh = (fila_desc["Balance diario almacenamiento (Wh)"] - fila_const["Balance diario almacenamiento (Wh)"]) / 1000.0
        mejora_m30 = fila_desc["Masa_min_conservadora_30d (kg)"] - fila_const["Masa_min_conservadora_30d (kg)"]

        if int(fila_const["Autosostenible diario (1/0)"]) == 1 and int(fila_desc["Autosostenible diario (1/0)"]) == 1:
            texto = (
                f"{base}: ambos escenarios cierran el día; el descenso 20→15 km "
                f"reduce la masa mínima de batería a 30 días en {abs(mejora_m30):.2f} kg."
            )
        elif int(fila_const["Autosostenible diario (1/0)"]) == 0 and int(fila_desc["Autosostenible diario (1/0)"]) == 0:
            texto = (
                f"{base}: el descenso 20→15 km mejora el balance diario en {mejora_balance_kwh:.2f} kWh/d, "
                f"pero el caso sigue sin cerrar energéticamente."
            )
        else:
            texto = (
                f"{base}: el descenso cambia la autosostenibilidad diaria y conviene revisar este caso con más detalle."
            )

        conclusiones.append(texto)

    return conclusiones


def insertar_tabla_docx_generica(doc, df_tabla, titulo_tabla, widths_cm, font_size_pt=8.5):
    p_head = doc.add_paragraph()
    p_head.paragraph_format.space_before = Pt(8)
    p_head.paragraph_format.space_after = Pt(4)
    r_head = p_head.add_run(titulo_tabla)
    r_head.bold = True
    r_head.font.size = Pt(11)

    table = doc.add_table(rows=1, cols=len(df_tabla.columns))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    header_cells = table.rows[0].cells
    for j, col_name in enumerate(df_tabla.columns):
        escribir_celda(header_cells[j], col_name, bold=True, align="center")
        sombrear_celda(header_cells[j], "D9EAF7")
        header_cells[j].width = Cm(widths_cm[j])

    for _, row in df_tabla.iterrows():
        row_cells = table.add_row().cells
        for j, col_name in enumerate(df_tabla.columns):
            align = "left" if col_name == "Lectura" else "center"
            escribir_celda(row_cells[j], row[col_name], bold=False, align=align)
            row_cells[j].width = Cm(widths_cm[j])
            for paragraph in row_cells[j].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(font_size_pt)


def exportar_resumen_compacto_word(
    ruta_docx,
    titulo,
    subtitulo,
    df_energetica,
    df_bateria_30d,
    df_comparacion_presentacion,
    conclusiones_cortas,
):
    if not DOCX_AVAILABLE:
        return

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.6)
    section.right_margin = Cm(1.6)

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run(titulo)
    r_title.bold = True
    r_title.font.size = Pt(14)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_sub.add_run(subtitulo)
    r_sub.font.size = Pt(9)

    p_intro = doc.add_paragraph()
    p_intro.paragraph_format.space_after = Pt(6)
    r_intro = p_intro.add_run(
        "Resumen compacto orientado a lectura rápida del TFG. "
        "Se eliminan filas de detalle menos útiles y se resaltan energía diaria, "
        "viabilidad y traducción masa→capacidad de batería."
    )
    r_intro.font.size = Pt(9)

    insertar_tabla_docx_generica(
        doc,
        df_energetica,
        "1. Resumen energético y viabilidad",
        widths_cm=[1.3, 2.0, 1.9, 2.1, 2.3, 1.0, 1.1, 1.1, 4.0],
        font_size_pt=8.2,
    )

    insertar_tabla_docx_generica(
        doc,
        df_bateria_30d,
        "2. Interpretación del pack mínimo de batería para 30 días",
        widths_cm=[1.3, 2.0, 1.9, 1.2, 1.8, 1.9, 1.8, 1.6, 1.5, 1.2],
        font_size_pt=8.0,
    )

    if len(df_comparacion_presentacion) > 0:
        insertar_tabla_docx_generica(
            doc,
            df_comparacion_presentacion,
            "3. Efecto del descenso nocturno frente a crucero constante",
            widths_cm=[1.3, 2.0, 2.4, 2.2, 2.2, 4.0],
            font_size_pt=8.2,
        )

    p_cons = doc.add_paragraph()
    p_cons.paragraph_format.space_before = Pt(8)
    r_cons = p_cons.add_run("4. Conclusiones automáticas a partir de los resultados")
    r_cons.bold = True
    r_cons.font.size = Pt(11)

    for texto in conclusiones_cortas:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(texto)
        r.font.size = Pt(9)

    doc.save(ruta_docx)


"""
BLOQUE NUEVO: guardado de figuras sin depender de que el backend sea interactivo.

POR QUÉ SE AÑADE:
- En Spyder te interesa poder ver las figuras con plt.show().
- Pero cuando el script se ejecuta con un backend no interactivo (por ejemplo Agg
  en pruebas automáticas), llamar a show() solo genera avisos y no aporta nada.
- Esta función guarda siempre la figura en PNG y solo llama a show() cuando el
  backend realmente puede mostrar ventanas o paneles gráficos.
"""


def guardar_figura_y_mostrar_si_procede(fig, ruta_png):
    fig.tight_layout()
    fig.savefig(ruta_png, dpi=200, bbox_inches="tight")
    backend = str(plt.get_backend()).lower()
    if "agg" not in backend:
        plt.show()
    plt.close(fig)



"""
FUNCIÓN NUEVA: resolver la geometría base del caso en el punto de diseño.

ORIGEN:
- Se mantiene exactamente la lógica que ya tenías en el script base: el ala se
  dimensiona en el punto de crucero nominal (20 km, velocidad fijada) usando la
  polar parabólica conceptual y la condición de mínima potencia requerida.
- Esta parte enlaza con Gundlach, Ecs. (3.40) a (3.46), donde la potencia del
  sistema eléctrico se relaciona con el empuje, la resistencia y el parámetro
  C_L^(3/2) / C_D.

POR QUÉ SE SEPARA EN UNA FUNCIÓN:
- Porque, una vez fijada la geometría del caso, el escenario de descenso nocturno
  NO debe volver a redimensionar el ala a 15 km. Lo correcto es congelar la
  geometría del avión y recalcular después cómo cambian CL, CD, D y la potencia
  requerida cuando cambia la altitud del vuelo.
"""

def resolver_geometria_crucero_base(W_TO, b, rho_design):
    W_N = W_TO * g
    AR = (((0.5 * rho_design * (v_crucero ** 2) * (b ** 2)) / W_N) ** 2) * 3 * PI * e * C_D0
    K = 1.0 / (PI * AR * e)
    C_L = math.sqrt((3.0 * C_D0) / K)
    S_alar = b ** 2 / AR
    C_D = C_D0 + K * (C_L ** 2)
    E_aero = C_L / C_D
    D_crucero = W_N / E_aero
    T_crucero = D_crucero
    P_thrust_crucero = T_crucero * v_crucero
    P_elec_crucero = P_thrust_crucero / eta_prop_global

    return {
        "W_N": W_N,
        "AR": AR,
        "K": K,
        "S_alar": S_alar,
        "CL_design": C_L,
        "CD_design": C_D,
        "E_aero_design": E_aero,
        "D_crucero_design": D_crucero,
        "T_crucero_design": T_crucero,
        "P_thrust_crucero_design": P_thrust_crucero,
        "P_elec_crucero_design": P_elec_crucero,
    }


"""
FUNCIÓN AUXILIAR SIMPLE: construir un perfil de potencia constante.

ORIGEN:
- No es una ecuación de Gundlach.
- Es simplemente una utilidad numérica para representar cargas constantes del bus,
  como la aviónica o una payload todavía no definida en detalle.

POR QUÉ SE USA:
- En esta fase quieres mantener el programa lo más claro posible.
- Por eso la PL se trata como una potencia constante editable, en vez de introducir
  una lógica temporal específica de cámara o sensor.
"""
def construir_perfil_constante(times_h, valor_constante_w):
    return [valor_constante_w for _ in times_h]



def construir_perfil_operativo_propulsion(
    times_h,            #la malla temporal del día, por ejemplo cada 5 min.
    t_sr_panel,         #hora solar de amanecer útil del panel.
    t_ss_panel,         #hora solar de puesta útil del panel.
    W_N,
    S_alar,
    K,
    escenario_dict,
):
    tipo = escenario_dict["tipo"]           
    h_max_m = escenario_dict["h_max_m"]
    h_min_m = escenario_dict["h_min_m"]

    day_hours = t_ss_panel - t_sr_panel         #horas útiles de panel
    night_hours = 24.0 - day_hours              #horas quedan de noche
    delta_h_m = h_max_m - h_min_m

    perfil_alt_m = []
    perfil_rho = []
    perfil_roc_mps = []
    perfil_cl = []
    perfil_cd = []
    perfil_ld = []
    perfil_drag_n = []
    perfil_thrust_n = []
    perfil_p_thrust_w = []
    perfil_p_elec_w = []
    perfil_modo = []

    for t_h in times_h:
        if tipo == "crucero_constante":
            h_m = h_max_m
            roc_mps = 0.0
            modo_vuelo = "crucero_constante"

        elif tipo == "descenso_nocturno_y_reascenso_diurno":
            tau_h = (t_h - t_sr_panel) % 24.0   #establece la hora de amanecer útil

            if tau_h < day_hours:
                frac = tau_h / day_hours if day_hours > 0.0 else 0.0    # % avanzado del dia
                h_m = h_min_m + delta_h_m * frac                            #sube h linealmente
                roc_mps = delta_h_m / (day_hours * 3600.0) if day_hours > 0.0 else 0.0  #razon de ascenso cte
                modo_vuelo = "ascenso_diurno"
            else:
                tau_night_h = tau_h - day_hours
                frac = tau_night_h / night_hours if night_hours > 0.0 else 0.0
                h_m = h_max_m - delta_h_m * frac
                roc_mps = -delta_h_m / (night_hours * 3600.0) if night_hours > 0.0 else 0.0
                modo_vuelo = "descenso_nocturno"
        else:
            raise ValueError("Tipo de escenario operativo no reconocido.")

        rho_local = atmosfera_isa(h_m)
        q_local = 0.5 * rho_local * (v_crucero ** 2)

        # MECÁNICA DE VUELO BÁSICA:
        # L = q*S*CL y, para este modelo cuasi estacionario, L ≈ W.
        # De aquí sale CL = W / (q*S). Esta igualdad no es una novedad de Gundlach,
        # sino la ecuación clásica de sustentación aplicada al punto de vuelo local.
        C_L_local = W_N / (q_local * S_alar)

        # POLAR PARABÓLICA CONCEPTUAL:
        # CD = CD0 + K*CL^2
        # Esta es la misma forma de polar usada en el dimensionado conceptual
        # que venías empleando a partir de Gundlach.
        C_D_local = C_D0 + K * (C_L_local ** 2)

        E_local = C_L_local / C_D_local

        # DEFINICIÓN AERODINÁMICA DE RESISTENCIA:
        # D = q*S*CD
        D_local = q_local * S_alar * C_D_local

        # GUNDLACH Eq. (9.26):
        # T_req = D + W*ROC/V
        # Si ROC > 0, el ascenso exige potencia adicional.
        # Si ROC < 0, parte del peso ayuda a sostener el vuelo y el empuje requerido baja.
        T_req_local = D_local + W_N * roc_mps / v_crucero
        T_req_local = max(0.0, T_req_local)     #si la T sale negativa=apagar motores=planeo

        # GUNDLACH Eq. (3.40) + Eq. (9.26):
        # P_thrust = T*V  ->  P_thrust_req = D*V + W*ROC
        P_thrust_local = T_req_local * v_crucero
        P_elec_local = P_thrust_local / eta_prop_global

        perfil_alt_m.append(h_m)
        perfil_rho.append(rho_local)
        perfil_roc_mps.append(roc_mps)
        perfil_cl.append(C_L_local)
        perfil_cd.append(C_D_local)
        perfil_ld.append(E_local)
        perfil_drag_n.append(D_local)
        perfil_thrust_n.append(T_req_local)
        perfil_p_thrust_w.append(P_thrust_local)
        perfil_p_elec_w.append(P_elec_local)
        perfil_modo.append(modo_vuelo)

    return {
        "altitud_m": perfil_alt_m,
        "rho_kgm3": perfil_rho,
        "roc_mps": perfil_roc_mps,
        "cl": perfil_cl,
        "cd": perfil_cd,
        "ld": perfil_ld,
        "drag_n": perfil_drag_n,
        "thrust_n": perfil_thrust_n,
        "p_thrust_w": perfil_p_thrust_w,
        "p_elec_w": perfil_p_elec_w,
        "modo_vuelo": perfil_modo,
        "day_hours": day_hours,
        "night_hours": night_hours,
    }


"""
FUNCIÓN REEMPLAZADA: construir el perfil energético de la batería en intervalos.

NOVEDAD RESPECTO A LA VERSIÓN ANTERIOR:
- Antes la potencia total del bus se trataba como constante.
- Ahora se acepta un perfil temporal completo de potencia total, porque la payload
  puede encenderse solo durante una ventana del día y porque el escenario de
  descenso nocturno / reascenso diurno hace variar la potencia propulsiva.

RELACIÓN CON GUNDLACH:
- Eq. (8.104): la diferencia entre la potencia del array y la demanda del avión
  define el exceso o déficit del sistema.
- Eq. (8.105): si el exceso es positivo, parte se almacena con la eficiencia de
  carga.
- Eqs. (8.106) a (8.108): si el balance es negativo, la batería cubre la noche o
  los tramos de dawn/dusk con insuficiencia solar.
"""

def construir_perfil_bateria_intervalos(
    times_h,
    el_list_deg,
    p_array_electric_w,
    p_propulsion_bus_w_points,
    p_sistemas_bus_w_points,
    p_payload_bus_w_points,
    altitud_m_points,
    roc_mps_points,
    modo_vuelo_points,
    eta_power_conditioning_local,
    eta_batt_charge_local,
    eta_batt_discharge_local,
):
    registros = []

    for i in range(len(times_h) - 1):
        t_ini = times_h[i]
        t_fin = times_h[i + 1]
        dt_h = t_fin - t_ini
        t_mid = 0.5 * (t_ini + t_fin)

        el_mid_deg = 0.5 * (el_list_deg[i] + el_list_deg[i + 1])
        p_array_bus_w = 0.5 * (p_array_electric_w[i] + p_array_electric_w[i + 1]) * eta_power_conditioning_local

        p_prop_mid_w = 0.5 * (p_propulsion_bus_w_points[i] + p_propulsion_bus_w_points[i + 1])
        p_sys_mid_w = 0.5 * (p_sistemas_bus_w_points[i] + p_sistemas_bus_w_points[i + 1])
        p_pl_mid_w = 0.5 * (p_payload_bus_w_points[i] + p_payload_bus_w_points[i + 1])
        p_total_bus_w = p_prop_mid_w + p_sys_mid_w + p_pl_mid_w
        p_net_bus_w = p_array_bus_w - p_total_bus_w

        alt_mid_m = 0.5 * (altitud_m_points[i] + altitud_m_points[i + 1])
        roc_mid_mps = 0.5 * (roc_mps_points[i] + roc_mps_points[i + 1])

        if modo_vuelo_points[i] == modo_vuelo_points[i + 1]:
            modo_vuelo_intervalo = modo_vuelo_points[i]
        else:
            modo_vuelo_intervalo = f"{modo_vuelo_points[i]}->{modo_vuelo_points[i + 1]}"

        if p_net_bus_w >= 0.0:
            modo_bateria = "carga"
            p_storage_internal_w = p_net_bus_w * eta_batt_charge_local
            e_storage_interval_wh = p_storage_internal_w * dt_h
            e_out_dd_wh = 0.0
            e_out_night_wh = 0.0
        else:
            p_storage_internal_w = p_net_bus_w / eta_batt_discharge_local
            e_storage_interval_wh = p_storage_internal_w * dt_h

            if el_mid_deg > 0.0:
                modo_bateria = "descarga_amanecer_atardecer"
                e_out_dd_wh = -e_storage_interval_wh
                e_out_night_wh = 0.0
            else:
                modo_bateria = "descarga_noche"
                e_out_dd_wh = 0.0
                e_out_night_wh = -e_storage_interval_wh

        registros.append({
            "t_ini_h": t_ini,
            "t_fin_h": t_fin,
            "t_mid_h": t_mid,
            "dt_h": dt_h,
            "Elevacion_solar_media_deg": el_mid_deg,
            "Altitud_media_m": alt_mid_m,
            "ROC_media_m_s": roc_mid_mps,
            "Modo_vuelo_intervalo": modo_vuelo_intervalo,
            "P_array_bus_W": p_array_bus_w,
            "P_propulsion_bus_W": p_prop_mid_w,
            "P_sistemas_bus_W": p_sys_mid_w,
            "P_payload_bus_W": p_pl_mid_w,
            "P_total_bus_W": p_total_bus_w,
            "P_neta_bus_W": p_net_bus_w,
            "P_storage_internal_W": p_storage_internal_w,
            "E_storage_interval_Wh": e_storage_interval_wh,
            "E_out_dd_interval_Wh": e_out_dd_wh,
            "E_out_night_interval_Wh": e_out_night_wh,
            "Modo_intervalo": modo_bateria,
        })

    df_batt = pd.DataFrame(registros)
    df_batt["E_storage_cumulative_midnight_ref_Wh"] = df_batt["E_storage_interval_Wh"].cumsum()
    return df_batt


def capacidad_bateria_usable_multidia_wh(df_batt, mission_days):
    dE_list = df_batt["E_storage_interval_Wh"].tolist()
    cumulative_list = [0.0]
    acumulado = 0.0
    for dE in dE_list:
        acumulado += dE
        cumulative_list.append(acumulado)

    idx_inicio_llena = max(range(len(cumulative_list)), key=lambda k: cumulative_list[k])
    dE_rot = dE_list[idx_inicio_llena:] + dE_list[:idx_inicio_llena]

    deuda_wh = 0.0
    capacidad_requerida_wh = 0.0
    historial_deuda = []

    for _ in range(mission_days):
        for dE in dE_rot:
            if dE >= 0.0:
                deuda_wh = max(0.0, deuda_wh - dE)
            else:
                deuda_wh += -dE
                capacidad_requerida_wh = max(capacidad_requerida_wh, deuda_wh)
            historial_deuda.append(deuda_wh)

    hora_inicio_llena_h = 0.0
    if len(df_batt) > 0:
        idx_hora = min(idx_inicio_llena, len(df_batt) - 1)
        hora_inicio_llena_h = float(df_batt.iloc[idx_hora]["t_mid_h"])

    return {
        "capacidad_requerida_wh": capacidad_requerida_wh,
        "idx_inicio_llena": idx_inicio_llena,
        "hora_inicio_llena_h": hora_inicio_llena_h,
        "historial_deuda_wh": historial_deuda,
    }


"""
FUNCIÓN NUEVA: traducir requisitos energéticos del HAPS a masa de batería.

ORIGEN:
- Gundlach Eq. (8.69) / Eq. (3.38): la energía utilizable del pack escala con la
  energía específica de la tecnología y con la masa de batería.
- Gundlach Tabla 8.2: además de energía específica, también da potencia
  específica. Eso permite verificar que el pack no solo "guarda" suficiente energía,
  sino que además puede entregar la potencia pico requerida.

POR QUÉ SE HACE ASÍ:
- Para energía: masa = E_requerida / (E_específica * fracción_utilizable)
- Para potencia: masa = P_pico / P_específica
- La masa final del pack viene impuesta por la más exigente de las dos.

NOTA SOBRE f_usable_batt:
- Esta fracción representa qué parte de la energía nominal del pack estás dispuesto
  a usar realmente en misión. Cuanto menor sea, más conservador será el diseño
  y mayor será la masa resultante.
"""

def diagnostico_viabilidad_preliminar(autosostenible_diario, masa_bateria_kg, mtow_kg, mission_days):
    """
    DIAGNÓSTICO SIMPLE DE VIABILIDAD PRELIMINAR.

    QUÉ SIGNIFICA CADA CRITERIO:
    - Para 1 día:
        * Se comprueba solo si la batería mínima conservadora cabe o no dentro del
          MTOW como criba básica de imposibilidad.
    - Para 30 y 60 días:
        * Además de lo anterior, se exige que el día representativo sea
          autosostenible energéticamente (balance diario de almacenamiento >= 0),
          porque si no lo es, repetir ese día durante 30 o 60 ciclos arrastra un
          déficit acumulativo.

    IMPORTANTE:
    - Esta viabilidad es PRELIMINAR.
    - No sustituye una iteración completa del MTOW del HAPS con el peso real de
      estructura, placas, aviónica, baterías y demás subsistemas.
    """
    if mission_days == 1:
        if masa_bateria_kg <= mtow_kg:
            return 1, "SI (preliminar: la bateria minima cabe dentro del MTOW)"
        return 0, "NO (preliminar: la bateria minima supera el MTOW)"

    if autosostenible_diario < 1:
        return 0, "NO (el dia representativo no se autosostiene energeticamente)"

    if masa_bateria_kg <= mtow_kg:
        return 1, "SI (autosostenible y bateria minima <= MTOW)"
    return 0, "NO (autosostenible, pero la bateria minima supera el MTOW)"



def construir_tabla_consola_principal(df_resultados):
    cols = [
        "Caso",
        "Escenario",
        "MTOW (kg)",
        "P_total_avg (W)",
        "Energia electrica diaria array (Wh)",
        "Balance diario almacenamiento (Wh)",
        "Autosostenible diario (1/0)",
        "Viable_preliminar_30d (1/0)",
        "Viable_preliminar_60d (1/0)",
        "Masa_min_conservadora_30d (kg)",
        "Masa_min_conservadora_60d (kg)",
    ]
    dfc = df_resultados[cols].copy()
    dfc = dfc.rename(columns={
        "MTOW (kg)": "MTOW_kg",
        "P_total_avg (W)": "Ptot_avg_W",
        "Energia electrica diaria array (Wh)": "Earray_dia_Wh",
        "Balance diario almacenamiento (Wh)": "Bal_almac_Wh",
        "Autosostenible diario (1/0)": "Auto_dia",
        "Viable_preliminar_30d (1/0)": "Viable_30d",
        "Viable_preliminar_60d (1/0)": "Viable_60d",
        "Masa_min_conservadora_30d (kg)": "Mbat_30d_kg",
        "Masa_min_conservadora_60d (kg)": "Mbat_60d_kg",
    })
    num_cols = dfc.select_dtypes(include="number").columns
    dfc[num_cols] = dfc[num_cols].round(2)
    return dfc



def construir_tabla_consola_bateria_mejor(df_baterias):
    if df_baterias.empty:
        return df_baterias.copy()

    idx = df_baterias.groupby(["Caso", "Escenario", "Dias_mision"])["Masa_total_conservadora_kg"].idxmin()
    dfc = df_baterias.loc[idx, [
        "Caso",
        "Escenario",
        "Dias_mision",
        "Tecnologia_bateria",
        "E_batt_usable_req (Wh)",
        "Masa_total_conservadora_kg",
        "Fraccion_MTOW_conservadora",
        "Viable_mision_preliminar (1/0)",
        "Diagnostico_viabilidad",
    ]].copy()

    dfc = dfc.sort_values(["Caso", "Escenario", "Dias_mision"]).reset_index(drop=True)
    dfc = dfc.rename(columns={
        "Dias_mision": "Dias",
        "Tecnologia_bateria": "Tec",
        "E_batt_usable_req (Wh)": "Ebat_req_Wh",
        "Masa_total_conservadora_kg": "Mbat_cons_kg",
        "Fraccion_MTOW_conservadora": "Frac_MTOW",
        "Viable_mision_preliminar (1/0)": "Viable",
        "Diagnostico_viabilidad": "Diagnostico",
    })
    num_cols = dfc.select_dtypes(include="number").columns
    dfc[num_cols] = dfc[num_cols].round(2)
    return dfc



def imprimir_diagnostico_resumido(df_resultados):
    print("\nDIAGNOSTICO PRELIMINAR DE VIABILIDAD:\n")
    for _, row in df_resultados.iterrows():
        print(f"{row['Caso']} | {row['Escenario']}")
        print(
            f"  - Balance diario almacenamiento: {row['Balance diario almacenamiento (Wh)']:.2f} Wh "
            f"-> {'autosostenible' if row['Autosostenible diario (1/0)'] == 1 else 'NO autosostenible'}"
        )
        print(
            f"  - 1 dia : {row['Diagnostico_1d']}"
        )
        print(
            f"  - 30 dias: {row['Diagnostico_30d']}"
        )
        print(
            f"  - 60 dias: {row['Diagnostico_60d']}"
        )
        print("")



def evaluar_tecnologia_bateria(
    tecnologia,
    datos_tecnologia,
    E_req_wh,
    P_pico_batt_bus_w,
    MTOW_kg,
    f_usable_batt_local,
):
    espec_min = datos_tecnologia["Espec_min_Whkg"]
    espec_max = datos_tecnologia["Espec_max_Whkg"]
    pspec_min = datos_tecnologia["Pspec_min_Wkg"]
    pspec_max = datos_tecnologia["Pspec_max_Wkg"]

    masa_energia_conservadora_kg = E_req_wh / (espec_min * f_usable_batt_local)
    masa_energia_optimista_kg = E_req_wh / (espec_max * f_usable_batt_local)
    masa_potencia_conservadora_kg = P_pico_batt_bus_w / pspec_min
    masa_potencia_optimista_kg = P_pico_batt_bus_w / pspec_max

    masa_total_conservadora_kg = max(masa_energia_conservadora_kg, masa_potencia_conservadora_kg)
    masa_total_optimista_kg = max(masa_energia_optimista_kg, masa_potencia_optimista_kg)

    return {
        "Tecnologia_bateria": tecnologia,
        "Espec_min_Whkg": espec_min,
        "Espec_max_Whkg": espec_max,
        "Pspec_min_Wkg": pspec_min,
        "Pspec_max_Wkg": pspec_max,
        "Masa_energia_conservadora_kg": masa_energia_conservadora_kg,
        "Masa_energia_optimista_kg": masa_energia_optimista_kg,
        "Masa_potencia_conservadora_kg": masa_potencia_conservadora_kg,
        "Masa_potencia_optimista_kg": masa_potencia_optimista_kg,
        "Masa_total_conservadora_kg": masa_total_conservadora_kg,
        "Masa_total_optimista_kg": masa_total_optimista_kg,
        "Fraccion_MTOW_conservadora": masa_total_conservadora_kg / MTOW_kg,
        "Fraccion_MTOW_optimista": masa_total_optimista_kg / MTOW_kg,
    }



"""
BLOQUE NUEVO: métricas del pack de batería útiles para estudiar ciclos y degradación.

QUÉ HACE:
- A partir de la masa final del pack calculada por energía/potencia, traduce esa
  masa a:
    * energía nominal total del pack,
    * energía de reserva no utilizable,
    * SOC mínimo diario,
    * DoD diario,
    * EFC diarios equivalentes basados en la descarga diaria,
    * C-rate pico equivalente,
    * y limitante principal del pack (energía o potencia).

POR QUÉ ES ÚTIL:
- Con esto ya puedes responder preguntas del tipo:
    "Si el pack pesa X kg, ¿a cuántos kWh nominales equivale?"
    "¿Qué porcentaje del pack estoy ciclando cada día?"
    "¿Cuánta reserva real dejo sin tocar si f_usable_batt = 0.80?"
- Esto no selecciona por sí solo una química final; solo traduce a magnitudes
  más interpretables el tamaño de pack que ya sale del modelo conceptual.

NOTA SOBRE C-RATE:
- Sin fijar todavía tensión nominal de pack ni arquitectura serie/paralelo, el
  script calcula un C-rate pico equivalente aproximado como P_pico / E_nominal.
- No es un cálculo electroquímico detallado, pero sí una señal conceptual útil.
"""


def enriquecer_metricas_pack_bateria(
    masa_total_conservadora_kg,
    masa_energia_conservadora_kg,
    masa_potencia_conservadora_kg,
    espec_min_whkg,
    f_usable_batt_local,
    E_req_1dia_wh,
    E_req_mision_wh,
    e_descarga_diaria_wh,
    P_pico_batt_bus_w,
):
    e_nominal_pack_wh = masa_total_conservadora_kg * espec_min_whkg
    e_usable_pack_wh = e_nominal_pack_wh * f_usable_batt_local
    e_reserva_pack_wh = e_nominal_pack_wh - e_usable_pack_wh

    dod_diario_total_pct = 100.0 * E_req_1dia_wh / e_nominal_pack_wh if e_nominal_pack_wh > 0.0 else 0.0
    dod_mision_total_pct = 100.0 * E_req_mision_wh / e_nominal_pack_wh if e_nominal_pack_wh > 0.0 else 0.0

    uso_ventana_usable_diario_pct = 100.0 * E_req_1dia_wh / e_usable_pack_wh if e_usable_pack_wh > 0.0 else 0.0
    uso_ventana_usable_mision_pct = 100.0 * E_req_mision_wh / e_usable_pack_wh if e_usable_pack_wh > 0.0 else 0.0

    soc_min_diario_pct = max(0.0, 100.0 - dod_diario_total_pct)
    soc_min_mision_pct = max(0.0, 100.0 - dod_mision_total_pct)

    efc_diarios_equiv = e_descarga_diaria_wh / e_nominal_pack_wh if e_nominal_pack_wh > 0.0 else 0.0
    c_rate_pico_equiv = P_pico_batt_bus_w / e_nominal_pack_wh if e_nominal_pack_wh > 0.0 else 0.0

    if abs(masa_energia_conservadora_kg - masa_potencia_conservadora_kg) <= 1e-9:
        limitante_pack = "Energia y potencia"
    elif masa_energia_conservadora_kg > masa_potencia_conservadora_kg:
        limitante_pack = "Energia"
    else:
        limitante_pack = "Potencia"

    return {
        "E_nominal_pack_conservadora_Wh": e_nominal_pack_wh,
        "E_usable_pack_conservadora_Wh": e_usable_pack_wh,
        "E_reserva_pack_Wh": e_reserva_pack_wh,
        "DoD_diario_total_pct": dod_diario_total_pct,
        "DoD_mision_total_pct": dod_mision_total_pct,
        "Uso_ventana_usable_diario_pct": uso_ventana_usable_diario_pct,
        "Uso_ventana_usable_mision_pct": uso_ventana_usable_mision_pct,
        "SOC_min_diario_pct": soc_min_diario_pct,
        "SOC_min_mision_pct": soc_min_mision_pct,
        "EFC_diarios_equiv": efc_diarios_equiv,
        "C_rate_pico_equiv_h_inv": c_rate_pico_equiv,
        "Limitante_pack": limitante_pack,
    }


def construir_historial_deuda_desde_bateria_llena(df_batt, mission_days):
    """
    PERFIL AUXILIAR PARA GRÁFICAS DE SOC / DEUDA.

    QUÉ HACE:
    - Reordena el día representativo para que el tiempo t = 0 coincida con el
      instante en que la batería está más llena.
    - A partir de ahí acumula la deuda energética interna durante 1, 30 o 60 días.

    POR QUÉ SE USA:
    - Así las gráficas de SOC y de deuda salen continuas y con una lectura física
      más clara: empiezan en el instante "mejor cargado" del ciclo.
    """
    salida = capacidad_bateria_usable_multidia_wh(df_batt, mission_days)
    idx_inicio = int(salida["idx_inicio_llena"])

    if len(df_batt) == 0:
        return pd.DataFrame({"t_rel_h": [0.0], "deuda_wh": [0.0]})

    df_rot = pd.concat([df_batt.iloc[idx_inicio:], df_batt.iloc[:idx_inicio]], ignore_index=True)

    tiempos = [0.0]
    deudas = [0.0]
    deuda = 0.0
    t_acum = 0.0

    for _ in range(mission_days):
        for _, row in df_rot.iterrows():
            dE = float(row["E_storage_interval_Wh"])
            dt_h = float(row["dt_h"])

            if dE >= 0.0:
                deuda = max(0.0, deuda - dE)
            else:
                deuda += -dE

            t_acum += dt_h
            tiempos.append(t_acum)
            deudas.append(deuda)

    return pd.DataFrame({"t_rel_h": tiempos, "deuda_wh": deudas})



# -----------------------------------------------------------------------------
# 4. CÁLCULO DE CADA CASO HAPS Y DE CADA ESCENARIO OPERATIVO
# -----------------------------------------------------------------------------
rho_crucero = atmosfera_isa(h_crucero)
El_horiz_rad = horizon_elevation_rad(h_crucero / 1000.0, h_cloud_km)

resultados = []
resultados_baterias = []
perfiles_solares = {}
perfiles_operacion = {}
perfiles_bateria = {}

print("--- COMPARACIÓN DE 4 CASOS HAPS CON BATERÍAS Y ESTRATEGIA DE PLANEO ---")
print(f"Densidad ISA a {h_crucero:.0f} m: {rho_crucero:.5f} kg/m^3")
print(f"eta_prop_global = {eta_prop_global:.2f}")
print(f"eta_array = {eta_array*100:.2f} %")
print(f"Cobertura ala = {f_cobertura_max*100:.0f} %")
print(f"P_sistemas_const_W = {P_sistemas_const_W:.2f} W")
print(f"P_payload_const_W = {P_payload_const_W:.2f} W")
print(f"eta_batt_charge = {eta_batt_charge:.2f}")
print(f"eta_batt_discharge = {eta_batt_discharge:.2f}")
print(f"f_usable_batt = {f_usable_batt:.2f}")
print()

for nombre_caso, datos_caso in casos_haps.items():
    W_TO = datos_caso["W_TO"]
    b = datos_caso["b"]
    day_of_year = datos_caso["day_of_year"]
    days_since_vernal = datos_caso["days_since_vernal"]

    geom = resolver_geometria_crucero_base(W_TO=W_TO, b=b, rho_design=rho_crucero)
    W_N = geom["W_N"]
    AR = geom["AR"]
    K = geom["K"]
    S_alar = geom["S_alar"]

    # --- Solar base del caso (común a los dos escenarios operativos) ---
    S_ala_extrados_geom = S_alar
    S_array_max = (S_ala_extrados_geom + S_cola_extra) * f_cobertura_max

    decl_rad_local = declination_rad(days_since_vernal)
    I_space_day = extraterrestrial_irradiance(day_of_year)

    v_sr_alt = sunrise_hour_angle_rad(lat_rad, decl_rad_local, El_horiz_rad)
    t_sr_alt = sunrise_time_from_hour_angle(v_sr_alt)
    t_ss_alt = sunset_time_from_sunrise(t_sr_alt)
    t_day_visible = t_ss_alt - t_sr_alt

    v_sr_panel = sunrise_hour_angle_rad(lat_rad, decl_rad_local, 0.0)
    t_sr_panel = sunrise_time_from_hour_angle(v_sr_panel)
    t_ss_panel = sunset_time_from_sunrise(t_sr_panel)
    t_day_panel = t_ss_panel - t_sr_panel
    t_night_panel = 24.0 - t_day_panel

    times_h = []
    el_list_deg = []
    i_panel_wm2 = []
    p_array_incident_w = []
    p_array_electric_w = []

    n_steps = int(round(24.0 * 60.0 / dt_min)) + 1
    for k in range(n_steps):
        t_h = k * dt_min / 60.0
        if t_h > 24.0:
            t_h = 24.0

        i_panel, el_rad = direct_irradiance_on_horizontal_panel(day_of_year, decl_rad_local, t_h)
        p_incident = i_panel * S_array_max
        p_electric = p_incident * eta_array

        times_h.append(t_h)
        el_list_deg.append(rad2deg(el_rad))
        i_panel_wm2.append(i_panel)
        p_array_incident_w.append(p_incident)
        p_array_electric_w.append(p_electric)

    E_incident_wh_m2 = integrar_energia_wh(times_h, i_panel_wm2)
    E_incident_wh_array = integrar_energia_wh(times_h, p_array_incident_w)
    E_electric_wh_array = integrar_energia_wh(times_h, p_array_electric_w)
    Pmax_incident_w_m2 = max(i_panel_wm2)
    Pmax_incident_array_w = max(p_array_incident_w)
    Pmax_electric_array_w = max(p_array_electric_w)

    perfiles_solares[nombre_caso] = pd.DataFrame({
        "Tiempo solar (h)": times_h,
        "Elevacion solar (deg)": el_list_deg,
        "Irradiancia directa en panel horizontal (W/m2)": i_panel_wm2,
        "Potencia solar incidente en array (W)": p_array_incident_w,
        "Potencia electrica del array (W)": p_array_electric_w,
    })

    p_sistemas_profile = construir_perfil_constante(times_h, P_sistemas_const_W)
    p_payload_profile = construir_perfil_constante(times_h, P_payload_const_W)

    E_sistemas_24h = integrar_energia_wh(times_h, p_sistemas_profile)
    E_payload_24h = integrar_energia_wh(times_h, p_payload_profile)
    P_payload_avg_equiv = E_payload_24h / 24.0

    for nombre_escenario, datos_escenario in escenarios_operativos.items():
        perfil_oper = construir_perfil_operativo_propulsion(
            times_h=times_h,
            t_sr_panel=t_sr_panel,
            t_ss_panel=t_ss_panel,
            W_N=W_N,
            S_alar=S_alar,
            K=K,
            escenario_dict=datos_escenario,
        )

        p_propulsion_profile = perfil_oper["p_elec_w"]
        p_total_profile = [
            p_propulsion_profile[i] + p_sistemas_profile[i] + p_payload_profile[i]
            for i in range(len(times_h))
        ]

        E_propulsion_24h = integrar_energia_wh(times_h, p_propulsion_profile)
        E_total_24h = integrar_energia_wh(times_h, p_total_profile)
        P_propulsion_avg = E_propulsion_24h / 24.0
        P_total_avg = E_total_24h / 24.0
        margen_diario_wh = E_electric_wh_array - E_total_24h

        """
        BLOQUE NUEVO AMPLIADO: balance detallado de batería con potencia total variable.

        DIFERENCIA RESPECTO AL BLOQUE ANTERIOR:
        - Ahora la demanda del bus ya no es constante, porque:
            1) la potencia propulsiva cambia si se aplica el ciclo de descenso
               nocturno y reascenso diurno,
            2) la payload de cámara solo está activa en una parte del día.
        - Esto mejora la fidelidad física del balance diario frente a asumir una
          carga uniforme de 24 horas.

        CRITERIO DE VIABILIDAD SOLAR MULTIDIURNA:
        - Si el balance diario de almacenamiento sale no negativo, entonces el día
          representativo puede repetirse indefinidamente en este modelo sin agotar
          la batería, y por tanto la misión de 1-2 meses es energéticamente
          autosostenible en crucero.
        - Si sale negativo, el escenario no cierra energéticamente día a día y la
          capacidad requerida crecerá con el número de días.
        """
        df_batt = construir_perfil_bateria_intervalos(
            times_h=times_h,
            el_list_deg=el_list_deg,
            p_array_electric_w=p_array_electric_w,
            p_propulsion_bus_w_points=p_propulsion_profile,
            p_sistemas_bus_w_points=p_sistemas_profile,
            p_payload_bus_w_points=p_payload_profile,
            altitud_m_points=perfil_oper["altitud_m"],
            roc_mps_points=perfil_oper["roc_mps"],
            modo_vuelo_points=perfil_oper["modo_vuelo"],
            eta_power_conditioning_local=eta_power_conditioning,
            eta_batt_charge_local=eta_batt_charge,
            eta_batt_discharge_local=eta_batt_discharge,
        )

        E_stor_day_internal_Wh = float(df_batt.loc[df_batt["E_storage_interval_Wh"] > 0.0, "E_storage_interval_Wh"].sum())
        E_out_dd_Wh = float(df_batt["E_out_dd_interval_Wh"].sum())
        E_out_night_Wh = float(df_batt["E_out_night_interval_Wh"].sum())
        E_out_total_Wh = E_out_dd_Wh + E_out_night_Wh
        balance_storage_day_Wh = E_stor_day_internal_Wh - E_out_total_Wh

        P_pico_batt_bus_W = float((-df_batt["P_neta_bus_W"]).clip(lower=0.0).max())
        P_pico_batt_internal_W = float((-df_batt["P_storage_internal_W"]).clip(lower=0.0).max())

        capacidad_dict = {}
        hora_bateria_llena_h = None
        for mission_days in mission_days_list:
            salida_cap = capacidad_bateria_usable_multidia_wh(df_batt, mission_days)
            capacidad_dict[mission_days] = salida_cap["capacidad_requerida_wh"]
            if hora_bateria_llena_h is None:
                hora_bateria_llena_h = salida_cap["hora_inicio_llena_h"]

        nombre_perfil = f"{nombre_caso}__{nombre_escenario}"
        perfiles_operacion[nombre_perfil] = pd.DataFrame({
            "Tiempo solar (h)": times_h,
            "Altitud (m)": perfil_oper["altitud_m"],
            "Densidad aire (kg/m3)": perfil_oper["rho_kgm3"],
            "ROC (m/s)": perfil_oper["roc_mps"],
            "CL (-)": perfil_oper["cl"],
            "CD (-)": perfil_oper["cd"],
            "L/D (-)": perfil_oper["ld"],
            "Drag (N)": perfil_oper["drag_n"],
            "Thrust requerido (N)": perfil_oper["thrust_n"],
            "P_thrust requerida (W)": perfil_oper["p_thrust_w"],
            "P_elec_propulsion (W)": perfil_oper["p_elec_w"],
            "P_sistemas (W)": p_sistemas_profile,
            "P_payload (W)": p_payload_profile,
            "P_total_bus (W)": p_total_profile,
            "Modo_vuelo": perfil_oper["modo_vuelo"],
        })
        perfiles_bateria[nombre_perfil] = df_batt.copy()

        filas_batt_locales = []
        for tecnologia, datos_tecnologia in battery_technologies.items():
            fila_base = {
                "Caso": nombre_caso,
                "Escenario": nombre_escenario,
                "MTOW (kg)": W_TO,
                "P_total_avg (W)": P_total_avg,
                "Pico potencia bateria bus (W)": P_pico_batt_bus_W,
                "Pico potencia bateria interna (W)": P_pico_batt_internal_W,
                "Balance diario almacenamiento (Wh)": balance_storage_day_Wh,
                "Autosostenible diario (1/0)": int(balance_storage_day_Wh >= 0.0),
                "Hora bateria llena (h solar)": hora_bateria_llena_h,
                "f_usable_batt": f_usable_batt,
                "eta_batt_charge": eta_batt_charge,
                "eta_batt_discharge": eta_batt_discharge,
                "E_batt_usable_1dia_ref (Wh)": capacidad_dict[1],
                "Throughput_diario_batt_Wh": E_stor_day_internal_Wh + E_out_total_Wh,
                "E_out_total_batt_Wh": E_out_total_Wh,
            }

            for mission_days in mission_days_list:
                E_req_wh = capacidad_dict[mission_days]
                eval_batt = evaluar_tecnologia_bateria(
                    tecnologia=tecnologia,
                    datos_tecnologia=datos_tecnologia,
                    E_req_wh=E_req_wh,
                    P_pico_batt_bus_w=P_pico_batt_bus_W,
                    MTOW_kg=W_TO,
                    f_usable_batt_local=f_usable_batt,
                )

                viable_tecnologia_flag, diagnostico_tecnologia = diagnostico_viabilidad_preliminar(
                    autosostenible_diario=int(balance_storage_day_Wh >= 0.0),
                    masa_bateria_kg=eval_batt["Masa_total_conservadora_kg"],
                    mtow_kg=W_TO,
                    mission_days=mission_days,
                )

                fila = fila_base.copy()
                fila.update({
                    "Tecnologia_bateria": tecnologia,
                    "Dias_mision": mission_days,
                    "E_batt_usable_req (Wh)": E_req_wh,
                    "Espec_min_Whkg": eval_batt["Espec_min_Whkg"],
                    "Espec_max_Whkg": eval_batt["Espec_max_Whkg"],
                    "Pspec_min_Wkg": eval_batt["Pspec_min_Wkg"],
                    "Pspec_max_Wkg": eval_batt["Pspec_max_Wkg"],
                    "Masa_energia_conservadora_kg": eval_batt["Masa_energia_conservadora_kg"],
                    "Masa_energia_optimista_kg": eval_batt["Masa_energia_optimista_kg"],
                    "Masa_potencia_conservadora_kg": eval_batt["Masa_potencia_conservadora_kg"],
                    "Masa_potencia_optimista_kg": eval_batt["Masa_potencia_optimista_kg"],
                    "Masa_total_conservadora_kg": eval_batt["Masa_total_conservadora_kg"],
                    "Masa_total_optimista_kg": eval_batt["Masa_total_optimista_kg"],
                    "Fraccion_MTOW_conservadora": eval_batt["Fraccion_MTOW_conservadora"],
                    "Fraccion_MTOW_optimista": eval_batt["Fraccion_MTOW_optimista"],
                    "Criba_masa_pack_leq_MTOW (1/0)": int(eval_batt["Masa_total_conservadora_kg"] <= W_TO),
                    "Viable_mision_preliminar (1/0)": viable_tecnologia_flag,
                    "Diagnostico_viabilidad": diagnostico_tecnologia,
                })
                fila.update(
                    enriquecer_metricas_pack_bateria(
                        masa_total_conservadora_kg=eval_batt["Masa_total_conservadora_kg"],
                        masa_energia_conservadora_kg=eval_batt["Masa_energia_conservadora_kg"],
                        masa_potencia_conservadora_kg=eval_batt["Masa_potencia_conservadora_kg"],
                        espec_min_whkg=eval_batt["Espec_min_Whkg"],
                        f_usable_batt_local=f_usable_batt,
                        E_req_1dia_wh=capacidad_dict[1],
                        E_req_mision_wh=E_req_wh,
                        e_descarga_diaria_wh=E_out_total_Wh,
                        P_pico_batt_bus_w=P_pico_batt_bus_W,
                    )
                )
                resultados_baterias.append(fila)
                filas_batt_locales.append(fila)

        minimos_por_dia = {}
        for mission_days in mission_days_list:
            subset = [fila for fila in filas_batt_locales if fila["Dias_mision"] == mission_days]
            mejor = min(subset, key=lambda x: x["Masa_total_conservadora_kg"])
            minimos_por_dia[mission_days] = mejor

        autosostenible_diario_flag = int(balance_storage_day_Wh >= 0.0)

        viable_1d_flag, diagnostico_1d = diagnostico_viabilidad_preliminar(
            autosostenible_diario=autosostenible_diario_flag,
            masa_bateria_kg=minimos_por_dia[1]["Masa_total_conservadora_kg"],
            mtow_kg=W_TO,
            mission_days=1,
        )
        viable_30d_flag, diagnostico_30d = diagnostico_viabilidad_preliminar(
            autosostenible_diario=autosostenible_diario_flag,
            masa_bateria_kg=minimos_por_dia[30]["Masa_total_conservadora_kg"],
            mtow_kg=W_TO,
            mission_days=30,
        )
        viable_60d_flag, diagnostico_60d = diagnostico_viabilidad_preliminar(
            autosostenible_diario=autosostenible_diario_flag,
            masa_bateria_kg=minimos_por_dia[60]["Masa_total_conservadora_kg"],
            mtow_kg=W_TO,
            mission_days=60,
        )

        resultados.append({
            "Caso": nombre_caso,
            "Escenario": nombre_escenario,
            "MTOW (kg)": W_TO,
            "b (m)": b,
            "DOY": day_of_year,
            "AR (-)": AR,
            "S_alar (m2)": S_alar,
            "CL_diseno_20km (-)": geom["CL_design"],
            "CD_diseno_20km (-)": geom["CD_design"],
            "E_diseno_20km = CL/CD (-)": geom["E_aero_design"],
            "T_crucero_20km_base (N)": geom["T_crucero_design"],
            "P_thrust_crucero_20km_base (W)": geom["P_thrust_crucero_design"],
            "P_elec_crucero_20km_base (W)": geom["P_elec_crucero_design"],
            "S_array_max (m2)": S_array_max,
            "Declinacion (deg)": rad2deg(decl_rad_local),
            "I_espacio (W/m2)": I_space_day,
            "Horas visibles a altitud (h)": t_day_visible,
            "Horas utiles panel horizontal (h)": t_day_panel,
            "Horas nocturnas (h)": t_night_panel,
            "Pico electrico array (W)": Pmax_electric_array_w,
            "Energia electrica diaria array (Wh)": E_electric_wh_array,
            "P_sistemas_const (W)": P_sistemas_const_W,
            "P_payload_const (W)": P_payload_const_W,
            "P_payload_avg_equiv (W)": P_payload_avg_equiv,
            "E_payload_24h (Wh)": E_payload_24h,
            "P_propulsion_avg (W)": P_propulsion_avg,
            "P_total_avg (W)": P_total_avg,
            "Consumo total 24h (Wh)": E_total_24h,
            "Margen diario simple (Wh)": margen_diario_wh,
            "Altitud minima operacion (m)": min(perfil_oper["altitud_m"]),
            "Altitud maxima operacion (m)": max(perfil_oper["altitud_m"]),
            "ROC_min (m/s)": min(perfil_oper["roc_mps"]),
            "ROC_max (m/s)": max(perfil_oper["roc_mps"]),
            "P_propulsion_min (W)": min(perfil_oper["p_elec_w"]),
            "P_propulsion_max (W)": max(perfil_oper["p_elec_w"]),
            "T_req_min_perfil (N)": min(perfil_oper["thrust_n"]),
            "T_req_max_perfil (N)": max(perfil_oper["thrust_n"]),
            "P_thrust_min_perfil (W)": min(perfil_oper["p_thrust_w"]),
            "P_thrust_max_perfil (W)": max(perfil_oper["p_thrust_w"]),
            "E_stor_dia_batt (Wh)": E_stor_day_internal_Wh,
            "E_out_dd_batt (Wh)": E_out_dd_Wh,
            "E_out_noche_batt (Wh)": E_out_night_Wh,
            "Balance diario almacenamiento (Wh)": balance_storage_day_Wh,
            "Pico potencia bateria bus (W)": P_pico_batt_bus_W,
            "Pico potencia bateria interna (W)": P_pico_batt_internal_W,
            "Hora bateria llena (h solar)": hora_bateria_llena_h,
            "E_batt_usable_1dia (Wh)": capacidad_dict[1],
            "E_batt_usable_30dias (Wh)": capacidad_dict[30],
            "E_batt_usable_60dias (Wh)": capacidad_dict[60],
            "Autosostenible diario (1/0)": autosostenible_diario_flag,
            "Viable_preliminar_1d (1/0)": viable_1d_flag,
            "Viable_preliminar_30d (1/0)": viable_30d_flag,
            "Viable_preliminar_60d (1/0)": viable_60d_flag,
            "Diagnostico_1d": diagnostico_1d,
            "Diagnostico_30d": diagnostico_30d,
            "Diagnostico_60d": diagnostico_60d,
            "Tecnologia_min_masa_1d": minimos_por_dia[1]["Tecnologia_bateria"],
            "Masa_min_conservadora_1d (kg)": minimos_por_dia[1]["Masa_total_conservadora_kg"],
            "Tecnologia_min_masa_30d": minimos_por_dia[30]["Tecnologia_bateria"],
            "Masa_min_conservadora_30d (kg)": minimos_por_dia[30]["Masa_total_conservadora_kg"],
            "Frac_MTOW_min_30d (-)": minimos_por_dia[30]["Fraccion_MTOW_conservadora"],
            "E_nominal_pack_30d (Wh)": minimos_por_dia[30]["E_nominal_pack_conservadora_Wh"],
            "E_usable_pack_30d (Wh)": minimos_por_dia[30]["E_usable_pack_conservadora_Wh"],
            "E_reserva_pack_30d (Wh)": minimos_por_dia[30]["E_reserva_pack_Wh"],
            "DoD_diario_pack_30d (%)": minimos_por_dia[30]["DoD_diario_total_pct"],
            "SOC_min_diario_pack_30d (%)": minimos_por_dia[30]["SOC_min_diario_pct"],
            "EFC_diarios_pack_30d (-)": minimos_por_dia[30]["EFC_diarios_equiv"],
            "C_rate_pico_eq_pack_30d (-)": minimos_por_dia[30]["C_rate_pico_equiv_h_inv"],
            "Limitante_pack_30d": minimos_por_dia[30]["Limitante_pack"],
            "Tecnologia_min_masa_60d": minimos_por_dia[60]["Tecnologia_bateria"],
            "Masa_min_conservadora_60d (kg)": minimos_por_dia[60]["Masa_total_conservadora_kg"],
            "Frac_MTOW_min_60d (-)": minimos_por_dia[60]["Fraccion_MTOW_conservadora"],
            "E_nominal_pack_60d (Wh)": minimos_por_dia[60]["E_nominal_pack_conservadora_Wh"],
            "E_usable_pack_60d (Wh)": minimos_por_dia[60]["E_usable_pack_conservadora_Wh"],
            "E_reserva_pack_60d (Wh)": minimos_por_dia[60]["E_reserva_pack_Wh"],
            "DoD_diario_pack_60d (%)": minimos_por_dia[60]["DoD_diario_total_pct"],
            "SOC_min_diario_pack_60d (%)": minimos_por_dia[60]["SOC_min_diario_pct"],
            "EFC_diarios_pack_60d (-)": minimos_por_dia[60]["EFC_diarios_equiv"],
            "C_rate_pico_eq_pack_60d (-)": minimos_por_dia[60]["C_rate_pico_equiv_h_inv"],
            "Limitante_pack_60d": minimos_por_dia[60]["Limitante_pack"],
            "Criba_masa_min_30d_leq_MTOW (1/0)": int(minimos_por_dia[30]["Masa_total_conservadora_kg"] <= W_TO),
            "Criba_masa_min_60d_leq_MTOW (1/0)": int(minimos_por_dia[60]["Masa_total_conservadora_kg"] <= W_TO),
        })

# -----------------------------------------------------------------------------
# 5. TABLAS DE RESULTADOS, POSTPROCESADO Y SIZING INICIAL
# -----------------------------------------------------------------------------
df_resultados = pd.DataFrame(resultados)
df_baterias = pd.DataFrame(resultados_baterias)

# -----------------------------------------------------------------------------
# 5A bis. DIMENSIONADO PARAMÉTRICO PRELIMINAR DE FUSELAJE Y COLA EN T
# -----------------------------------------------------------------------------
"""
BLOQUE NUEVO: dimensionado conceptual de fuselaje esbelto + cola en T.

QUÉ HACE:
- Usa los resultados ya calculados del ala: b, S, AR y cuerda media.
- Calcula una longitud preliminar de fuselaje/boom por criterio de brazo de cola.
- Dimensiona cola horizontal y vertical mediante coeficientes de volumen de cola.
- Dimensiona la cápsula delantera a partir de envolventes internas de payload,
  aviónica y electrónica central.
- Genera una tabla de secciones preliminares para meter el fuselaje manualmente
  en XFLR5 como cuerpo por estaciones.

IMPORTANTE:
- Es una Iteración 0. No sustituye un cálculo completo de estabilidad ni CFD.
- Los coeficientes son hipótesis editables, pensadas para justificar una primera
  geometría y poder hacer sensibilidad en Spyder.
"""

# --- Hipótesis editables de arquitectura / estabilidad preliminar ---
hipotesis_fuselaje_cola = {
    # Longitud mínima del fuselaje/boom. Se toma el máximo entre ambos criterios.
    "k_Lf_sobre_b": 0.35,          # Lf >= 0.35*b, criterio preliminar para UAV HALE/MALE de gran AR
    "k_Lf_sobre_cbar": 3.00,      # Lf >= 3*c_bar, evita brazos de cola demasiado cortos

    # Brazos de cola respecto a la longitud total de fuselaje.
    "fraccion_brazo_cola_H_sobre_Lf": 0.65,
    "fraccion_brazo_cola_V_sobre_Lf": 0.60,
    "x_tail_ac_frac_Lf": 0.93,    # posición aproximada del AC de la cola desde la nariz
    "x_ac_ala_sobre_cuerda": 0.25,

    # Coeficientes de volumen de cola.
    "VH_base": 0.45,              # coeficiente de volumen horizontal preliminar
    "factor_cola_T": 0.95,        # pequeña reducción por cola en T / flujo más limpio / end-plate effect
    "AR_H": 4.50,                 # alargamiento preliminar cola horizontal
    "AR_V": 1.50,                 # alargamiento preliminar cola vertical

    # Cápsula / fuselaje delantero.
    "fineness_ratio_capsula": 4.50,       # L_capsula / D_eq; cuerpo carenado moderadamente esbelto
    "factor_volumetrico_capsula": 0.65,   # V ~= k_v*Amax*L para forma carenada, no caja
    "eta_volumen_util_interno": 0.65,     # fracción del volumen externo realmente utilizable
    "margen_integracion_interno": 0.30,   # margen para soportes, cableado, accesibilidad, curvatura, etc.
    "relacion_H_W_capsula": 0.85,         # sección elíptica: H = 0.85*W

    # Boom trasero.
    "factor_diam_boom_sobre_Deq": 0.50,
    "diam_boom_min_m": 0.12,

    # Ventana inicial de CG respecto a la cuerda media del ala.
    "cg_min_sobre_cuerda": 0.25,
    "cg_max_sobre_cuerda": 0.35,
}

# --- Envolventes internas preliminares: cambia estas cajas si cambias equipos ---
# Formato: (longitud, anchura, altura) en metros.
envolventes_componentes_m = {
    "payload": (0.40, 0.25, 0.20),              # 20 L aprox.
    "avionica": (0.30, 0.18, 0.12),             # 6.5 L aprox.
    "electronica_central": (0.25, 0.16, 0.12),  # 4.8 L aprox.
}


def volumen_caja_litros(dimensiones_m):
    L, W, H = dimensiones_m
    return L * W * H * 1000.0


def coeficiente_volumen_vertical_hale(AR_ala):
    """
    Coeficiente vertical reducido para alas de gran alargamiento.

    Hipótesis simple:
    - AR <= 15  -> CVT = 0.015
    - AR >= 25  -> CVT = 0.009
    - entre medias, interpolación lineal.

    Esto evita que el área de deriva salga desproporcionada cuando el ala tiene
    mucha envergadura y el coeficiente vertical clásico se aplica sin adaptar.
    """
    if AR_ala <= 15.0:
        return 0.015
    if AR_ala >= 25.0:
        return 0.009
    return 0.015 + (AR_ala - 15.0) * (0.009 - 0.015) / (25.0 - 15.0)


def calcular_dimensionado_fuselaje_cola(row, hip, envolventes):
    """
    Calcula una geometría preliminar de fuselaje + cola para una fila de df_resultados.
    """
    caso = row["Caso"]
    escenario = row["Escenario"]
    mtow = float(row["MTOW (kg)"])
    b = float(row["b (m)"])
    S_w = float(row["S_alar (m2)"])
    AR_w = float(row["AR (-)"])
    c_bar = S_w / b

    # 1) Longitud preliminar por brazo de cola.
    Lf_por_b = hip["k_Lf_sobre_b"] * b
    Lf_por_c = hip["k_Lf_sobre_cbar"] * c_bar
    L_f = max(Lf_por_b, Lf_por_c)

    # 2) Brazos de cola.
    l_H = hip["fraccion_brazo_cola_H_sobre_Lf"] * L_f
    l_V = hip["fraccion_brazo_cola_V_sobre_Lf"] * L_f

    # 3) Cola horizontal por coeficiente de volumen.
    V_H_usado = hip["VH_base"] * hip["factor_cola_T"]
    S_H = V_H_usado * S_w * c_bar / l_H
    b_H = math.sqrt(hip["AR_H"] * S_H)
    c_H = S_H / b_H

    # 4) Cola vertical por coeficiente de volumen adaptado a alas de gran AR.
    C_VT_base = coeficiente_volumen_vertical_hale(AR_w)
    C_VT_usado = C_VT_base * hip["factor_cola_T"]
    S_V = C_VT_usado * S_w * b / l_V
    h_V = math.sqrt(hip["AR_V"] * S_V)
    c_V = S_V / h_V

    # 5) Volumen interno útil reservado en cápsula.
    V_componentes_L = sum(volumen_caja_litros(dim) for dim in envolventes.values())
    V_util_int_L = V_componentes_L * (1.0 + hip["margen_integracion_interno"])
    V_ext_capsula_m3 = (V_util_int_L / 1000.0) / hip["eta_volumen_util_interno"]

    # 6) Cápsula carenada. Se resuelve con fineness ratio y factor volumétrico.
    FR = hip["fineness_ratio_capsula"]
    k_v = hip["factor_volumetrico_capsula"]
    D_eq_capsula = (V_ext_capsula_m3 / (k_v * (math.pi / 4.0) * FR)) ** (1.0 / 3.0)
    L_capsula = FR * D_eq_capsula

    H_sobre_W = hip["relacion_H_W_capsula"]
    W_capsula = math.sqrt((D_eq_capsula ** 2) / H_sobre_W)
    H_capsula = H_sobre_W * W_capsula
    A_frontal_capsula = math.pi * W_capsula * H_capsula / 4.0

    # 7) Boom trasero fino.
    D_boom = max(hip["diam_boom_min_m"], hip["factor_diam_boom_sobre_Deq"] * D_eq_capsula)
    L_boom = max(0.0, L_f - L_capsula)

    # 8) Posición longitudinal inicial de ala y CG.
    x_tail_ac = hip["x_tail_ac_frac_Lf"] * L_f
    x_ac_ala = x_tail_ac - l_H
    x_le_ala = x_ac_ala - hip["x_ac_ala_sobre_cuerda"] * c_bar

    # Borde de salida aproximado del ala.
    # En esta fase se aproxima usando la cuerda media c_bar.
    x_te_ala = x_le_ala + c_bar

    x_cg_min = x_le_ala + hip["cg_min_sobre_cuerda"] * c_bar
    x_cg_max = x_le_ala + hip["cg_max_sobre_cuerda"] * c_bar

    # 9) Tramos longitudinales auxiliares para interpretar el fuselaje/boom.
    # L_boom_aprox = L_f - L_capsula incluye también el tramo que atraviesa la zona del ala.
    L_tramo_capsula_a_LE = max(0.0, x_le_ala - L_capsula)
    L_tramo_bajo_ala = max(0.0, min(L_f, x_te_ala) - max(L_capsula, x_le_ala))
    L_boom_detras_ala = max(0.0, L_f - x_te_ala)
    L_TE_ala_a_tail_AC = max(0.0, x_tail_ac - x_te_ala)

    return {
        "Caso": caso,
        "Escenario": escenario,
        "MTOW (kg)": mtow,
        "b_ref (m)": b,
        "S_ref (m2)": S_w,
        "AR_ref (-)": AR_w,
        "c_bar_ref (m)": c_bar,
        "Lf_por_0p35b (m)": Lf_por_b,
        "Lf_por_3cbar (m)": Lf_por_c,
        "Lf_fuselaje_total (m)": L_f,
        "lH_brazo_cola_H (m)": l_H,
        "lV_brazo_cola_V (m)": l_V,
        "VH_usado (-)": V_H_usado,
        "S_H_cola_horizontal (m2)": S_H,
        "AR_H_usado (-)": hip["AR_H"],
        "b_H_cola_horizontal (m)": b_H,
        "c_H_media (m)": c_H,
        "CVT_usado (-)": C_VT_usado,
        "S_V_cola_vertical (m2)": S_V,
        "AR_V_usado (-)": hip["AR_V"],
        "h_V_deriva (m)": h_V,
        "c_V_media (m)": c_V,
        "V_componentes_sin_margen (L)": V_componentes_L,
        "V_util_interno_con_margen (L)": V_util_int_L,
        "V_ext_capsula (L)": V_ext_capsula_m3 * 1000.0,
        "V_geom_capsula_bruto (L)": V_ext_capsula_m3 * 1000.0,
        "D_eq_capsula (m)": D_eq_capsula,
        "L_capsula (m)": L_capsula,
        "W_capsula_max (m)": W_capsula,
        "H_capsula_max (m)": H_capsula,
        "A_frontal_capsula (m2)": A_frontal_capsula,
        "D_boom_equiv (m)": D_boom,
        "L_boom_aprox (m)": L_boom,
        "x_tail_AC_desde_nariz (m)": x_tail_ac,
        "x_ala_AC_desde_nariz (m)": x_ac_ala,
        "x_ala_LE_desde_nariz (m)": x_le_ala,
        "x_ala_TE_aprox_desde_nariz (m)": x_te_ala,
        "L_tramo_capsula_a_LE (m)": L_tramo_capsula_a_LE,
        "L_tramo_bajo_ala (m)": L_tramo_bajo_ala,
        "L_boom_detras_ala (m)": L_boom_detras_ala,
        "L_TE_ala_a_tail_AC (m)": L_TE_ala_a_tail_AC,
        "x_CG_obj_min_desde_nariz (m)": x_cg_min,
        "x_CG_obj_max_desde_nariz (m)": x_cg_max,
    }


def generar_estaciones_fuselaje_xflr5(row_dim):
    """
    Genera estaciones simples para introducir manualmente el fuselaje en XFLR5.

    Una estación es una sección transversal del fuselaje definida por:
    - posición x desde la nariz,
    - anchura total,
    - altura total.

    La tabla no sustituye a un CAD. Sirve para reconstruir una Iteración 0
    suave del cuerpo: nariz -> cápsula -> unión con ala -> boom -> cola.
    """
    Lf = float(row_dim["Lf_fuselaje_total (m)"])
    Lcap = float(row_dim["L_capsula (m)"])
    Wmax = float(row_dim["W_capsula_max (m)"])
    Hmax = float(row_dim["H_capsula_max (m)"])
    Dboom = float(row_dim["D_boom_equiv (m)"])
    x_le = float(row_dim.get("x_ala_LE_desde_nariz (m)", 0.40 * Lf))
    x_te = float(row_dim.get("x_ala_TE_aprox_desde_nariz (m)", 0.55 * Lf))
    x_tail = float(row_dim.get("x_tail_AC_desde_nariz (m)", 0.93 * Lf))

    puntos = [
        (0.00, 0.05 * Wmax, 0.05 * Hmax, "nariz cerrada"),
        (0.20 * Lcap, 0.65 * Wmax, 0.65 * Hmax, "crecimiento nariz"),
        (0.45 * Lcap, 1.00 * Wmax, 1.00 * Hmax, "seccion maxima capsula"),
        (0.85 * Lcap, 0.80 * Wmax, 0.80 * Hmax, "cierre capsula"),
        (Lcap, Dboom, Dboom, "transicion capsula-boom"),
        (x_le, Dboom, Dboom, "borde ataque ala aprox"),
        (x_te, Dboom, Dboom, "borde salida ala aprox"),
        (0.40 * Lf, Dboom, Dboom, "boom"),
        (0.70 * Lf, Dboom, Dboom, "boom"),
        (x_tail, Dboom, Dboom, "AC cola aprox"),
        (Lf, 0.75 * Dboom, 0.75 * Dboom, "extremo trasero"),
    ]

    puntos_unicos = []
    vistos = set()
    for x, w, h, nota in sorted(puntos, key=lambda p: p[0]):
        x_round = round(x, 4)
        if x_round in vistos:
            continue
        vistos.add(x_round)
        puntos_unicos.append({
            "x_desde_nariz_m": x,
            "anchura_total_m": w,
            "altura_total_m": h,
            "semianchura_y_m": 0.5 * w,
            "semialtura_z_m": 0.5 * h,
            "nota": nota,
        })

    return pd.DataFrame(puntos_unicos)

def limpiar_nombre_archivo(nombre):
    return "".join(c if c.isalnum() or c in "-_" else "_" for c in nombre)


# --- Cálculo para todos los casos ya existentes en df_resultados ---
filas_fuselaje_cola = []
for _, row in df_resultados.iterrows():
    filas_fuselaje_cola.append(
        calcular_dimensionado_fuselaje_cola(
            row=row,
            hip=hipotesis_fuselaje_cola,
            envolventes=envolventes_componentes_m,
        )
    )

df_fuselaje_cola = pd.DataFrame(filas_fuselaje_cola)

# Añadir columnas al dataframe general para que queden disponibles en el resto del script.
df_resultados = df_resultados.merge(
    df_fuselaje_cola.drop(columns=["MTOW (kg)"]),
    on=["Caso", "Escenario"],
    how="left",
)

# Guardado de tabla de dimensionado y estaciones XFLR5.
output_dir_fuselaje = Path(__file__).resolve().parent if "__file__" in globals() else Path.cwd()
df_fuselaje_cola.to_csv(output_dir_fuselaje / "dimensionado_fuselaje_cola_T.csv", index=False, encoding="utf-8-sig")

for _, row_dim in df_fuselaje_cola.drop_duplicates(subset=["Caso"]).iterrows():
    df_estaciones = generar_estaciones_fuselaje_xflr5(row_dim)
    nombre_archivo = f"xflr5_estaciones_fuselaje_{limpiar_nombre_archivo(row_dim['Caso'])}.csv"
    df_estaciones.to_csv(output_dir_fuselaje / nombre_archivo, index=False, encoding="utf-8-sig")

# Presentación detallada en consola de la nueva parte fuselaje + cola.
# La geometría de fuselaje/cola depende del caso geométrico, no del escenario operativo.
# Por eso se eliminan duplicados de "Caso" para no imprimir geometrías repetidas.
pd.set_option("display.max_columns", None)
pd.set_option("display.width", 220)

df_fuselaje_cola_console = df_fuselaje_cola.drop_duplicates(subset=["Caso"]).copy()

cols_fus_general = [
    "Caso",
    "b_ref (m)",
    "S_ref (m2)",
    "AR_ref (-)",
    "c_bar_ref (m)",
    "Lf_por_0p35b (m)",
    "Lf_por_3cbar (m)",
    "Lf_fuselaje_total (m)",
    "L_capsula (m)",
    "L_boom_aprox (m)",
    "D_boom_equiv (m)",
]

print("\n--- DIMENSIONADO PRELIMINAR FUSELAJE + COLA EN T: GEOMETRIA GENERAL ---\n")
print(df_fuselaje_cola_console[cols_fus_general].round(3).to_string(index=False))

cols_capsula = [
    "Caso",
    "V_componentes_sin_margen (L)",
    "V_util_interno_con_margen (L)",
    "V_geom_capsula_bruto (L)",
    "D_eq_capsula (m)",
    "L_capsula (m)",
    "W_capsula_max (m)",
    "H_capsula_max (m)",
    "A_frontal_capsula (m2)",
]

print("\n--- CAPSULA DELANTERA: VOLUMEN, FORMA Y SECCION ---\n")
print(df_fuselaje_cola_console[cols_capsula].round(3).to_string(index=False))

cols_cola = [
    "Caso",
    "lH_brazo_cola_H (m)",
    "VH_usado (-)",
    "S_H_cola_horizontal (m2)",
    "AR_H_usado (-)",
    "b_H_cola_horizontal (m)",
    "c_H_media (m)",
    "lV_brazo_cola_V (m)",
    "CVT_usado (-)",
    "S_V_cola_vertical (m2)",
    "AR_V_usado (-)",
    "h_V_deriva (m)",
    "c_V_media (m)",
]

print("\n--- EMPENAJE: COLA HORIZONTAL Y COLA VERTICAL ---\n")
print(df_fuselaje_cola_console[cols_cola].round(3).to_string(index=False))

cols_posiciones = [
    "Caso",
    "x_ala_LE_desde_nariz (m)",
    "x_ala_AC_desde_nariz (m)",
    "x_ala_TE_aprox_desde_nariz (m)",
    "x_tail_AC_desde_nariz (m)",
    "x_CG_obj_min_desde_nariz (m)",
    "x_CG_obj_max_desde_nariz (m)",
    "L_tramo_capsula_a_LE (m)",
    "L_tramo_bajo_ala (m)",
    "L_boom_detras_ala (m)",
    "L_TE_ala_a_tail_AC (m)",
]

print("\n--- POSICIONES LONGITUDINALES Y TRAMOS DEL FUSELAJE ---\n")
print(df_fuselaje_cola_console[cols_posiciones].round(3).to_string(index=False))

print("\n--- ESTACIONES SIMPLIFICADAS PARA XFLR5 ---")
print("Cada estacion es una seccion transversal: x desde nariz, anchura total y altura total.")
print("Estos datos se guardan en CSV y sirven para reconstruir manualmente el cuerpo en XFLR5.\n")

cols_estaciones_print = [
    "x_desde_nariz_m",
    "anchura_total_m",
    "altura_total_m",
    "semianchura_y_m",
    "semialtura_z_m",
    "nota",
]

for _, row_dim in df_fuselaje_cola_console.iterrows():
    df_estaciones_console = generar_estaciones_fuselaje_xflr5(row_dim)
    print(f"\nEstaciones XFLR5 - {row_dim['Caso']}:\n")
    print(df_estaciones_console[cols_estaciones_print].round(3).to_string(index=False))

print("\nArchivos generados de fuselaje/cola:")
print(f"- {output_dir_fuselaje / 'dimensionado_fuselaje_cola_T.csv'}")
print(f"- {output_dir_fuselaje / 'xflr5_estaciones_fuselaje_<Caso>.csv'}")

# -----------------------------------------------------------------------------
# 5A ter. SISTEMA FOTOVOLTAICO EN ALA: AREA, MASA Y CENTROIDE PRELIMINAR
# -----------------------------------------------------------------------------
"""
Este bloque convierte la cobertura fotovoltaica del ala en masa y posiciones
preliminares para empezar el mass build-up y el calculo de CG.

Hipotesis principales:
- Se mantiene f_cobertura_max = 0.80, ya definida en el bloque solar.
- De momento solo se instalan placas en el ala, no en cola ni fuselaje.
- La masa superficial PV se toma como celulas + encapsulado.
- La cobertura chordwise se mueve hacia delante para dejar margen en borde de
  salida ante posibles alerones/flaps futuros.

IMPORTANTE:
- La masa PV se modela como masa distribuida. Para el CG se sustituye por dos
  masas equivalentes, una por semiala, situadas en el centroide del area cubierta.
- Si en una fase posterior se definen superficies moviles reales, esta distribucion
  debera revisarse.
"""

# Masa superficial del sistema fotovoltaico.
# Valores tipo Noth/Sky-Sailor: celulas + encapsulado.
sigma_cells_PV_kg_m2 = 0.32
sigma_encaps_PV_kg_m2 = 0.22
sigma_PV_kg_m2 = sigma_cells_PV_kg_m2 + sigma_encaps_PV_kg_m2

# Margen opcional para cableado local, busbars, adhesivos, conectores o integracion.
# Se deja en 0.00 para no duplicar margenes en esta primera iteracion.
margen_PV_instalacion = 0.15

# Cobertura sobre el ala.
f_PV_wing = f_cobertura_max

# Distribucion chordwise de placas.
# Caso base: x/c = 0.02 a 0.87 para reservar un 13% de cuerda en borde de salida.

x_PV_c_start = 0.02
x_PV_c_end = 0.87
x_PV_c_frac = x_PV_c_end - x_PV_c_start
x_PV_c_centroid = 0.5 * (x_PV_c_start + x_PV_c_end)

# Area y masa PV total.
df_resultados["f_PV_wing (-)"] = f_PV_wing
df_resultados["A_PV_wing (m2)"] = f_PV_wing * df_resultados["S_alar (m2)"]
df_resultados["sigma_PV (kg_m2)"] = sigma_PV_kg_m2
df_resultados["m_PV_wing_base (kg)"] = (
    df_resultados["A_PV_wing (m2)"] * sigma_PV_kg_m2
)
df_resultados["m_PV_wing_instalada (kg)"] = (
    df_resultados["m_PV_wing_base (kg)"] * (1.0 + margen_PV_instalacion)
)

# Reparto simetrico izquierda/derecha.
df_resultados["A_PV_semiala (m2)"] = df_resultados["A_PV_wing (m2)"] / 2.0
df_resultados["m_PV_left (kg)"] = df_resultados["m_PV_wing_instalada (kg)"] / 2.0
df_resultados["m_PV_right (kg)"] = df_resultados["m_PV_wing_instalada (kg)"] / 2.0

# Centroide longitudinal de la masa PV equivalente.
df_resultados["x_PV_c_start (-)"] = x_PV_c_start
df_resultados["x_PV_c_end (-)"] = x_PV_c_end
df_resultados["x_PV_c_centroid (-)"] = x_PV_c_centroid
df_resultados["x_PV_desde_nariz (m)"] = (
    df_resultados["x_ala_LE_desde_nariz (m)"]
    + x_PV_c_centroid * df_resultados["c_bar_ref (m)"]
)

# Longitud spanwise necesaria por semiala para conseguir el area PV definida.
# Aproximacion valida para ala rectangular/equivalente con cuerda media c_bar_ref.
df_resultados["y_PV_longitud_semiala (m)"] = (
    df_resultados["A_PV_semiala (m2)"]
    / (x_PV_c_frac * df_resultados["c_bar_ref (m)"])
)
df_resultados["y_PV_semispan (m)"] = df_resultados["b (m)"] / 2.0

# Si la longitud requerida no cabe en la semiala, se marca. En el caso base debe caber.
df_resultados["PV_cabe_en_semiala (1/0)"] = (
    df_resultados["y_PV_longitud_semiala (m)"] <= df_resultados["y_PV_semispan (m)"]
).astype(int)

# Para que el centroide lateral no quede pegado a raiz ni punta, se reparte la franja
# dejando margen simetrico en la semiala siempre que sea posible.
df_resultados["y_PV_margen_total_semiala (m)"] = (
    df_resultados["y_PV_semispan (m)"] - df_resultados["y_PV_longitud_semiala (m)"]
).clip(lower=0.0)
df_resultados["y_PV_inicio_desde_raiz (m)"] = 0.5 * df_resultados["y_PV_margen_total_semiala (m)"]
df_resultados["y_PV_fin_desde_raiz (m)"] = (
    df_resultados["y_PV_inicio_desde_raiz (m)"]
    + df_resultados["y_PV_longitud_semiala (m)"]
)
df_resultados["y_PV_abs_centroid (m)"] = 0.5 * (
    df_resultados["y_PV_inicio_desde_raiz (m)"]
    + df_resultados["y_PV_fin_desde_raiz (m)"]
)

# Margenes chordwise expresados en metros para interpretacion geometrica.
df_resultados["margen_LE_PV (m)"] = x_PV_c_start * df_resultados["c_bar_ref (m)"]
df_resultados["margen_TE_PV (m)"] = (1.0 - x_PV_c_end) * df_resultados["c_bar_ref (m)"]
df_resultados["longitud_cuerda_PV (m)"] = x_PV_c_frac * df_resultados["c_bar_ref (m)"]

cols_pv_console = [
    "Caso",
    "S_alar (m2)",
    "f_PV_wing (-)",
    "A_PV_wing (m2)",
    "sigma_PV (kg_m2)",
    "m_PV_wing_instalada (kg)",
    "m_PV_left (kg)",
    "m_PV_right (kg)",
    "x_PV_c_start (-)",
    "x_PV_c_end (-)",
    "margen_TE_PV (m)",
    "x_PV_desde_nariz (m)",
    "y_PV_inicio_desde_raiz (m)",
    "y_PV_fin_desde_raiz (m)",
    "y_PV_abs_centroid (m)",
    "PV_cabe_en_semiala (1/0)",
]

df_pv_console = df_resultados.drop_duplicates(subset=["Caso"])[cols_pv_console].copy()

# Guardado de tabla PV para revisar en Excel/CSV.
archivo_pv_wing = output_dir_fuselaje / "masa_y_distribucion_PV_ala.csv"
df_pv_console.round(4).to_csv(archivo_pv_wing, index=False, encoding="utf-8-sig")

print("\n" + "=" * 78)
print("SISTEMA FOTOVOLTAICO EN ALA: AREA, MASA Y DISTRIBUCION PRELIMINAR")
print("=" * 78)
print(f"Cobertura PV en ala: {f_PV_wing:.2f} de S_alar")
print(f"Masa superficial PV: {sigma_PV_kg_m2:.3f} kg/m2")
print(f"Distribucion chordwise base: x/c = {x_PV_c_start:.2f} a {x_PV_c_end:.2f}")
print("Interpretacion: se reserva margen en borde de salida para posibles superficies moviles.")
print(f"Tabla completa guardada en CSV: {archivo_pv_wing}")

for _, r in df_pv_console.iterrows():
    print("\n" + "-" * 78)
    print(f"Caso: {r['Caso']}")
    print("-" * 78)
    print(f"Superficie alar:                  {r['S_alar (m2)']:.3f} m2")
    print(f"Area PV en ala:                   {r['A_PV_wing (m2)']:.3f} m2")
    print(f"Masa PV total ala:                {r['m_PV_wing_instalada (kg)']:.3f} kg")
    print(f"Masa PV por semiala:              {r['m_PV_left (kg)']:.3f} kg")
    print(f"Rango chordwise PV:               x/c = {r['x_PV_c_start (-)']:.2f} a {r['x_PV_c_end (-)']:.2f}")
    print(f"Margen borde salida equivalente:  {r['margen_TE_PV (m)']:.3f} m")
    print(f"Centroide x PV desde nariz:       {r['x_PV_desde_nariz (m)']:.3f} m")
    print(f"Franja PV por semiala:            y = {r['y_PV_inicio_desde_raiz (m)']:.3f} a {r['y_PV_fin_desde_raiz (m)']:.3f} m")
    print(f"Centroide |y| PV:                 {r['y_PV_abs_centroid (m)']:.3f} m")
    print(f"PV cabe en semiala:               {int(r['PV_cabe_en_semiala (1/0)'])}")



# -----------------------------------------------------------------------------
# 5A quater. VOLUMEN INTERNO DEL ALA PARA INTEGRACION DE BATERIAS
# -----------------------------------------------------------------------------
"""
BLOQUE NUEVO: calculo del volumen interno disponible en el ala leyendo el
perfil SD7032 desde un archivo .dat.

QUE HACE:
- Lee el perfil SD7032 desde un .dat en formato Selig/UIUC o Lednicer.
- Separa extradós e intradós.
- Integra el espesor local y_u - y_l en una bahia de cuerda seleccionada.
- Convierte el area 2D de la bahia en volumen por metro de envergadura.
- Aplica un factor de aprovechamiento interno para tener en cuenta estructura,
  costillas, largueros, cableado, fijaciones y margen de seguridad.
- Compara ese volumen util con el volumen de bateria requerido.

IMPORTANTE:
- Es una comprobacion geometrica preliminar, no un diseno estructural del ala.
- Las baterias se suponen modulares/adaptables a la bahia del ala.
- El factor eta_bay_bateria debe tratarse como hipotesis editable.
"""

# --- Hipotesis editables para bahia de baterias en ala ---
airfoil_dat_filename = "sd7032.dat"   # pon este archivo .dat en la misma carpeta que el script

# Bahia de cuerda considerada para baterias.
# Se evita borde de ataque, borde de salida y zonas futuras de superficies moviles.
x_batt_c_start = 0.20
x_batt_c_end = 0.60

# Factor de aprovechamiento real del volumen geometrico de la bahia.
# Incluye perdidas por estructura, costillas, largueros, cableado, separacion,
# accesibilidad, fijaciones y margen de seguridad del pack.
eta_bay_bateria = 0.35

# Densidad energetica volumetrica pack-level usada para traducir energia a volumen.
rhoE_batt_pack_base_Wh_L = 245.0
rhoE_batt_pack_conservadora_Wh_L = 200.0
margen_volumen_bateria_integracion = 0.20

# Posicion lateral inicial de comienzo de la bahia de bateria desde la raiz.
# Se deja un margen inicial en la raiz para union ala-fuselaje, cableado, estructura, etc.
y_batt_inicio_desde_raiz_m = 0.75


def _limpiar_puntos_duplicados_por_x(x_arr, y_arr):
    """Ordena por x y promedia y si hay x duplicadas."""
    x_arr = np.asarray(x_arr, dtype=float)
    y_arr = np.asarray(y_arr, dtype=float)
    orden = np.argsort(x_arr)
    x_ord = x_arr[orden]
    y_ord = y_arr[orden]

    xs = []
    ys = []
    for x_val in np.unique(x_ord):
        mask = np.isclose(x_ord, x_val, rtol=0.0, atol=1e-10)
        xs.append(float(x_val))
        ys.append(float(np.mean(y_ord[mask])))
    return np.array(xs, dtype=float), np.array(ys, dtype=float)


def leer_perfil_dat_generico(ruta_dat):
    """
    Lee un perfil .dat en formato Selig/UIUC o Lednicer.

    Devuelve:
    - xu, yu: extradós ordenado de LE a TE.
    - xl, yl: intradós ordenado de LE a TE.

    Formato Selig/UIUC habitual:
    nombre
    TE superior -> LE -> TE inferior

    Formato Lednicer habitual:
    nombre
    n_upper n_lower
    puntos extradós LE -> TE
    puntos intradós LE -> TE
    """
    ruta_dat = Path(ruta_dat)
    lineas = ruta_dat.read_text(encoding="utf-8", errors="ignore").splitlines()

    pares_numericos = []
    for linea in lineas:
        s = linea.strip().replace(",", " ")
        if not s:
            continue
        partes = s.split()
        if len(partes) < 2:
            continue
        try:
            x_val = float(partes[0])
            y_val = float(partes[1])
            pares_numericos.append((x_val, y_val))
        except ValueError:
            continue

    if len(pares_numericos) < 10:
        raise ValueError(f"No se han encontrado suficientes puntos numericos en {ruta_dat}")

    # Deteccion simple de Lednicer: primera linea numerica suele ser n_upper n_lower,
    # ambos enteros y mayores que 2.
    primero = pares_numericos[0]
    n_up = int(round(primero[0]))
    n_low = int(round(primero[1]))
    es_lednicer = (
        abs(primero[0] - n_up) < 1e-8
        and abs(primero[1] - n_low) < 1e-8
        and n_up > 2
        and n_low > 2
        and len(pares_numericos) >= 1 + n_up + n_low
    )

    if es_lednicer:
        pts = np.array(pares_numericos[1:], dtype=float)
        upper = pts[:n_up]
        lower = pts[n_up:n_up + n_low]
        xu, yu = _limpiar_puntos_duplicados_por_x(upper[:, 0], upper[:, 1])
        xl, yl = _limpiar_puntos_duplicados_por_x(lower[:, 0], lower[:, 1])
        return xu, yu, xl, yl, "Lednicer"

    # Si no es Lednicer, se asume Selig/UIUC como contorno cerrado.
    pts = np.array(pares_numericos, dtype=float)
    i_le = int(np.argmin(pts[:, 0]))

    # Superficie superior: normalmente TE -> LE, se invierte para LE -> TE.
    upper = pts[:i_le + 1][::-1]
    # Superficie inferior: normalmente LE -> TE.
    lower = pts[i_le:]

    xu, yu = _limpiar_puntos_duplicados_por_x(upper[:, 0], upper[:, 1])
    xl, yl = _limpiar_puntos_duplicados_por_x(lower[:, 0], lower[:, 1])
    return xu, yu, xl, yl, "Selig/UIUC"


def integrar_bahia_perfil_normalizado(xu, yu, xl, yl, x_start, x_end, n_grid=2001):
    """
    Integra el espesor normalizado y_u - y_l entre x_start y x_end.

    Devuelve:
    - area_norm: integral adimensional del espesor.
    - x_centroid_norm: centroide en x/c de la bahia integrada.
    - t_prom_norm: espesor medio normalizado en la bahia.
    """
    x_start = float(max(0.0, min(1.0, x_start)))
    x_end = float(max(0.0, min(1.0, x_end)))
    if x_end <= x_start:
        raise ValueError("x_end debe ser mayor que x_start para integrar la bahia")

    x_grid = np.linspace(x_start, x_end, int(n_grid))
    yu_grid = np.interp(x_grid, xu, yu)
    yl_grid = np.interp(x_grid, xl, yl)
    espesor = yu_grid - yl_grid
    espesor = np.maximum(espesor, 0.0)

    area_norm = float(np.trapezoid(espesor, x_grid))
    if area_norm > 0.0:
        x_centroid_norm = float(np.trapezoid(x_grid * espesor, x_grid) / area_norm)
    else:
        x_centroid_norm = float("nan")
    t_prom_norm = area_norm / (x_end - x_start)
    return area_norm, x_centroid_norm, t_prom_norm


def buscar_archivo_perfil_dat(nombre_archivo):
    """Busca el .dat en la carpeta del script y en el directorio de trabajo."""
    base_script = Path(__file__).resolve().parent if "__file__" in globals() else Path.cwd()
    candidatos = [
        base_script / nombre_archivo,
        base_script / "sd7032(1).dat",
        base_script / "sd7032.dat",
        Path.cwd() / nombre_archivo,
        Path.cwd() / "sd7032(1).dat",
        Path.cwd() / "sd7032.dat",
    ]
    vistos = []
    for cand in candidatos:
        if cand not in vistos:
            vistos.append(cand)
            if cand.exists():
                return cand
    return None


ruta_perfil_bateria = buscar_archivo_perfil_dat(airfoil_dat_filename)
volumen_ala_disponible = ruta_perfil_bateria is not None

if volumen_ala_disponible:
    xu_sd, yu_sd, xl_sd, yl_sd, formato_perfil_detectado = leer_perfil_dat_generico(ruta_perfil_bateria)
    A_perfil_total_norm, x_perfil_total_centroid_norm, t_prom_total_norm = integrar_bahia_perfil_normalizado(
        xu_sd, yu_sd, xl_sd, yl_sd, 0.0, 1.0
    )
    A_bay_norm, x_bay_centroid_norm, t_bay_prom_norm = integrar_bahia_perfil_normalizado(
        xu_sd, yu_sd, xl_sd, yl_sd, x_batt_c_start, x_batt_c_end
    )
else:
    formato_perfil_detectado = "NO ENCONTRADO"
    A_perfil_total_norm = float("nan")
    x_perfil_total_centroid_norm = float("nan")
    t_prom_total_norm = float("nan")
    A_bay_norm = float("nan")
    x_bay_centroid_norm = 0.5 * (x_batt_c_start + x_batt_c_end)
    t_bay_prom_norm = float("nan")

# Geometria del volumen interno.
df_resultados["airfoil_dat_usado"] = str(ruta_perfil_bateria) if ruta_perfil_bateria is not None else "NO ENCONTRADO"
df_resultados["airfoil_formato_detectado"] = formato_perfil_detectado
df_resultados["x_batt_c_start (-)"] = x_batt_c_start
df_resultados["x_batt_c_end (-)"] = x_batt_c_end
df_resultados["x_batt_c_centroid (-)"] = x_bay_centroid_norm
df_resultados["eta_bay_bateria (-)"] = eta_bay_bateria

df_resultados["A_perfil_total_norm (-)"] = A_perfil_total_norm
df_resultados["A_bay_batt_norm (-)"] = A_bay_norm
df_resultados["t_bay_prom_norm (-)"] = t_bay_prom_norm

df_resultados["A_bay_batt (m2)"] = A_bay_norm * df_resultados["c_bar_ref (m)"] ** 2
df_resultados["V_bay_geom_L_por_m (L_m)"] = df_resultados["A_bay_batt (m2)"] * 1000.0
df_resultados["V_bay_util_L_por_m (L_m)"] = (
    df_resultados["V_bay_geom_L_por_m (L_m)"] * eta_bay_bateria
)
df_resultados["V_bay_util_semiala_total (L)"] = (
    df_resultados["V_bay_util_L_por_m (L_m)"] * df_resultados["b (m)"] / 2.0
)
df_resultados["V_bay_util_ala_total (L)"] = (
    df_resultados["V_bay_util_L_por_m (L_m)"] * df_resultados["b (m)"]
)
df_resultados["V_ala_geom_total_perfil (L)"] = (
    A_perfil_total_norm * df_resultados["c_bar_ref (m)"] ** 2 * df_resultados["b (m)"] * 1000.0
)

df_resultados["x_batt_desde_nariz (m)"] = (
    df_resultados["x_ala_LE_desde_nariz (m)"]
    + x_bay_centroid_norm * df_resultados["c_bar_ref (m)"]
)

# Volumen de bateria a partir del pack nominal calculado previamente.
# Se toma la mision de 30 dias como referencia de integracion de bateria.
df_resultados["rhoE_batt_pack_base (Wh_L)"] = rhoE_batt_pack_base_Wh_L
df_resultados["rhoE_batt_pack_conservadora (Wh_L)"] = rhoE_batt_pack_conservadora_Wh_L
df_resultados["margen_vol_batt_integracion (-)"] = margen_volumen_bateria_integracion

df_resultados["V_batt_pack_30d_245WhL (L)"] = (
    df_resultados["E_nominal_pack_30d (Wh)"] / rhoE_batt_pack_base_Wh_L
)
df_resultados["V_batt_pack_30d_200WhL (L)"] = (
    df_resultados["E_nominal_pack_30d (Wh)"] / rhoE_batt_pack_conservadora_Wh_L
)
df_resultados["V_batt_pack_30d_200WhL_margen (L)"] = (
    df_resultados["V_batt_pack_30d_200WhL (L)"]
    * (1.0 + margen_volumen_bateria_integracion)
)
df_resultados["V_batt_pack_semiala_req (L)"] = (
    df_resultados["V_batt_pack_30d_200WhL_margen (L)"] / 2.0
)

# Longitud de bahia requerida por semiala.
df_resultados["L_bahia_batt_req_semiala (m)"] = (
    df_resultados["V_batt_pack_semiala_req (L)"]
    / df_resultados["V_bay_util_L_por_m (L_m)"]
)
df_resultados["y_batt_inicio_desde_raiz (m)"] = y_batt_inicio_desde_raiz_m
df_resultados["y_batt_fin_desde_raiz (m)"] = (
    df_resultados["y_batt_inicio_desde_raiz (m)"]
    + df_resultados["L_bahia_batt_req_semiala (m)"]
)
df_resultados["y_batt_abs_centroid (m)"] = 0.5 * (
    df_resultados["y_batt_inicio_desde_raiz (m)"]
    + df_resultados["y_batt_fin_desde_raiz (m)"]
)
df_resultados["bateria_cabe_en_semiala (1/0)"] = (
    df_resultados["y_batt_fin_desde_raiz (m)"] <= (df_resultados["b (m)"] / 2.0)
).astype(int)
df_resultados["margen_volumen_bahia_vs_bateria (-)"] = (
    df_resultados["V_bay_util_semiala_total (L)"]
    / df_resultados["V_batt_pack_semiala_req (L)"]
)

cols_volumen_ala_bateria = [
    "Caso",
    "Escenario",
    "airfoil_dat_usado",
    "airfoil_formato_detectado",
    "c_bar_ref (m)",
    "b (m)",
    "A_perfil_total_norm (-)",
    "V_ala_geom_total_perfil (L)",
    "x_batt_c_start (-)",
    "x_batt_c_end (-)",
    "x_batt_c_centroid (-)",
    "A_bay_batt_norm (-)",
    "A_bay_batt (m2)",
    "eta_bay_bateria (-)",
    "V_bay_geom_L_por_m (L_m)",
    "V_bay_util_L_por_m (L_m)",
    "V_bay_util_semiala_total (L)",
    "E_nominal_pack_30d (Wh)",
    "Masa_min_conservadora_30d (kg)",
    "V_batt_pack_30d_245WhL (L)",
    "V_batt_pack_30d_200WhL (L)",
    "V_batt_pack_30d_200WhL_margen (L)",
    "V_batt_pack_semiala_req (L)",
    "L_bahia_batt_req_semiala (m)",
    "x_batt_desde_nariz (m)",
    "y_batt_inicio_desde_raiz (m)",
    "y_batt_fin_desde_raiz (m)",
    "y_batt_abs_centroid (m)",
    "bateria_cabe_en_semiala (1/0)",
    "margen_volumen_bahia_vs_bateria (-)",
]

archivo_volumen_ala_batt = output_dir_fuselaje / "volumen_interno_ala_baterias.csv"
df_volumen_ala_bateria = df_resultados[cols_volumen_ala_bateria].copy()
df_volumen_ala_bateria.round(5).to_csv(archivo_volumen_ala_batt, index=False, encoding="utf-8-sig")

print("\n" + "=" * 78)
print("VOLUMEN INTERNO DEL ALA PARA INTEGRACION DE BATERIAS")
print("=" * 78)
if ruta_perfil_bateria is None:
    print("AVISO: no se encontro el archivo .dat del perfil. Revisa airfoil_dat_filename y la carpeta del script.")
else:
    print(f"Perfil usado: {ruta_perfil_bateria}")
    print(f"Formato detectado: {formato_perfil_detectado}")
print(f"Bahia de bateria considerada: x/c = {x_batt_c_start:.2f} a {x_batt_c_end:.2f}")
print(f"Factor de aprovechamiento eta_bay: {eta_bay_bateria:.2f}")
print(f"Tabla completa guardada en CSV: {archivo_volumen_ala_batt}")

for _, r in df_resultados.drop_duplicates(subset=["Caso", "Escenario"]).iterrows():
    print("\n" + "-" * 78)
    print(f"Caso: {r['Caso']}")
    print(f"Escenario: {r['Escenario']}")
    print("-" * 78)
    print(f"Cuerda media:                         {r['c_bar_ref (m)']:.3f} m")
    print(f"Area perfil total normalizada:        {r['A_perfil_total_norm (-)']:.5f}")
    print(f"Volumen geom total ala perfil:        {r['V_ala_geom_total_perfil (L)']:.1f} L")
    print(f"Area bahia bateria normalizada:       {r['A_bay_batt_norm (-)']:.5f}")
    print(f"Area bahia bateria dimensional:       {r['A_bay_batt (m2)']:.4f} m2")
    print(f"Volumen geometrico bahia por metro:   {r['V_bay_geom_L_por_m (L_m)']:.1f} L/m")
    print(f"Volumen util bahia por metro:         {r['V_bay_util_L_por_m (L_m)']:.1f} L/m")
    print(f"Volumen util disponible por semiala:  {r['V_bay_util_semiala_total (L)']:.1f} L")
    print(f"Volumen bateria 30d 200 Wh/L+margen:  {r['V_batt_pack_30d_200WhL_margen (L)']:.1f} L")
    print(f"Volumen bateria requerido por lado:   {r['V_batt_pack_semiala_req (L)']:.1f} L")
    print(f"Longitud bahia requerida por semiala: {r['L_bahia_batt_req_semiala (m)']:.2f} m")
    print(f"Centroide bateria x desde nariz:      {r['x_batt_desde_nariz (m)']:.3f} m")
    print(f"Franja bateria por semiala:           y = {r['y_batt_inicio_desde_raiz (m)']:.3f} a {r['y_batt_fin_desde_raiz (m)']:.3f} m")
    print(f"Centroide |y| bateria:                {r['y_batt_abs_centroid (m)']:.3f} m")
    print(f"Bateria cabe en semiala:              {int(r['bateria_cabe_en_semiala (1/0)'])}")
    print(f"Margen volumen semiala/bateria:       {r['margen_volumen_bahia_vs_bateria (-)']:.2f}")


"""
POSTPROCESADO NUEVO: métricas de sizing inicial ligadas a Gundlach.

QUÉ SE CALCULA AQUÍ:
1) MF_Energy:
   - A partir de la masa mínima conservadora del pack, se calcula la fracción
     de masa energética respecto al MTOW.
   - Para baterías químicas, MF_Energy = MF_Batt = W_Batt / W_TO.

2) MF_Struct y MF_Subs asumidas:
   - Se dejan visibles como hipótesis provisionales por similitud, NO como
     resultados calculados por el script.
   - Sirven para ordenar el cierre del sizing conceptual mientras no exista
     un mass build-up estructural detallado.

3) Ratios aircraft-level P/W y T/W:
   - P/W es el parámetro principal en aeronaves propulsadas por hélice.
   - T/W se mantiene como magnitud auxiliar.
   - Ambos salen directamente de las potencias y empujes que el script ya calcula
     en crucero y a lo largo del perfil operativo.

4) Punto 8 del sizing:
   - Se deja montada la estructura para convertir potencia requerida del avión
     a masa del sistema propulsivo usando benchmarks de tecnología.
   - Si no se han rellenado todavía los datos tecnológicos, el DataFrame queda
     vacío y el programa lo indica explícitamente en consola.
"""

# 5A. Fracción de masa energética del pack
df_resultados["MF_Energy_1d (-)"] = (
    df_resultados["Masa_min_conservadora_1d (kg)"] / df_resultados["MTOW (kg)"]
)
df_resultados["MF_Energy_30d (-)"] = (
    df_resultados["Masa_min_conservadora_30d (kg)"] / df_resultados["MTOW (kg)"]
)
df_resultados["MF_Energy_60d (-)"] = (
    df_resultados["Masa_min_conservadora_60d (kg)"] / df_resultados["MTOW (kg)"]
)

# 5B. Hipótesis provisionales de MF_Struct y MF_Subs
mf_struct_guess = {
    "24m": 0.25,   # editable; solo como arranque conceptual inspirado en Gundlach
    "32m": 0.25,
}
mf_subs_guess = {
    "24m": 0.05,   # editable; solo como arranque conceptual inspirado en Gundlach
    "32m": 0.05,
}

df_resultados["MF_Struct_asumida (-)"] = df_resultados["Caso"].apply(
    lambda s: mf_struct_guess[clave_plataforma_desde_caso(s)]
)
df_resultados["MF_Subs_asumida (-)"] = df_resultados["Caso"].apply(
    lambda s: mf_subs_guess[clave_plataforma_desde_caso(s)]
)

# 5C. Ratios de sizing del avión (aircraft-level)
W_N_series = df_resultados["MTOW (kg)"] * g

df_resultados["P_W_elec_crucero_Wkg"] = (
    df_resultados["P_elec_crucero_20km_base (W)"] / df_resultados["MTOW (kg)"]
)
df_resultados["P_W_elec_max_perfil_Wkg"] = (
    df_resultados["P_propulsion_max (W)"] / df_resultados["MTOW (kg)"]
)
df_resultados["P_W_thrust_max_perfil_Wkg"] = (
    df_resultados["P_thrust_max_perfil (W)"] / df_resultados["MTOW (kg)"]
)
df_resultados["T_W_crucero_20km_base (-)"] = (
    df_resultados["T_crucero_20km_base (N)"] / W_N_series
)
df_resultados["T_W_max_perfil (-)"] = (
    df_resultados["T_req_max_perfil (N)"] / W_N_series
)


# 5C bis. Sistema propulsivo comercial de referencia
# ---------------------------------------------------
# Benchmark preliminar: 2 x Hobbywing C6225-200KV + 2 x XRotor Pro H130A
# + 2 helices de referencia. Se usa para masa, corriente y packaging,
# no como validacion final del rendimiento a 20 km.
# IMPORTANTE: el bloque 5C ter calcula el diametro de disco requerido por TCM.
# Ese diametro puede ser mayor que la helice comercial 21x10, por lo que la masa
# de 2.11 kg debe leerse como benchmark provisional, no como propulsor final cerrado.
propulsion_ref = {
    "nombre": "Hobbywing C6225-200KV + XRotor Pro H130A 14S + helice 20x10/21x10",
    "N_motores": 2,
    "m_motor_kg": 0.530,
    "m_ESC_kg": 0.150,
    "m_helice_kg": 0.200,
    "factor_instalacion": 1.20,
    "V_bus_V": 48.0,
    "P_input_nominal_motor_W": 710.0,
    "P_output_nominal_motor_W": 580.0,
    "I_ESC_continua_A": 60.0,
}

N_motores_ref = propulsion_ref["N_motores"]
m_propulsion_ref_kg = (
    N_motores_ref
    * (
        propulsion_ref["m_motor_kg"]
        + propulsion_ref["m_ESC_kg"]
        + propulsion_ref["m_helice_kg"]
    )
    * propulsion_ref["factor_instalacion"]
)

df_resultados["N_motores_ref (-)"] = N_motores_ref
df_resultados["M_propulsion_ref_kg"] = m_propulsion_ref_kg
df_resultados["MF_Prop_ref (-)"] = (
    df_resultados["M_propulsion_ref_kg"] / df_resultados["MTOW (kg)"]
)
df_resultados["P_elec_crucero_por_motor_ref (W)"] = (
    df_resultados["P_elec_crucero_20km_base (W)"] / N_motores_ref
)
df_resultados["I_crucero_por_motor_ref (A)"] = (
    df_resultados["P_elec_crucero_por_motor_ref (W)"] / propulsion_ref["V_bus_V"]
)
df_resultados["P_elec_max_por_motor_ref (W)"] = (
    df_resultados["P_propulsion_max (W)"] / N_motores_ref
)
df_resultados["I_max_por_motor_ref (A)"] = (
    df_resultados["P_elec_max_por_motor_ref (W)"] / propulsion_ref["V_bus_V"]
)
df_resultados["I_nominal_input_motor_ref (A)"] = (
    propulsion_ref["P_input_nominal_motor_W"] / propulsion_ref["V_bus_V"]
)
df_resultados["Margen_I_crucero_vs_I_nominal (-)"] = (
    df_resultados["I_nominal_input_motor_ref (A)"]
    / df_resultados["I_crucero_por_motor_ref (A)"]
)
df_resultados["Margen_I_max_vs_I_nominal (-)"] = (
    df_resultados["I_nominal_input_motor_ref (A)"]
    / df_resultados["I_max_por_motor_ref (A)"]
)
df_resultados["Margen_I_crucero_vs_ESC60A (-)"] = (
    propulsion_ref["I_ESC_continua_A"]
    / df_resultados["I_crucero_por_motor_ref (A)"]
)
df_resultados["Margen_I_max_vs_ESC60A (-)"] = (
    propulsion_ref["I_ESC_continua_A"]
    / df_resultados["I_max_por_motor_ref (A)"]
)

# Tabla compacta para consola.
cols_prop_ref_console = [
    "Caso",
    "Escenario",
    "M_propulsion_ref_kg",
    "MF_Prop_ref (-)",
    "P_elec_crucero_por_motor_ref (W)",
    "I_crucero_por_motor_ref (A)",
    "P_elec_max_por_motor_ref (W)",
    "I_max_por_motor_ref (A)",
    "I_nominal_input_motor_ref (A)",
    "Margen_I_crucero_vs_I_nominal (-)",
    "Margen_I_max_vs_I_nominal (-)",
    "Margen_I_crucero_vs_ESC60A (-)",
    "Margen_I_max_vs_ESC60A (-)",
]
df_console_prop_ref = df_resultados[cols_prop_ref_console].copy()

# 5C ter. Dimensionado preliminar de helices mediante teoria de cantidad de movimiento
# -----------------------------------------------------------------------------------
"""
Este bloque dimensiona el area de disco necesaria de las helices usando la
teoria de cantidad de movimiento (TCM) de las diapositivas de helices.

OBJETIVO:
- Comprobar si la helice comercial usada como referencia inicial (21 in) es
  razonable para operar a 20 km.
- Obtener un diametro preliminar de helice para representar en OpenVSP con
  configuracion bimotor.
- Comparar rapidamente la configuracion de 2 motores frente a 4 motores.

HIPOTESIS:
- Se modela cada helice como disco actuador.
- Velocidad de vuelo usada: v_crucero = 23 m/s, interpretada como TAS a 20 km.
- Densidad local: rho_crucero, calculada por ISA.
- Se adopta un rendimiento ideal de disco/helice eta_i = U_inf/U_1.

ECUACIONES:
    G = rho * U_1 * A
    T = 2 * rho * U_1 * A * (U_1 - U_inf)
    eta_i = U_inf / U_1

Despejando:
    U_1 = U_inf / eta_i
    A = T_i / [2 * rho * U_1 * (U_1 - U_inf)]
    D = sqrt(4A/pi)

LECTURA:
- Si D requerido para 2 motores sale muy grande, la alternativa conceptual es
  aumentar el numero de motores para reducir el empuje y el area de disco por helice.
"""

eta_i_helice_TCM = 0.80
N_motores_TCM_base = 2
T_total_helice_diseno_N = 35.0      # empuje total de diseno para dimension visual OpenVSP
T_total_helice_sensibilidad_N = [26.7, 30.0, 35.0, 40.0]
N_motores_helice_sensibilidad = [2, 4]
D_helice_comercial_21in_m = 21.0 * 0.0254


def dimensionar_disco_helice_TCM(T_total_N, N_motores, rho, V_inf, eta_i):
    """Dimensiona area y diametro de disco por helice con teoria de cantidad de movimiento."""
    T_total_N = float(T_total_N)
    N_motores = int(N_motores)
    rho = float(rho)
    V_inf = float(V_inf)
    eta_i = float(eta_i)

    if N_motores <= 0:
        raise ValueError("N_motores debe ser positivo")
    if rho <= 0 or V_inf <= 0:
        raise ValueError("rho y V_inf deben ser positivos")
    if not (0.0 < eta_i < 1.0):
        raise ValueError("eta_i debe estar entre 0 y 1")

    T_por_helice_N = T_total_N / N_motores
    U1_m_s = V_inf / eta_i
    vi_m_s = U1_m_s - V_inf
    A_helice_m2 = T_por_helice_N / (2.0 * rho * U1_m_s * vi_m_s)
    D_helice_m = math.sqrt(4.0 * A_helice_m2 / math.pi)
    mdot_helice_kg_s = rho * U1_m_s * A_helice_m2
    P_ideal_por_helice_W = T_por_helice_N * U1_m_s
    P_util_por_helice_W = T_por_helice_N * V_inf

    return {
        "T_total_N": T_total_N,
        "N_motores": N_motores,
        "T_por_helice_N": T_por_helice_N,
        "rho_kg_m3": rho,
        "V_inf_TAS_m_s": V_inf,
        "eta_i (-)": eta_i,
        "U1_disco_m_s": U1_m_s,
        "vi_inducida_m_s": vi_m_s,
        "A_disco_helice_m2": A_helice_m2,
        "D_helice_m": D_helice_m,
        "D_helice_in": D_helice_m / 0.0254,
        "mdot_helice_kg_s": mdot_helice_kg_s,
        "P_ideal_por_helice_W": P_ideal_por_helice_W,
        "P_util_por_helice_W": P_util_por_helice_W,
    }


def eta_ideal_para_diametro_helice(T_total_N, N_motores, D_helice_m, rho, V_inf):
    """Calcula eta_i implicita para una helice de diametro dado, usando TCM."""
    T_por_helice_N = float(T_total_N) / int(N_motores)
    A = math.pi * float(D_helice_m) ** 2 / 4.0
    # T = 2 rho A (V + vi) vi = 2 rho A (vi^2 + V vi)
    # Se resuelve la raiz positiva de vi^2 + V vi - T/(2 rho A) = 0.
    vi = (-float(V_inf) + math.sqrt(float(V_inf) ** 2 + 4.0 * T_por_helice_N / (2.0 * float(rho) * A))) / 2.0
    U1 = float(V_inf) + vi
    eta_i = float(V_inf) / U1
    return eta_i, vi, A


# Sensibilidad general 2 vs 4 motores para empujes de referencia.
filas_helices_sens = []
for T_total in T_total_helice_sensibilidad_N:
    for N_mot in N_motores_helice_sensibilidad:
        fila = dimensionar_disco_helice_TCM(
            T_total_N=T_total,
            N_motores=N_mot,
            rho=rho_crucero,
            V_inf=v_crucero,
            eta_i=eta_i_helice_TCM,
        )
        eta_21, vi_21, A_21 = eta_ideal_para_diametro_helice(
            T_total_N=T_total,
            N_motores=N_mot,
            D_helice_m=D_helice_comercial_21in_m,
            rho=rho_crucero,
            V_inf=v_crucero,
        )
        fila["D_21in_m"] = D_helice_comercial_21in_m
        fila["eta_i_implicita_21in (-)"] = eta_21
        fila["vi_implicita_21in_m_s"] = vi_21
        filas_helices_sens.append(fila)


df_helices_sensibilidad = pd.DataFrame(filas_helices_sens)
archivo_helices_sens = output_dir_fuselaje / "dimensionado_helices_TCM_sensibilidad.csv"
df_helices_sensibilidad.round(5).to_csv(archivo_helices_sens, index=False, encoding="utf-8-sig")

# Diametro recomendado para la geometria OpenVSP con 2 motores.
helice_openvsp_diseno = dimensionar_disco_helice_TCM(
    T_total_N=T_total_helice_diseno_N,
    N_motores=N_motores_TCM_base,
    rho=rho_crucero,
    V_inf=v_crucero,
    eta_i=eta_i_helice_TCM,
)
D_helice_openvsp_TCM_m = helice_openvsp_diseno["D_helice_m"]
D_helice_openvsp_TCM_in = helice_openvsp_diseno["D_helice_in"]
A_helice_openvsp_TCM_m2 = helice_openvsp_diseno["A_disco_helice_m2"]

# Resultados por caso/escenario usando los empujes calculados por el script.
filas_helices_resultados = []
for _, row in df_resultados.iterrows():
    for etiqueta_T, col_T in [
        ("crucero_20km_base", "T_crucero_20km_base (N)"),
        ("max_perfil_operativo", "T_req_max_perfil (N)"),
    ]:
        fila = dimensionar_disco_helice_TCM(
            T_total_N=row[col_T],
            N_motores=N_motores_TCM_base,
            rho=rho_crucero,
            V_inf=v_crucero,
            eta_i=eta_i_helice_TCM,
        )
        eta_21, vi_21, A_21 = eta_ideal_para_diametro_helice(
            T_total_N=row[col_T],
            N_motores=N_motores_TCM_base,
            D_helice_m=D_helice_comercial_21in_m,
            rho=rho_crucero,
            V_inf=v_crucero,
        )
        fila.update({
            "Caso": row["Caso"],
            "Escenario": row["Escenario"],
            "Punto_T": etiqueta_T,
            "D_21in_m": D_helice_comercial_21in_m,
            "eta_i_implicita_21in (-)": eta_21,
            "vi_implicita_21in_m_s": vi_21,
        })
        filas_helices_resultados.append(fila)


df_helices_resultados = pd.DataFrame(filas_helices_resultados)
archivo_helices_resultados = output_dir_fuselaje / "dimensionado_helices_TCM_por_caso.csv"
df_helices_resultados.round(5).to_csv(archivo_helices_resultados, index=False, encoding="utf-8-sig")

# Columnas resumen en df_resultados para poder exportar/usar posteriormente.
df_resultados["N_motores_TCM_base (-)"] = N_motores_TCM_base
df_resultados["eta_i_helice_TCM (-)"] = eta_i_helice_TCM
df_resultados["T_total_helice_diseno_OpenVSP_N"] = T_total_helice_diseno_N
df_resultados["D_helice_OpenVSP_TCM_m"] = D_helice_openvsp_TCM_m
df_resultados["D_helice_OpenVSP_TCM_in"] = D_helice_openvsp_TCM_in
df_resultados["A_disco_helice_OpenVSP_TCM_m2"] = A_helice_openvsp_TCM_m2
df_resultados["D_helice_comercial_21in_m"] = D_helice_comercial_21in_m

# Tambien se añaden los resultados propios del empuje de cada fila.
map_D_cruc = df_helices_resultados[df_helices_resultados["Punto_T"] == "crucero_20km_base"].set_index(["Caso", "Escenario"])
map_D_max = df_helices_resultados[df_helices_resultados["Punto_T"] == "max_perfil_operativo"].set_index(["Caso", "Escenario"])
idx_df = pd.MultiIndex.from_frame(df_resultados[["Caso", "Escenario"]])
df_resultados["D_helice_TCM_crucero_m"] = map_D_cruc.reindex(idx_df)["D_helice_m"].to_numpy()
df_resultados["D_helice_TCM_max_perfil_m"] = map_D_max.reindex(idx_df)["D_helice_m"].to_numpy()
df_resultados["eta_i_21in_crucero (-)"] = map_D_cruc.reindex(idx_df)["eta_i_implicita_21in (-)"].to_numpy()
df_resultados["eta_i_21in_max_perfil (-)"] = map_D_max.reindex(idx_df)["eta_i_implicita_21in (-)"].to_numpy()

# 5D. Punto 8: masa propulsiva a partir de benchmarks de tecnología
filas_propulsion = []
for _, row in df_resultados.iterrows():
    for tech_name, tech_data in propulsion_technologies.items():
        if not tecnologia_propulsiva_completa(tech_data):
            continue
        filas_propulsion.append(
            evaluar_tecnologia_propulsiva(
                row=row,
                tech_name=tech_name,
                tech_dict=tech_data,
                eta_propulsor_para_sizing_local=eta_propulsor_para_sizing,
            )
        )

df_propulsion_sizing = pd.DataFrame(filas_propulsion)

# 5E. Cierre preliminar de WTO con los 8 bloques previos
df_wto_preliminar = construir_dataframe_wto_preliminar(
    df_resultados=df_resultados,
    df_propulsion_sizing=df_propulsion_sizing,
    fixed_weights_mode_local=wto_fixed_weights_mode,
    mission_days_ref=wto_energy_reference_days,
    mf_prop_assumed_cases_local=mf_prop_assumed_cases,
    fixed_weights_manual_local=fixed_weights_manual,
    mf_prop_ref_local=wto_reference_mf_prop_for_implied_fixed,
    tol_kg_local=wto_secant_tol_kg,
    max_iter_local=wto_secant_max_iter,
)

"""
TABLA NUEVA DE COMPARACIÓN ENTRE ESTRATEGIAS.

QUÉ HACE:
- Para cada caso geométrico/estacional compara el escenario de descenso nocturno
  contra el de crucero constante a 20 km.
- Un delta negativo en energía de batería requerida o masa de batería significa
  que el planeo/descenso nocturno ayuda a reducir el almacenamiento necesario.
"""
comparaciones = []
for nombre_caso in df_resultados["Caso"].unique():
    df_caso = df_resultados[df_resultados["Caso"] == nombre_caso].copy()
    if {"Constante_20km", "DescensoNocturno_20a15km"}.issubset(set(df_caso["Escenario"])):
        fila_const = df_caso[df_caso["Escenario"] == "Constante_20km"].iloc[0]
        fila_planeo = df_caso[df_caso["Escenario"] == "DescensoNocturno_20a15km"].iloc[0]

        comparaciones.append({
            "Caso": nombre_caso,
            "Delta_Balance_almacenamiento_Wh (planeo-constante)": fila_planeo["Balance diario almacenamiento (Wh)"] - fila_const["Balance diario almacenamiento (Wh)"],
            "Delta_E_batt_30d_Wh (planeo-constante)": fila_planeo["E_batt_usable_30dias (Wh)"] - fila_const["E_batt_usable_30dias (Wh)"],
            "Delta_E_batt_60d_Wh (planeo-constante)": fila_planeo["E_batt_usable_60dias (Wh)"] - fila_const["E_batt_usable_60dias (Wh)"],
            "Delta_Masa_min_30d_kg (planeo-constante)": fila_planeo["Masa_min_conservadora_30d (kg)"] - fila_const["Masa_min_conservadora_30d (kg)"],
            "Delta_Masa_min_60d_kg (planeo-constante)": fila_planeo["Masa_min_conservadora_60d (kg)"] - fila_const["Masa_min_conservadora_60d (kg)"],
            "Mejora_autosostenibilidad (1/0)": int(
                (fila_planeo["Autosostenible diario (1/0)"] == 1) and
                (fila_const["Autosostenible diario (1/0)"] == 0)
            ),
        })

df_comparacion = pd.DataFrame(comparaciones)

print("Resumen energético y viabilidad (tabla compacta):\n")
df_console_energetica = construir_tabla_presentacion_energetica(df_resultados)
print(df_console_energetica.to_string(index=False))

print("\nInterpretación del pack mínimo de batería para 30 días:\n")
df_console_batt30 = construir_tabla_presentacion_bateria_30d(df_resultados)
print(df_console_batt30.to_string(index=False))

if len(df_comparacion) > 0:
    print("\nComparación entre estrategias:\n")
    df_console_cmp = construir_tabla_presentacion_comparacion(df_comparacion)
    print(df_console_cmp.to_string(index=False))

conclusiones_cortas = construir_conclusiones_por_caso(df_resultados)
print("\nLectura rápida automática de los resultados:\n")
for texto in conclusiones_cortas:
    print(f"- {texto}")


print("\nEficiencia aerodinámica en el punto de diseño (20 km):\n")
df_console_eff = construir_tabla_presentacion_eficiencia_aero(df_resultados)
print(df_console_eff.to_string(index=False))

print("\nPotencias propulsivas de referencia:\n")
df_console_pot = construir_tabla_presentacion_potencias_referencia(df_resultados)
print(df_console_pot.to_string(index=False))

print("\nResumen por plataforma de las potencias eléctricas clave:\n")
df_console_pot_resumen = construir_resumen_potencias_por_plataforma(df_resultados)
if len(df_console_pot_resumen) > 0:
    print(df_console_pot_resumen.to_string(index=False))

print("\nFracción de masa energética del pack mínimo:\n")
df_console_mf_energy = construir_tabla_presentacion_mf_energy(df_resultados)
print(df_console_mf_energy.to_string(index=False))

print("\nHipótesis provisionales de MF_Struct y MF_Subs frente a MF_Energy:\n")
df_console_mf_struct_subs = construir_tabla_presentacion_mf_struct_subs(df_resultados)
print(df_console_mf_struct_subs.to_string(index=False))

print("\nRatios de sizing del avión (P/W y T/W):\n")
df_console_pw_tw = construir_tabla_presentacion_pw_tw(df_resultados)
print(df_console_pw_tw.to_string(index=False))


print("\n" + "=" * 78)
print("SISTEMA PROPULSIVO COMERCIAL DE REFERENCIA")
print("=" * 78)
print(f"Referencia: {propulsion_ref['nombre']}")
print(f"Masa propulsiva instalada estimada: {m_propulsion_ref_kg:.2f} kg")
print(f"Numero de motores: {propulsion_ref['N_motores']}")
print(f"Bus electrico considerado: {propulsion_ref['V_bus_V']:.1f} V")
print(f"Potencia nominal input motor ref: {propulsion_ref['P_input_nominal_motor_W']:.1f} W")
print(f"Potencia nominal output motor ref: {propulsion_ref['P_output_nominal_motor_W']:.1f} W")
print(f"Corriente continua ESC ref: {propulsion_ref['I_ESC_continua_A']:.1f} A")

# Guardado de la tabla completa para revisarla sin depender de la anchura de consola.
archivo_prop_ref = output_dir_fuselaje / "resumen_propulsion_ref.csv"
df_console_prop_ref.round(4).to_csv(archivo_prop_ref, index=False, encoding="utf-8-sig")
print(f"Tabla completa guardada en CSV: {archivo_prop_ref}")

print("\nResumen legible por caso/escenario:")
for _, r in df_console_prop_ref.iterrows():
    print("\n" + "-" * 78)
    print(f"Caso: {r['Caso']}")
    print(f"Escenario: {r['Escenario']}")
    print("-" * 78)
    print(f"Masa propulsiva total ref:        {r['M_propulsion_ref_kg']:.3f} kg")
    print(f"Fraccion MF_Prop_ref:             {r['MF_Prop_ref (-)']:.4f}")
    print(f"Potencia crucero por motor:       {r['P_elec_crucero_por_motor_ref (W)']:.1f} W")
    print(f"Corriente crucero por motor:      {r['I_crucero_por_motor_ref (A)']:.2f} A")
    print(f"Potencia max perfil por motor:    {r['P_elec_max_por_motor_ref (W)']:.1f} W")
    print(f"Corriente max perfil por motor:   {r['I_max_por_motor_ref (A)']:.2f} A")
    print(f"Corriente nominal input motor:    {r['I_nominal_input_motor_ref (A)']:.2f} A")
    print(f"Margen I_crucero vs I_nominal:    {r['Margen_I_crucero_vs_I_nominal (-)']:.2f}")
    print(f"Margen I_max vs I_nominal:        {r['Margen_I_max_vs_I_nominal (-)']:.2f}")
    print(f"Margen I_crucero vs ESC 60A:      {r['Margen_I_crucero_vs_ESC60A (-)']:.2f}")
    print(f"Margen I_max vs ESC 60A:          {r['Margen_I_max_vs_ESC60A (-)']:.2f}")

print("\n" + "=" * 78)
print("DIMENSIONADO PRELIMINAR DE HELICES POR TCM")
print("=" * 78)
print("Modelo: disco actuador / teoria de cantidad de movimiento")
print(f"Velocidad usada: {v_crucero:.2f} m/s TAS a {h_crucero/1000:.1f} km")
print(f"Densidad ISA local: {rho_crucero:.5f} kg/m3")
print(f"Rendimiento ideal objetivo eta_i = U_inf/U1: {eta_i_helice_TCM:.2f}")
print(f"Diametro helice comercial de referencia 21 in: {D_helice_comercial_21in_m:.3f} m")
print(f"Tabla sensibilidad guardada en CSV: {archivo_helices_sens}")
print(f"Tabla por caso/escenario guardada en CSV: {archivo_helices_resultados}")

print("\nSensibilidad 2 motores vs 4 motores:")
cols_helices_sens_console = [
    "T_total_N",
    "N_motores",
    "T_por_helice_N",
    "A_disco_helice_m2",
    "D_helice_m",
    "D_helice_in",
    "eta_i_implicita_21in (-)",
]
print(df_helices_sensibilidad[cols_helices_sens_console].round(3).to_string(index=False))

print("\nRecomendacion OpenVSP para la configuracion bimotor de esta iteracion:")
print(f"T_total diseno:                  {T_total_helice_diseno_N:.1f} N")
print(f"Numero motores:                  {N_motores_TCM_base:d}")
print(f"Diametro por helice:             {D_helice_openvsp_TCM_m:.3f} m")
print(f"Diametro por helice:             {D_helice_openvsp_TCM_in:.1f} in")
print(f"Area de disco por helice:        {A_helice_openvsp_TCM_m2:.3f} m2")
print("Valor a poner en OpenVSP:         Diameter = " + f"{D_helice_openvsp_TCM_m:.3f}")

print("\nLectura rapida:")
print("- Con 2 motores, 21 in queda pequeno para eta_i~0.80 a 20 km.")
print("- Para T_total=35 N y 2 motores, sale D~0.875 m.")
print("- 4 motores permitirian diametros bastante menores, pero aumentan masa, cableado y drag de gondolas.")

print("\nPunto 8: estimación de masa propulsiva por tecnología:\n")
if df_propulsion_sizing.empty:
    print(
        "Aún no se calcula MF_Prop por benchmark porque propulsion_technologies sigue con NaN.\n"
        "Mientras tanto, el cierre preliminar de WTO usa los tres casos asumidos de MF_Prop\n"
        "(0.05 / 0.10 / 0.15), definidos arriba como sensibilidad conceptual.\n"
        "Cuando tengas benchmark de motor / ESC / hélice, rellena ese diccionario y este\n"
        "bloque empezará a devolver masa propulsiva instalada y MF_Prop calculado."
    )
else:
    df_console_propulsion = construir_tabla_presentacion_propulsion(df_propulsion_sizing)
    print(df_console_propulsion.to_string(index=False))

print("\nPunto 9: cierre preliminar de WTO a partir de los 8 bloques previos:\n")
if df_wto_preliminar.empty:
    print("No se ha podido construir el cierre preliminar de WTO.")
else:
    df_console_wto = construir_tabla_presentacion_wto_preliminar(df_wto_preliminar)
    print(df_console_wto.to_string(index=False))


"""
ATENCIÓN:
- El criterio masa_pack <= MTOW que se exporta en las tablas es solo una criba
  mínima para detectar casos obviamente imposibles.
- No sustituye una iteración completa del MTOW del HAPS con el peso real de
  baterías, estructura, placas, aviónica y resto de subsistemas.
"""

# -----------------------------------------------------------------------------
# 6. GUARDADO
# -----------------------------------------------------------------------------
output_dir = Path(__file__).resolve().parent
"""
EXPORTACIÓN ACTIVADA SOLO PARA EL RESUMEN DOCX COMPACTO.

MOTIVO:
- Ahora mismo lo que más te interesa es un Word limpio y legible para revisar
  resultados y enseñarlos, no una colección de tablas largas en CSV/TSV/HTML.
- Por eso se deja activa únicamente la exportación DOCX del resumen compacto.
- Si más adelante quieres reactivar CSV/TSV/HTML, esta zona sigue siendo el sitio
  natural para hacerlo.
"""

tabla_docx = output_dir / "tabla_resumen_4casos_haps_baterias_planeo.docx"

if DOCX_AVAILABLE:
    exportar_resumen_compacto_word(
        ruta_docx=tabla_docx,
        titulo="Resumen comparativo HAPS con baterías y planeo",
        subtitulo=(
            f"Latitud {lat_deg:.1f}°N | Altitud alta {h_crucero/1000.0:.1f} km | "
            f"Altitud baja 15.0 km | Cobertura ala {f_cobertura_max*100:.0f}% | "
            f"eta_prop {eta_prop_global:.2f} | f_usable_batt {f_usable_batt:.2f}"
        ),
        df_energetica=df_console_energetica,
        df_bateria_30d=df_console_batt30,
        df_comparacion_presentacion=construir_tabla_presentacion_comparacion(df_comparacion),
        conclusiones_cortas=conclusiones_cortas,
    )


# -----------------------------------------------------------------------------
# 7. PERFILES AUXILIARES PARA GRÁFICAS
# -----------------------------------------------------------------------------


"""
BLOQUE NUEVO: perfiles auxiliares para gráficas de SOC y deuda multidiaria.

QUÉ HACE:
- Construye perfiles de deuda energética a partir del instante de batería llena.
- Usa el pack mínimo seleccionado para 30 días como referencia para traducir
  deuda -> SOC total del pack.
- Esto sirve para visualizar mejor:
    * la profundidad diaria de descarga,
    * y la deriva descendente en invierno cuando el día no cierra.
"""

perfiles_soc_30d = {}
perfiles_deuda_30d = {}

for _, row in df_resultados.iterrows():
    nombre_perfil = f"{row['Caso']}__{row['Escenario']}"
    df_batt_ref = perfiles_bateria[nombre_perfil]

    hist_1d = construir_historial_deuda_desde_bateria_llena(df_batt_ref, 1)
    hist_30d = construir_historial_deuda_desde_bateria_llena(df_batt_ref, 30)

    e_nom_pack_30d_wh = float(row["E_nominal_pack_30d (Wh)"])

    hist_1d["SOC_total_pack_30d_pct"] = 100.0 * (1.0 - hist_1d["deuda_wh"] / e_nom_pack_30d_wh)
    hist_30d["SOC_total_pack_30d_pct"] = 100.0 * (1.0 - hist_30d["deuda_wh"] / e_nom_pack_30d_wh)

    perfiles_soc_30d[nombre_perfil] = hist_1d
    perfiles_deuda_30d[nombre_perfil] = hist_30d


# -----------------------------------------------------------------------------
# 8. GRÁFICAS
# -----------------------------------------------------------------------------

"""
GRÁFICA 1:
- Potencia eléctrica diaria del array para cada caso geométrico/estacional.
- Sirve para ver rápidamente la diferencia fuerte entre equinoccio e invierno.
"""
fig, ax = plt.subplots(figsize=(10, 6))
for nombre_caso, df_caso in perfiles_solares.items():
    ax.plot(
        df_caso["Tiempo solar (h)"],
        df_caso["Potencia electrica del array (W)"],
        linewidth=2,
        label=etiqueta_caso_breve(nombre_caso),
    )
ax.set_title("Potencia eléctrica del array")
ax.set_xlabel("Tiempo solar local (h)")
ax.set_ylabel("Potencia eléctrica del array (W)")
ax.grid(True)
ax.legend()
guardar_figura_y_mostrar_si_procede(fig, output_dir / "potencia_array_haps.png")


"""
GRÁFICA 2:
- Potencia propulsiva eléctrica en el caso 32 m invierno.
- Si ese caso no está activo en casos_haps, se omite automáticamente.
"""
fig, ax = plt.subplots(figsize=(10, 6))
hay_datos_grafica2 = False
for nombre_perfil, df_perfil in perfiles_operacion.items():
    if "32m_invierno" in nombre_perfil:
        caso_ref, escenario_ref = nombre_perfil.split("__")
        ax.plot(
            df_perfil["Tiempo solar (h)"],
            df_perfil["P_elec_propulsion (W)"],
            linewidth=2,
            label=etiqueta_escenario_corta(escenario_ref),
        )
        hay_datos_grafica2 = True

if hay_datos_grafica2:
    ax.set_title("Potencia propulsiva eléctrica | 32 m invierno")
    ax.set_xlabel("Tiempo solar local (h)")
    ax.set_ylabel("Potencia propulsiva eléctrica (W)")
    ax.grid(True)
    ax.legend()
    guardar_figura_y_mostrar_si_procede(fig, output_dir / "potencia_propulsion_32m_invierno.png")
else:
    plt.close(fig)
    print("Grafica 2 omitida: no hay perfiles activos de 32 m invierno.")


"""
GRÁFICA 3:
- Comparación entre potencia disponible del array y potencia total demandada
  por el avión en 32 m invierno.
- Si el caso 32 m invierno no está activo, se omite automáticamente.
"""
caso_inv_32 = "Caso_2.2_32m_invierno"
if caso_inv_32 in perfiles_solares:
    df_solar_32_inv = perfiles_solares[caso_inv_32]

    fig, ax = plt.subplots(figsize=(10, 6))
    ax.plot(
        df_solar_32_inv["Tiempo solar (h)"],
        df_solar_32_inv["Potencia electrica del array (W)"],
        linewidth=2,
        label="Array solar",
    )

    hay_demanda_grafica3 = False
    for escenario_ref in ["Constante_20km", "DescensoNocturno_20a15km"]:
        nombre_perfil = f"{caso_inv_32}__{escenario_ref}"
        if nombre_perfil not in perfiles_operacion:
            continue
        df_perfil = perfiles_operacion[nombre_perfil]
        ax.plot(
            df_perfil["Tiempo solar (h)"],
            df_perfil["P_total_bus (W)"],
            linewidth=2,
            label=f"Demanda total | {etiqueta_escenario_corta(escenario_ref)}",
        )
        hay_demanda_grafica3 = True

    if hay_demanda_grafica3:
        ax.set_title("Array vs demanda total | 32 m invierno")
        ax.set_xlabel("Tiempo solar local (h)")
        ax.set_ylabel("Potencia en el bus (W)")
        ax.grid(True)
        ax.legend()
        guardar_figura_y_mostrar_si_procede(fig, output_dir / "balance_potencia_32m_invierno.png")
    else:
        plt.close(fig)
        print("Grafica 3 omitida: el caso 32 m invierno no tiene perfiles operativos activos.")
else:
    print("Grafica 3 omitida: el caso 32 m invierno no esta activo en casos_haps.")


"""
GRÁFICA 4:
- SOC total del pack mínimo seleccionado para 30 días, a lo largo de un día
  representativo y tomando t = 0 en el instante de batería llena.
- Ayuda a visualizar la profundidad diaria de descarga con el pack realmente
  elegido para 30 días.
"""
fig, ax = plt.subplots(figsize=(10, 6))
for _, row in df_resultados.iterrows():
    nombre_perfil = f"{row['Caso']}__{row['Escenario']}"
    hist = perfiles_soc_30d[nombre_perfil]
    ax.plot(
        hist["t_rel_h"],
        hist["SOC_total_pack_30d_pct"],
        linewidth=2,
        label=etiqueta_caso_breve(row["Caso"], row["Escenario"]),
    )
ax.set_title("SOC diario del pack mínimo de 30 días")
ax.set_xlabel("Tiempo desde batería llena (h)")
ax.set_ylabel("SOC total del pack (%)")
ax.set_xlim(0, 24)
ax.grid(True)
ax.legend(fontsize=8)
guardar_figura_y_mostrar_si_procede(fig, output_dir / "soc_diario_pack_30d.png")


"""
GRÁFICA 5:
- Deriva de deuda energética durante 30 días repetidos para los casos de invierno.
- Si no hay casos de invierno activos, se omite automáticamente.
"""
fig, ax = plt.subplots(figsize=(10, 6))
hay_datos_grafica5 = False
for _, row in df_resultados.iterrows():
    if "invierno" not in row["Caso"].lower():
        continue
    nombre_perfil = f"{row['Caso']}__{row['Escenario']}"
    if nombre_perfil not in perfiles_deuda_30d:
        continue
    hist = perfiles_deuda_30d[nombre_perfil]
    ax.plot(
        hist["t_rel_h"] / 24.0,
        hist["deuda_wh"] / 1000.0,
        linewidth=2,
        label=etiqueta_caso_breve(row["Caso"], row["Escenario"]),
    )
    hay_datos_grafica5 = True

if hay_datos_grafica5:
    ax.set_title("Deuda energética acumulada | invierno | 30 días")
    ax.set_xlabel("Días desde batería llena")
    ax.set_ylabel("Deuda energética interna (kWh)")
    ax.grid(True)
    ax.legend(fontsize=8)
    guardar_figura_y_mostrar_si_procede(fig, output_dir / "deuda_30dias_invierno.png")
else:
    plt.close(fig)
    print("Grafica 5 omitida: no hay casos de invierno activos.")


"""
GRÁFICA 6:
- Comparación de masas mínimas de batería para 30 y 60 días.
- Muy útil para detectar de un vistazo si el caso es estable (30d≈60d) o si
  arrastra déficit acumulado (60d >> 30d).
"""
labels = [etiqueta_caso_breve(row["Caso"], row["Escenario"]) for _, row in df_resultados.iterrows()]
x = list(range(len(labels)))
masas_30 = [float(v) for v in df_resultados["Masa_min_conservadora_30d (kg)"]]
masas_60 = [float(v) for v in df_resultados["Masa_min_conservadora_60d (kg)"]]
bar_w = 0.38

fig, ax = plt.subplots(figsize=(12, 6))
ax.bar([xi - bar_w / 2 for xi in x], masas_30, width=bar_w, label="30 días")
ax.bar([xi + bar_w / 2 for xi in x], masas_60, width=bar_w, label="60 días")
ax.set_title("Masa mínima del pack según la tecnología")
ax.set_ylabel("Masa de batería (kg)")
ax.set_xticks(x)
ax.set_xticklabels(labels, rotation=25, ha="right")
ax.grid(True, axis="y")
ax.legend()
guardar_figura_y_mostrar_si_procede(fig, output_dir / "masa_bateria_30d_60d.png")

"""
GRÁFICA ADICIONAL:
- Potencia neta en el bus para la configuración principal de 24 m en invierno.
- Compara crucero constante y descenso nocturno.
- P_neta > 0: existe excedente para cargar la batería.
- P_neta < 0: la batería debe alimentar el bus.
"""
caso_ref = "Caso_1.2_24m_invierno"

etiquetas_estrategia = {
    "Constante_20km": "Crucero constante a 20 km",
    "DescensoNocturno_20a15km": "Descenso nocturno 20→15 km",
}

fig, ax = plt.subplots(figsize=(10, 6))
hay_datos = False

for escenario_ref, etiqueta in etiquetas_estrategia.items():
    nombre_perfil = f"{caso_ref}__{escenario_ref}"

    if nombre_perfil not in perfiles_bateria:
        continue

    df_batt_ref = perfiles_bateria[nombre_perfil]

    ax.plot(
        df_batt_ref["t_mid_h"],
        df_batt_ref["P_neta_bus_W"],
        linewidth=2,
        label=etiqueta,
    )
    hay_datos = True

if hay_datos:
    ax.axhline(
    0.0,
    color="black",
    linestyle="--",
    linewidth=1.0,
    label="_nolegend_",
)
    ax.set_title("Potencia neta en el bus | 24 m | invierno")
    ax.set_xlabel("Tiempo solar local (h)")
    ax.set_ylabel("Potencia neta en el bus (W)")
    ax.set_xlim(0, 24)
    ax.grid(True)
    ax.legend()

    guardar_figura_y_mostrar_si_procede(
        fig,
        output_dir / "potencia_neta_24m_invierno.png",
    )
else:
    plt.close(fig)
    print("Gráfica de potencia neta omitida: no se han encontrado los perfiles.")
    
"""
GRÁFICA ADICIONAL:
- Comparación de masa de batería por tecnología.
- Caso: 24 m, equinoccio, descenso nocturno y misión de 30 días.
"""
df_tec = df_baterias[
    (df_baterias["Caso"] == "Caso_1.1_24m_equinoccio")
    & (
        df_baterias["Escenario"]
        == "DescensoNocturno_20a15km"
    )
    & (df_baterias["Dias_mision"] == 30)
].copy()

if not df_tec.empty:
    df_tec = df_tec.sort_values("Tecnologia_bateria").reset_index(drop=True)

    tecnologias = df_tec["Tecnologia_bateria"].tolist()
    masa_conservadora = df_tec["Masa_total_conservadora_kg"].to_numpy()
    masa_optimista = df_tec["Masa_total_optimista_kg"].to_numpy()

    x = np.arange(len(tecnologias))
    ancho = 0.36

    fig, ax = plt.subplots(figsize=(8, 5.5))

    barras_cons = ax.bar(
        x - ancho / 2,
        masa_conservadora,
        width=ancho,
        label="Estimación conservadora",
    )
    barras_opt = ax.bar(
        x + ancho / 2,
        masa_optimista,
        width=ancho,
        label="Estimación optimista",
    )

    ax.set_title("Masa de batería según la tecnología")
    ax.set_ylabel("Masa estimada del pack (kg)")
    ax.set_xticks(x)
    ax.set_xticklabels(tecnologias)
    ax.grid(True, axis="y")
    ax.legend()

    for barras in (barras_cons, barras_opt):
        for barra in barras:
            valor = barra.get_height()
            ax.text(
                barra.get_x() + barra.get_width() / 2,
                valor,
                f"{valor:.1f}".replace(".", ","),
                ha="center",
                va="bottom",
                fontsize=9,
            )

    guardar_figura_y_mostrar_si_procede(
        fig,
        output_dir / "masa_bateria_por_tecnologia.png",
    )
else:
    print("Gráfica por tecnología omitida: no se encontró el caso seleccionado.")
    
"""
GRÁFICA 4:
- Comparación del SOC diario entre un caso viable de equinoccio y el caso
  invernal equivalente.
- Ambos corresponden a la configuración de 24 m, 85 kg y descenso nocturno.
- Cada curva se normaliza respecto a la capacidad de su propio pack mínimo
  para una misión de 30 días.
"""
casos_soc = [
    (
        "Caso_1.1_24m_equinoccio",
        "DescensoNocturno_20a15km",
        "Equinoccio",
    ),
    (
        "Caso_1.2_24m_invierno",
        "DescensoNocturno_20a15km",
        "Invierno",
    ),
]

fig, ax = plt.subplots(figsize=(10, 6))
hay_datos_soc = False

for caso_ref, escenario_ref, etiqueta in casos_soc:
    nombre_perfil = f"{caso_ref}__{escenario_ref}"

    if nombre_perfil not in perfiles_soc_30d:
        continue

    hist = perfiles_soc_30d[nombre_perfil]

    ax.plot(
        hist["t_rel_h"],
        hist["SOC_total_pack_30d_pct"],
        linewidth=2,
        label=etiqueta,
    )
    hay_datos_soc = True

if hay_datos_soc:
    ax.axhline(
        20.0,
        linestyle="--",
        linewidth=1.0,
        label="SOC mínimo adoptado",
    )

    ax.set_title("SOC diario del pack mínimo de 30 días | 24 m")
    ax.set_xlabel("Tiempo desde batería llena (h)")
    ax.set_ylabel("SOC del pack (%)")
    ax.set_xlim(0, 24)
    ax.set_ylim(15, 102)
    ax.grid(True)
    ax.legend()

    guardar_figura_y_mostrar_si_procede(
        fig,
        output_dir / "soc_24m_equinoccio_invierno.png",
    )
else:
    plt.close(fig)
    print("Gráfica de SOC omitida: no se encontraron los perfiles seleccionados.")

# -----------------------------------------------------------------------------
# ÍNDICE RÁPIDO DEL SCRIPT (líneas aproximadas; actualizar si desplazas bloques)
# -----------------------------------------------------------------------------
# 0. IMPORT OPCIONAL PARA WORD .............................. línea 20
# 1. CONSTANTES GLOBALES .................................... línea 34
# 1A. Hipótesis potencia no propulsiva ...................... línea 80
# 1B. Hipótesis payload / baterías .......................... línea 106
# 1C. Tecnologías de batería ................................ línea 162
# 1D. Hipótesis propulsión / sizing ......................... línea 229
# 1E. Hipótesis cierre preliminar de WTO .................... línea 255
# 2. CASOS DE ESTUDIO ....................................... línea 335
# 3. FUNCIONES AUXILIARES ................................... línea 408
# 4. CÁLCULO DE CADA CASO Y ESCENARIO ....................... línea 2304
# 5. TABLAS, POSTPROCESADO Y SIZING INICIAL ................. línea 2683
# 5A bis. Dimensionado fuselaje + cola en T ................. línea 2689
# 5A ter. Sistema fotovoltaico en ala ....................... línea 3081
# 5A quater. Volumen interno ala / baterías ................. línea 3232
# 5A. MF_Energy ............................................. línea 3610
# 5B. MF_Struct / MF_Subs ................................... línea 3621
# 5C. Ratios P/W y T/W ...................................... línea 3638
# 5C bis. Sistema propulsivo comercial de referencia ........ línea 3658
# 5C ter. Dimensionado hélices por TCM ...................... línea 3742
# 5D. MF_Prop desde benchmark tecnológico ................... línea 3935
# 5E. Cierre preliminar de WTO .............................. línea 3952
# 5F. Tablas de consola y conclusiones generales ............ línea 3994
# 6. GUARDADO ............................................... línea 4140
# 7. PERFILES AUXILIARES PARA GRÁFICAS ...................... línea 4173
# 8. GRÁFICAS ............................................... línea 4209
# -----------------------------------------------------------------------------
